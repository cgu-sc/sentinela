import io
import os
import re
import uuid
from decimal import Decimal
from datetime import date, datetime
from pathlib import Path
import time
from typing import Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from data_cache import get_localidades_df

from .farmacia import get_dados_farmacia
from .dashboard import get_dashboard_data
from .crm import get_crm_data
from .socios import get_socios_farmacia
from .nota_tecnica_charts import (
    _add_figura_evolucao_financeira,
    _add_figura_percentil_risco,
    _add_figura_posicionamento_regional,
)
from .nota_tecnica_anexo_ii import _add_anexo_ii_memoria_calculo, _build_anexo_ii_context
from .nota_tecnica_anexos import _add_anexo_crm_evidencias, _add_anexo_falecidos
from .nota_tecnica_contexts import (
    _build_esocial_context,
    _build_evolucao_financeira_context,
    _build_gtin_sem_comprovacao_context,
    _build_percentil_risco_context,
    _build_posicionamento_regional_context,
    _build_regional_comparison_context,
    _build_socios_volume_atipico_context,
    _build_ultimo_mes_sav_context,
)
from .nota_tecnica_quadros import (
    _add_quadro_53,
    _add_quadro_comparativo_regional,
    _add_quadro_evolucao_financeira,
    _add_quadro_identificacao,
    _add_quadro_socios_volume_atipico,
    _add_tabela_gtins_sem_comprovacao,
    _add_tabela_medicamentos_aumento_atipico,
)
from .nota_tecnica_criticidades import (
    _SECAO5_MAP,
    _add_alto_custo_text,
    _add_dias_pico_text,
    _add_dispersao_geografica_text,
    _add_falecidos_criticidade_text,
    _add_indicador_regional_table,
    _add_indicadores_criticos_quadro,
    _add_incompatibilidade_patologica_text,
    _add_per_capita_text,
    _add_polimedicamento_text,
    _add_receita_paciente_text,
    _add_recorrencia_sistemica_text,
    _add_teto_text,
    _add_ticket_medio_text,
    _add_vendas_rapidas_text,
    _build_alto_custo_context,
    _build_dias_pico_context,
    _build_dispersao_geografica_context,
    _build_falecidos_context,
    _build_indicador_regional_context,
    _build_indicadores_criticos_quadro,
    _build_incompatibilidade_patologica_context,
    _build_per_capita_context,
    _build_polimedicamento_context,
    _build_receita_paciente_context,
    _build_recorrencia_sistemica_context,
    _build_teto_context,
    _build_ticket_medio_context,
    _build_vendas_rapidas_context,
    _count_incompatibilidade_patologica_tables,
    _get_criticos,
    _get_criticos_ordenados_por_risco,
)
from .nota_tecnica_crm import (
    _add_crms_irregulares_text,
    _add_hhi_crm_text,
    _build_crm_evidencias_complementares_context,
    _build_crms_irregulares_context,
    _build_hhi_crm_context,
)
from .nota_tecnica_esocial import _add_esocial_context_text
from .nota_tecnica_regionais import resolve_nota_tecnica_regional
from .nota_tecnica_formatters import (
    _format_decimal_pt,
    _format_full_date_long_pt,
    _format_list_pt,
)
from .nota_tecnica_docx_utils import (
    _add_bookmark,
    _cell_bg,
    _cell_bg_run,
    _cell_borders,
    _footnote_ref,
    _run,
    _set_table_fixed_widths,
    _tbl_no_borders,
)


def _project_root() -> Path:
    return Path(__file__).resolve().parents[4]


def _format_timing_ms(ms: float) -> str:
    return f"{ms / 1000:.3f}s" if ms >= 1000 else f"{ms:.1f}ms"


def _nota_tecnica_timing_enabled() -> bool:
    value = os.getenv("NOTA_TECNICA_TIMING_LOG")
    if value is None:
        env_path = _project_root() / ".env"
        if env_path.exists():
            for line in env_path.read_text(encoding="utf-8").splitlines():
                text = line.strip()
                if not text or text.startswith("#") or "=" not in text:
                    continue
                key, raw_value = text.split("=", 1)
                if key.strip() == "NOTA_TECNICA_TIMING_LOG":
                    value = raw_value.strip().strip('"').strip("'")
                    break
    return (value or "").strip().lower() in {"1", "true", "yes", "on"}


class _NotaTecnicaTiming:
    def __init__(self, cnpj: str, data_inicio: Optional[date], data_fim: Optional[date]):
        self.cnpj = cnpj
        self.data_inicio = data_inicio
        self.data_fim = data_fim
        self.started_at = datetime.now()
        self._start = time.perf_counter()
        self._last = self._start
        self.steps: list[tuple[str, float, float]] = []

    def mark(self, label: str):
        now = time.perf_counter()
        self.steps.append((label, (now - self._last) * 1000, (now - self._start) * 1000))
        self._last = now

    def write(self, status: str = "OK", error: str | None = None):
        total_ms = (time.perf_counter() - self._start) * 1000
        periodo = (
            f"{self.data_inicio or 'inicio-aberto'} a {self.data_fim or 'fim-aberto'}"
            if self.data_inicio or self.data_fim
            else "historico completo"
        )
        log_path = _project_root() / "logs" / "nota_tecnica_timing.log"
        log_path.parent.mkdir(parents=True, exist_ok=True)

        lines = [
            "=" * 88,
            f"{self.started_at:%Y-%m-%d %H:%M:%S} | Nota Tecnica | CNPJ {self.cnpj} | periodo {periodo} | status {status}",
            f"TOTAL: {_format_timing_ms(total_ms)}",
            "Etapas:",
        ]
        lines.extend(
            f"  - {label:<52} {_format_timing_ms(step_ms):>9} | acumulado {_format_timing_ms(total_step_ms)}"
            for label, step_ms, total_step_ms in self.steps
        )
        if error:
            lines.append(f"Erro: {error}")
        lines.append("")

        with log_path.open("a", encoding="utf-8") as fp:
            fp.write("\n".join(lines) + "\n")


def _risk_color(classificacao: str | None, score: float) -> tuple[str, str]:
    """Retorna (hex_6chars, label) baseado na classificação de risco do sistema."""
    c = (classificacao or '').upper()
    if 'CRÍTICO' in c or 'CRITICO' in c or 'ALTO' in c:
        return 'F87171', 'CRÍTICO'
    if 'MÉDIO' in c or 'MEDIO' in c or 'ATENÇÃO' in c or 'ATENCAO' in c:
        return 'F97316', 'ATENÇÃO'
    if 'BAIXO' in c or 'NORMAL' in c:
        return '334155', 'NORMAL'
    if score > 20:
        return 'F87171', 'CRÍTICO'
    if score > 10:
        return 'F97316', 'ATENÇÃO'
    return '334155', 'NORMAL'


def _vez_ou_vezes(value: float) -> str:
    return "vez" if abs(value) <= 1 else "vezes"


def _build_codigo_verificacao(cnpj: str, generated_at: datetime) -> str:
    cnpj_digits = ''.join(ch for ch in str(cnpj or '') if ch.isdigit()) or 'SEM-CNPJ'
    suffix = uuid.uuid4().hex[:8].upper()
    return f'NT-{cnpj_digits}-{generated_at:%Y%m%d}-{suffix}'


def _resolve_numero_nota_input(numero_nota: Optional[str]) -> tuple[str, bool]:
    text = str(numero_nota or "").strip()
    if not text:
        return "XXX", True
    if not text.isdigit():
        raise ValueError("Número da Nota Técnica deve conter apenas dígitos.")
    return text, False


def _resolve_numero_processo_input(numero_processo: Optional[str], ano_nota: int) -> tuple[str, bool]:
    text = str(numero_processo or "").strip()
    if not text:
        return f"00XXX.XXXXXX/{ano_nota}-XX", True

    if text.isdigit() and len(text) == 17:
        digits = text
    elif re.fullmatch(r"\d{5}\.\d{6}/\d{4}-\d{2}", text):
        digits = "".join(ch for ch in text if ch.isdigit())
    else:
        raise ValueError("Número do processo deve conter 17 dígitos ou estar no padrão 00000.000000/0000-00.")

    return f"{digits[:5]}.{digits[5:11]}/{digits[11:15]}-{digits[15:]}", False


def _resolve_assinantes_tecnicos(
    assinantes_tecnicos: Optional[list[dict[str, Any]]],
) -> list[dict[str, str]]:
    if assinantes_tecnicos is None:
        return [
            {"nome": "Fulano de Tal", "cargo": "Cargo"},
            {"nome": "Cicrano de Tal", "cargo": "Cargo"},
        ]
    if not isinstance(assinantes_tecnicos, list):
        raise ValueError("assinantes_tecnicos deve ser uma lista.")
    if len(assinantes_tecnicos) > 3:
        raise ValueError("Informe no maximo 3 assinaturas tecnicas.")

    normalized: list[dict[str, str]] = []
    for item in assinantes_tecnicos:
        if not isinstance(item, dict):
            raise ValueError("Cada assinatura tecnica deve ser um objeto.")
        nome = str(item.get("nome", "")).strip()
        cargo = str(item.get("cargo", "")).strip()
        if (nome and not cargo) or (cargo and not nome):
            raise ValueError("Cada assinatura tecnica deve conter nome e cargo.")
        if nome and cargo:
            normalized.append({"nome": nome, "cargo": cargo})

    return normalized or [
        {"nome": "Fulano de Tal", "cargo": "Cargo"},
        {"nome": "Cicrano de Tal", "cargo": "Cargo"},
    ]


def _add_sumario_official_header(
    doc,
    brasao_path: str,
    regional: dict[str, str],
    ano_nota: int,
    numero_nota: str,
    numero_nota_placeholder: bool,
    numero_processo: str,
    numero_processo_placeholder: bool,
):
    if not os.path.exists(brasao_path):
        raise FileNotFoundError(f'Brasao da CGU nao encontrado em {brasao_path}')

    placeholder_color = 'DC2626'
    default_color = '0F172A'

    p_brasao = doc.add_paragraph()
    p_brasao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brasao.paragraph_format.space_after = Pt(4)
    p_brasao.add_run().add_picture(brasao_path, width=Inches(1.15))

    def add_centered_line(parts: list[tuple[str, str, bool]], *, size: float = 9, space_after: float = 0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(space_after)
        for text, color, bold in parts:
            _run(p, text, color=color, size=size, bold=bold)

    add_centered_line([('CONTROLADORIA-GERAL DA UNIÃO', default_color, True)], size=14, space_after=1)
    add_centered_line([(regional["nome_unidade"], default_color, True)], size=9, space_after=1)
    add_centered_line([(regional["linha_endereco"], default_color, False)], size=8, space_after=1)
    add_centered_line([(regional["linha_contato"], default_color, False)], size=8, space_after=8)
    add_centered_line(
        [
            ('NOTA TÉCNICA Nº ', default_color, True),
            (numero_nota, placeholder_color if numero_nota_placeholder else default_color, True),
            (f'/{ano_nota}/NAE/{regional["codigo"]}/Regional/{regional["codigo"]}', default_color, True),
        ],
        size=10,
        space_after=1,
    )
    add_centered_line(
        [
            ('(PROCESSO Nº ', default_color, True),
            (numero_processo, placeholder_color if numero_processo_placeholder else default_color, True),
            (')', default_color, True),
        ],
        size=9,
        space_after=10,
    )


def _apply_codigo_verificacao_footer(doc, codigo_verificacao: str):
    footer_text = f'Sentinela • Código de verificação: {codigo_verificacao}'
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = ''
        _run(footer, footer_text, color='94A3B8', size=7)


def _configure_section(section, footer_lines: list[str] | None = None):
    """Configura margens e rodape independente para uma secao."""
    section.footer.is_linked_to_previous = False
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.text = ''
    for idx, line in enumerate(footer_lines or []):
        if idx > 0:
            footer.add_run('\n')
        _run(footer, line, color='64748B', size=8)


def _start_section(doc, *, footer_lines: list[str] | None = None, start=WD_SECTION.CONTINUOUS):
    """Inicia uma secao com rodape proprio."""
    section = doc.add_section(start)
    _configure_section(section, footer_lines)
    return section


def _format_main_heading(heading):
    """Padroniza o espacamento dos titulos principais da Nota Tecnica."""
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True
    return heading


def _add_confidential_watermark(section):
    """Adiciona marca d'agua discreta no cabecalho da secao."""
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    watermark_xml = """
    <w:r
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office"
        xmlns:w10="urn:schemas-microsoft-com:office:word">
      <w:pict>
        <v:shape id="sentinela_confidential_watermark"
          o:spid="_x0000_s1025"
          type="#_x0000_t136"
          style="position:absolute;margin-left:0;margin-top:0;width:650pt;height:150pt;rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
          fillcolor="#D9D9D9"
          stroked="f">
          <v:fill opacity="0.35"/>
          <v:textpath
            style="font-family:Arial;font-size:64pt;font-weight:bold"
            string="ACESSO RESTRITO"/>
          <w10:wrap type="none"/>
        </v:shape>
      </w:pict>
    </w:r>
    """
    paragraph._p.append(parse_xml(watermark_xml))


def _valor_estimado_por_percentual(total_mov: float, percentual: float) -> float:
    return total_mov * percentual / 100 if total_mov > 0 and percentual > 0 else 0.0


def _normalize_id_ibge7(value: Any) -> str:
    text = str(value or "").strip()
    if "." in text:
        text = text.split(".", 1)[0]
    return text.zfill(7) if text.isdigit() else text


def _resolve_unidade_pf(cadastro: dict[str, Any]) -> str:
    id_ibge7 = _normalize_id_ibge7(cadastro.get("id_ibge7"))
    if not id_ibge7:
        return "Delegacia de Polícia Federal competente"

    try:
        df_loc = get_localidades_df()
    except Exception as exc:
        print(f"[NOTA_TECNICA] Jurisdicao PF indisponivel para id_ibge7 {id_ibge7}: {exc}")
        return "Delegacia de Polícia Federal competente"

    if not {"id_ibge7", "unidade_pf"}.issubset(set(df_loc.columns)):
        return "Delegacia de Polícia Federal competente"
    for row in df_loc.iter_rows(named=True):
        if _normalize_id_ibge7(row.get("id_ibge7")) == id_ibge7:
            unidade_pf = str(row.get("unidade_pf") or "").strip(" .")
            return unidade_pf or "Delegacia de Polícia Federal competente"
    return "Delegacia de Polícia Federal competente"


def _add_resumo_criticidades_conclusao(doc, resumos: list[str]):
    if not resumos:
        return

    p_intro = doc.add_paragraph()
    _run(
        p_intro,
        'Além disso, foram identificadas, para o mesmo período, as seguintes criticidades que corroboram aquele achado principal:',
        color='0F172A',
        size=10,
    )
    for resumo in resumos:
        p_item = doc.add_paragraph(style='List Bullet')
        _run(p_item, resumo, color='0F172A', size=10)


def _build_resumo_falecidos(num: str, falecidos_comp: dict[str, Any]) -> str:
    return (
        f'[Subitem {num}]: Registros, {falecidos_comp["periodo_desc"]}, de '
        f'{falecidos_comp["total_autorizacoes"]:,}'.replace(',', '.')
        + ' vendas de medicamentos em data igual e/ou posterior ao registro de morte de '
        + f'{falecidos_comp["cpfs_distintos"]:,}'.replace(',', '.')
        + ' beneficiários. Estas vendas representaram um valor total de '
        + f'R$ {_format_decimal_pt(falecidos_comp["valor_total"], 2)};'
    )


def _build_resumo_criticidade(num: str, key: str, comp: dict[str, Any], total_mov: float) -> str | None:
    percentual = float(comp.get("percentual") or 0.0)
    multiplicador = float(comp.get("multiplicador_regiao") or 0.0)

    if key == "incompatibilidade_patologica":
        ranking_patologias = comp.get("ranking_patologias")
        if not ranking_patologias:
            return (
                f'[Subitem {num}]: Vendas de medicamentos com incompatibilidade patológica com percentual de '
                f'{_format_decimal_pt(percentual, 2)}% das vendas monitoradas pelo indicador, superior em '
                f'{_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região;'
            )

        valor_identificado = 0.0
        for item in ranking_patologias:
            if "valor_incompativel_pago" not in item or item["valor_incompativel_pago"] is None:
                raise RuntimeError("Valor clinico identificado obrigatorio ausente para resumo da Nota Tecnica.")
            try:
                valor_identificado += float(item["valor_incompativel_pago"])
            except (TypeError, ValueError) as exc:
                raise RuntimeError(
                    f"Valor clinico identificado invalido para resumo da Nota Tecnica: {item['valor_incompativel_pago']}"
                ) from exc

        return (
            f'[Subitem {num}]: Vendas de medicamentos com incompatibilidade patológica com percentual de '
            f'{_format_decimal_pt(percentual, 2)}% das vendas monitoradas pelo indicador, superior em '
            f'{_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região. '
            f'Estas vendas representaram um valor total identificado de R$ {_format_decimal_pt(valor_identificado, 2)};'
        )

    percentuais_templates = {
        "teto": "Vendas correspondentes ao limite máximo de retirada mensal de medicamento por cliente",
        "polimedicamento": "Vendas correspondentes a quatro ou mais itens de medicamentos por cupom",
        "alto_custo": "Vendas de medicamentos de alto custo",
        "vendas_rapidas": "Vendas de medicamentos em tempo inferior a 60 segundos",
        "recorrencia_sistemica": "Vendas de medicamentos com precisão absoluta de 30 dias",
        "dias_pico": "Vendas de medicamentos em dias de pico",
        "dispersao_geografica": "Vendas para pessoas residentes em outros Estados",
    }
    if key == "dispersao_geografica":
        percentual_financeiro = float(comp.get("percentual_financeiro_outra_uf") or 0.0)
        valor_outra_uf = float(comp.get("total_valor_outra_uf") or 0.0)
        return (
            f'[Subitem {num}]: {percentuais_templates[key]} com percentual financeiro de '
            f'{_format_decimal_pt(percentual_financeiro, 2)}% do valor autorizado total da farmácia no período. '
            f'Estas vendas representaram valor autorizado de R$ {_format_decimal_pt(valor_outra_uf, 2)};'
        )
    if key in percentuais_templates:
        valor_estimado = _valor_estimado_por_percentual(total_mov, percentual)
        return (
            f'[Subitem {num}]: {percentuais_templates[key]} com percentual de '
            f'{_format_decimal_pt(percentual, 2)}% de suas vendas totais, superior em '
            f'{_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região. '
            f'Estas vendas representaram um valor total estimado de R$ {_format_decimal_pt(valor_estimado, 2)};'
        )

    if key == "ticket_medio":
        return (
            f'[Subitem {num}]: Valor do ticket médio de medicamentos de R$ {_format_decimal_pt(comp.get("valor") or 0.0, 2)}, '
            f'superior em {_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região;'
        )
    if key == "receita_paciente":
        return (
            f'[Subitem {num}]: Faturamento médio mensal por cliente de R$ {_format_decimal_pt(comp.get("valor") or 0.0, 2)}, '
            f'superior em {_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região;'
        )
    if key == "per_capita":
        return (
            f'[Subitem {num}]: Faturamento mensal per capita de R$ {_format_decimal_pt(comp.get("valor") or 0.0, 2)}, '
            f'superior em {_format_decimal_pt(multiplicador, 2)} vezes à mediana correspondente aos estabelecimentos de sua região;'
        )
    if key == "hhi_crm":
        principal = comp.get("principal") or {}
        crm_ident = str(principal.get("id_medico") or "não informado")
        return (
            f'[Subitem {num}]: Concentração atípica de registros vinculados ao CRM {crm_ident}, com '
            f'{comp.get("principal_autorizacoes") or 0:,}'.replace(',', '.')
            + f' autorizações e valor associado de R$ {_format_decimal_pt(comp.get("principal_valor") or 0.0, 2)}, '
            + f'equivalente a {_format_decimal_pt(comp.get("pct_valor") or 0.0, 2)}% do valor pago pelo PFPB à farmácia no período;'
        )
    if key == "crms_irregulares":
        return (
            f'[Subitem {num}]: Vendas de medicamentos prescritos por médicos com CRMs irregulares ou inválidos, equivalentes a '
            f'{_format_decimal_pt(comp.get("pct_irregular") or 0.0, 2)}% das vendas totais. '
            f'Estas vendas representaram um valor total de R$ {_format_decimal_pt(comp.get("valor_irregular") or 0.0, 2)};'
        )
    return None


def _iter_criticidade_items(
    criticos: set[str],
    razao_social: str,
    *,
    start_index: int = 1,
    exclude_keys: set[str] | None = None,
    ordered_keys: list[str] | None = None,
) -> list[tuple[str, str, str]]:
    """Retorna criticidades criticas com numeracao da secao 7."""
    items: list[tuple[str, str, str]] = []
    exclude_keys = exclude_keys or set()
    titles_by_key = {key: title for key, _, title in _SECAO5_MAP}
    keys = ordered_keys if ordered_keys is not None else [key for key, _, _ in _SECAO5_MAP]
    for key in keys:
        if key in exclude_keys or key not in criticos:
            continue
        title = titles_by_key.get(key)
        if title is None:
            continue
        full_title = title.format(farmacia=razao_social) if '{farmacia}' in title else title
        items.append((key, f'7.{start_index + len(items)}', full_title))
    return items


def _add_toc_entry(doc, num: str, title: str, page: str = 'xx'):
    """Adiciona uma entrada no sumário com tab stop, líder de pontos e recuo para evitar sobreposição."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Recuo à direita para forçar o wrap antes de chegar no número da página
    p.paragraph_format.right_indent = Inches(0.7)
    
    # Recuo deslocado (Hanging Indent) para que a quebra de linha fique elegante
    # Ajustamos conforme o nível (se tem espaço ou não no início)
    indent_base = 0.4 if num.startswith(' ') else 0.2
    p.paragraph_format.left_indent = Inches(indent_base)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    _run(p, f"{num} {title}\t{page}", color='0F172A', size=10)
    p.paragraph_format.space_after = Pt(2)


def _build_sumario(
    doc,
    criticos: set[str],
    razao_social: str,
    cnpj_fmt: str,
    criticidade_order: list[str],
):
    """Constrói a página de sumário dinâmica."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_title, 'SUMÁRIO', color='0F172A', size=14, bold=True)
    doc.add_paragraph()

    _add_toc_entry(doc, '1.', 'ASSUNTO', page='3')
    _add_toc_entry(doc, '2.', 'REFERÊNCIAS', page='3')
    _add_toc_entry(doc, '3.', 'INTRODUÇÃO', page='4')
    _add_toc_entry(doc, '4.', 'SÍNTESE DO PROGRAMA FARMÁCIA POPULAR DO BRASIL E DA METODOLOGIA DESENVOLVIDA PELA CGU PARA SEU MONITORAMENTO', page='5')
    _add_toc_entry(doc, '  4.1', 'Sobre o Programa Farmácia Popular do Brasil', page='5')
    _add_toc_entry(doc, '  4.2', 'Sobre a metodologia desenvolvida pela CGU para apuração de possíveis “vendas sem comprovação”', page='5')

    _add_toc_entry(doc, '5.', f'SOBRE A FARMÁCIA {razao_social} (CNPJ {cnpj_fmt})', page='6')
    _add_toc_entry(doc, '  5.1', f'Informações sobre a Farmácia {razao_social} (CNPJ {cnpj_fmt})', page='6')
    _add_toc_entry(doc, '  5.2', 'Contexto trabalhista e assistência técnica farmacêutica', page='6')
    _add_toc_entry(doc, '6.', f'SOBRE “VENDAS SEM COMPROVAÇÃO” REALIZADAS PELA FARMÁCIA {razao_social}', page='6')
    _add_toc_entry(doc, '  6.1', f'Evolução das transferências do Programa Farmácia Popular do Brasil para a Farmácia {razao_social} e das possíveis “vendas sem comprovação” por ela realizadas', page='6')

    _add_toc_entry(doc, '7.', f'SOBRE OUTRAS CRITICIDADES RELATIVAS À FARMÁCIA {razao_social}, NO ÂMBITO DO PFPB', page='7')
    criticidade_start = 1
    criticidade_items = _iter_criticidade_items(
        criticos,
        razao_social,
        start_index=criticidade_start,
        ordered_keys=criticidade_order,
    )
    for _, num, full_title in criticidade_items:
        _add_toc_entry(doc, f'  {num}', full_title, page='7')

    _add_toc_entry(doc, '8.', 'CONCLUSÃO E ENCAMINHAMENTO', page='8')
    doc.add_page_break()


# ── Geração do documento ─────────────────────────────────────────────────────

def generate_nota_tecnica(
    db,
    cnpj: str,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
    regional_codigo: Optional[str] = None,
    numero_nota: Optional[str] = None,
    numero_processo: Optional[str] = None,
    assinantes_tecnicos: Optional[list[dict[str, Any]]] = None,
):
    """Gera a Nota Técnica Preliminar em formato .docx."""
    timing_log_enabled = _nota_tecnica_timing_enabled()
    timing = _NotaTecnicaTiming(cnpj, data_inicio, data_fim)
    generated_at = datetime.now()
    codigo_verificacao = _build_codigo_verificacao(cnpj, generated_at)
    regional_emissora = resolve_nota_tecnica_regional(regional_codigo)
    numero_nota_base, numero_nota_placeholder = _resolve_numero_nota_input(numero_nota)
    numero_processo_texto, numero_processo_placeholder = _resolve_numero_processo_input(
        numero_processo,
        generated_at.year,
    )
    assinantes_tecnicos_resolvidos = _resolve_assinantes_tecnicos(assinantes_tecnicos)
    numero_nota_tecnica = (
        f'{numero_nota_base}/{generated_at.year}/NAE/{regional_emissora["codigo"]}/Regional/{regional_emissora["codigo"]}'
    )

    # 1. Coleta de dados
    cadastro_obj = get_dados_farmacia(cnpj)
    cadastro = cadastro_obj.model_dump() if cadastro_obj is not None else {}
    timing.mark("dados cadastrais")

    resumo = get_dashboard_data(db, data_inicio, data_fim, cnpjs=[cnpj])
    cnpj_data_obj = resumo.resultado_cnpjs[0] if hasattr(resumo, 'resultado_cnpjs') and resumo.resultado_cnpjs else None
    cnpj_data = cnpj_data_obj.model_dump() if cnpj_data_obj is not None else {}
    timing.mark("dashboard / resumo do CNPJ")

    # Coleta de sócios
    socios_res = get_socios_farmacia(cnpj)
    socios_ativos = [s for s in socios_res.socios if not s.data_exclusao_sociedade]
    timing.mark("socios")

    # 2. Campos derivados
    razao_social = (cadastro.get('razao_social') or cnpj_data.get('razao_social') or 'NÃO INFORMADO').upper()
    nome_fantasia = str(cadastro.get('nome_fantasia') or '').strip()
    if nome_fantasia.upper() in {'', 'NONE', 'NULL', 'NAN'}:
        nome_fantasia = ''
    municipio = cnpj_data.get('municipio') or cadastro.get('municipio') or '—'
    uf = cnpj_data.get('uf') or cadastro.get('uf') or '—'
    unidade_pf = _resolve_unidade_pf(cadastro)
    cnpj_fmt = f'{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}' if len(cnpj) == 14 and cnpj.isdigit() else cnpj

    logradouro = ' '.join(p for p in [cadastro.get('tipo_logradouro') or '', cadastro.get('logradouro') or ''] if p).strip()
    numero = cadastro.get('numero') or ''
    bairro = cadastro.get('bairro') or ''
    cep = cadastro.get('cep') or ''
    endereco_parts = [f'{logradouro} {numero}'.strip(), bairro]
    endereco = ', '.join(p for p in endereco_parts if p and p not in ('None',))
    if cep and cep not in ('None', ''): endereco += f' — CEP {cep}'

    perc = float(cnpj_data.get('percValSemComp') or 0)
    score = float(cnpj_data.get('score_risco_final') or 0)
    classificacao = cnpj_data.get('classificacao_risco') or ''
    risco_hex, risco_label = _risk_color(classificacao, score)

    if data_inicio is not None and data_fim is not None:
        periodo_txt = f'{data_inicio.strftime("%d/%m/%Y")} a {data_fim.strftime("%d/%m/%Y")}'
    elif data_inicio is not None:
        periodo_txt = f'A partir de {data_inicio.strftime("%d/%m/%Y")}'
    elif data_fim is not None:
        periodo_txt = f'Até {data_fim.strftime("%d/%m/%Y")}'
    else:
        periodo_txt = 'Histórico completo'

    criticos = _get_criticos(cnpj, data_inicio, data_fim)
    criticidade_order = _get_criticos_ordenados_por_risco(cnpj, criticos, data_inicio, data_fim)
    timing.mark("criticidades matriz")
    falecidos_comp = _build_falecidos_context(cnpj, uf, data_inicio, data_fim) if 'falecidos' in criticos else None
    if 'falecidos' in criticos and not falecidos_comp:
        raise RuntimeError('Indicador falecidos classificado como critico, mas o detalhamento de falecidos nao foi encontrado para a Nota Tecnica.')
    timing.mark("contexto falecidos")
    anexo_ii_comp = _build_anexo_ii_context(cnpj, db)
    timing.mark("contexto anexo memoria de calculo")
    try:
        crm_data_comp = get_crm_data(
            cnpj,
            data_inicio.isoformat() if data_inicio else None,
            data_fim.isoformat() if data_fim else None,
            timing_log=timing_log_enabled,
        )
    except Exception as exc:
        print(f"[NOTA_TECNICA] CRM indisponivel para {cnpj}: {exc}")
        crm_data_comp = None
    timing.mark("contexto CRM base")
    crm_evidencias_comp = _build_crm_evidencias_complementares_context(cnpj, data_inicio, data_fim, crm_data=crm_data_comp)
    timing.mark("contexto evidencias CRM complementares")
    anexo_crm_num = 'II' if crm_evidencias_comp else None
    anexo_memoria_num = 'III' if crm_evidencias_comp else 'II'
    anexo_falecidos_num = 'IV' if crm_evidencias_comp else 'III'

    # 3. Documento e margens
    doc = Document()
    style_normal: Any = doc.styles['Normal']
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Aplica o tema Grafite Médio (Slate 700) para as Seções/Títulos
    heading_sizes = {1: 13, 2: 12, 3: 11}
    for i in range(1, 4):
        try:
            style_heading: Any = doc.styles[f'Heading {i}']
            style_heading.font.size = Pt(heading_sizes[i])
            style_heading.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            if i == 2:
                style_heading.paragraph_format.space_before = Pt(14)
                style_heading.paragraph_format.space_after = Pt(10)
                style_heading.paragraph_format.line_spacing = 1.0
                style_heading.paragraph_format.keep_with_next = True
        except Exception:
            pass

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    timing.mark("inicializacao DOCX")

    PAGE_W = Inches(7.1)
    BAR_W = Inches(0.18)
    MAIN_W = PAGE_W - BAR_W

    # PÁGINA 1 — CAPA (Estilo Oficial)
    # ── 4. Cabeçalho Oficial ───────────────────────────────────────────
    p_brasao = doc.add_paragraph()
    p_brasao.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Busca a raiz do projeto (d:\sentinela) a partir deste arquivo (backend/api/services/analytics/nota_tecnica.py)
    root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..', '..'))
    brasao_path = os.path.join(root_dir, 'frontend', 'public', 'img', 'brasao_republica_mini.jpg')
    
    if os.path.exists(brasao_path):
        p_brasao.add_run().add_picture(brasao_path, width=Inches(1.5))
    else:
        # Fallback para debug se não encontrar
        print(f"⚠️ Alerta: Brasão não encontrado em {brasao_path}")
    
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_header, 'CONTROLADORIA-GERAL DA UNIÃO', color='0F172A', size=14, bold=True)
    
    doc.add_paragraph('\n' * 3)
    
    # ── 5. Título da Nota Técnica ──────────────────────────────────────
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_titulo, 'NOTA TÉCNICA PRELIMINAR\n', color='0F172A', size=24, bold=True)
    _run(p_titulo, 'SENTINELA', color='0F172A', size=14, bold=True)
    _run(p_titulo, '\nPrograma Farmácia Popular do Brasil', color='64748B', size=10, italic=True)

    doc.add_paragraph('\n' * 2)

    # ── 6. Selo de Sigilo ──────────────────────────────────────────────
    p_sigilo = doc.add_paragraph()
    p_sigilo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sigilo = _run(p_sigilo, ' ACESSO RESTRITO ', color='FFFFFF', size=10, bold=True)
    _cell_bg_run(r_sigilo, 'EF4444') # Fundo vermelho para o selo de sigilo
    
    doc.add_paragraph('\n')


    # ── 6. Resumo Executivo da Auditoria (Capa) ──────────────────────────
    tbl_resumo = doc.add_table(rows=1, cols=2)
    tbl_resumo.autofit = False
    _set_table_fixed_widths(tbl_resumo, [Inches(4.44), Inches(2.56)])
    _tbl_no_borders(tbl_resumo)
    
    # Coluna 1: Dados do Estabelecimento
    c_info = tbl_resumo.rows[0].cells[0]
    _cell_borders(c_info, bottom={'sz': '6', 'color': 'CBD5E1'})
    p_info = c_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.space_before = Pt(16)
    _run(p_info, 'IDENTIFICAÇÃO DO ESTABELECIMENTO AUDITADO\n', color='64748B', size=7, bold=True)
    _run(p_info, f'{razao_social}\n', color='0F172A', size=13, bold=True)
    if nome_fantasia:
        _run(p_info, f'Nome fantasia: {nome_fantasia}\n', color='475569', size=9)
    _run(p_info, f'CNPJ {cnpj_fmt}  •  {municipio} / {uf}\n', color='475569', size=9)
    if endereco: 
        _run(p_info, f'{endereco}\n', color='64748B', size=8)
    _run(p_info, f'Período sob análise: {periodo_txt}', color='64748B', size=8)

    # Coluna 2: Status do Risco
    c_risk = tbl_resumo.rows[0].cells[1]
    _cell_bg(c_risk, 'F8FAFC') # Cinza suave para contraste
    _cell_borders(c_risk, left={'sz': '6', 'color': 'CBD5E1'}, bottom={'sz': '6', 'color': 'CBD5E1'})
    p_risk = c_risk.paragraphs[0]
    p_risk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_risk.paragraph_format.space_before = Pt(10)
    
    _run(p_risk, 'CLASSIFICAÇÃO DE RISCO\n', color='64748B', size=7, bold=True)
    _run(p_risk, f'{risco_label}\n', color=risco_hex, size=18, bold=True)
    
    # Pequena divisória interna por parágrafo
    p_metrics = c_risk.add_paragraph()
    p_metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_metrics, 'SCORE FINAL: ', color='64748B', size=7)
    _run(p_metrics, f'{score:.1f}   ', color='0F172A', size=10, bold=True)
    _run(p_metrics, 'SEM COMPROVAÇÃO: ', color='64748B', size=7)
    _run(p_metrics, f'{perc:.1f}%', color='0F172A', size=10, bold=True)

    doc.add_paragraph('\n')
    p_ts = doc.add_paragraph()
    p_ts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ts.paragraph_format.space_before = Pt(0)
    _run(p_ts, '\n\n\n\n', color='94A3B8', size=8)
    _run(p_ts, f'Relatório extraído do Sentinela em {generated_at.strftime("%d/%m/%Y às %H:%M")}\n', color='94A3B8', size=8, italic=True)
    _run(p_ts, f'Código de verificação: {codigo_verificacao}', color='64748B', size=8, bold=True)

    # ── 4. Seção 1: Sumário (Sem Rodapé) ──────────────────────────────────
    sec_sumario = doc.add_section()
    sec_sumario.footer.is_linked_to_previous = False
    _add_confidential_watermark(sec_sumario)
    
    sec_sumario.top_margin = Inches(0.5); sec_sumario.bottom_margin = Inches(0.5)
    sec_sumario.left_margin = Inches(0.7); sec_sumario.right_margin = Inches(0.7)

    _add_sumario_official_header(
        doc,
        brasao_path,
        regional_emissora,
        generated_at.year,
        numero_nota_base,
        numero_nota_placeholder,
        numero_processo_texto,
        numero_processo_placeholder,
    )

    # SUMÁRIO
    _build_sumario(
        doc,
        criticos,
        razao_social,
        cnpj_fmt,
        criticidade_order,
    )
    timing.mark("capa e sumario")

    # ── 5. Seção 2: Assunto e Referências (Rodapé 1) ────────────────────
    sec_ref = doc.add_section(WD_SECTION.CONTINUOUS)
    sec_ref.footer.is_linked_to_previous = False
    sec_ref.footer.paragraphs[0].text = ''
    
    sec_ref.top_margin = Inches(0.5); sec_ref.bottom_margin = Inches(0.5)
    sec_ref.left_margin = Inches(0.7); sec_ref.right_margin = Inches(0.7)

    # 1. ASSUNTO
    _format_main_heading(doc.add_heading('1. ASSUNTO', level=1))
    doc.add_paragraph(f'A presente Nota Técnica (NT), de caráter investigativo e sigiloso, tem como objetivo apresentar indícios de irregularidades na atuação da Farmácia {razao_social} (CNPJ {cnpj_fmt}), credenciada junto ao Programa Farmácia Popular do Brasil (PFPB), do Ministério da Saúde (MS), para a dispensação gratuita de medicamentos aos cidadãos.')

    # 2. REFERÊNCIAS
    h2 = _format_main_heading(doc.add_heading('2. REFERÊNCIAS', level=1))
    _footnote_ref(
        doc,
        h2,
        1,
        'De acordo com informações contidas no site do Ministério da Saúde a respeito do Programa Farmácia Popular do Brasil: '
        'https://www.gov.br/saude/pt-br/composicao/sectics/farmacia-popular/legislacao '
        f'(acessado em {date.today().strftime("%d/%m/%Y")}).',
    )
    doc.add_paragraph('As principais referências normativas e técnicas utilizadas nesta análise incluem:')
    ref_list = [
        'Lei nº 10.858, de 06.05.2004, que instituiu o Programa Farmácia Popular do Brasil (PFPB);',
        'Decreto nº 5.090, de 20.05.2004, que regulamentou o PFPB;',
        'Portaria GM/MS nº 491, de 09.03.2006, que habilitou farmácias e drogarias privadas;',
        'Portaria GM/MS nº 184, de 03.02.2011, que estabeleceu normas operacionais do Programa;',
        'Portaria de Consolidação GM/MS nº 5, de 28.09.2017, marco regulatório atual do Programa Farmácia Popular do Brasil;',
        'Portaria GM/MS nº 2.898, de 03.11.2021, que ampliou para dez anos o prazo de guarda da documentação comprobatória das dispensações;',
        'Portaria GM/MS nº 1.053, de 12.05.2022, que regulamentou o procedimento de averiguação de fatos relacionados a indícios ou notícias de irregularidades no âmbito do PFPB;',
        'Relatório de Apuração CGU nº 823121, publicado em 04.01.2024;',
        'Portaria GM/MS nº 6.613, de 13.02.2025, que extinguiu a modalidade de copagamento do Programa.',
    ]
    for ref in ref_list: doc.add_paragraph(ref, style='List Bullet')

    # ── 6. Seção 3: Introdução (Sem rodapé específico) ──────────────────
    sec_intro = doc.add_section()
    sec_intro.footer.is_linked_to_previous = False
    
    sec_intro.top_margin = Inches(0.5); sec_intro.bottom_margin = Inches(0.5)
    sec_intro.left_margin = Inches(0.7); sec_intro.right_margin = Inches(0.7)

    # 3. INTRODUÇÃO
    _format_main_heading(doc.add_heading('3. INTRODUÇÃO', level=1))
    doc.add_paragraph(f'No âmbito dos trabalhos de monitoramento e avaliação dos gastos do Ministério da Saúde com o Programa Farmácia Popular do Brasil, a presente Nota Técnica (NT) trata de indícios de fraudes cometidas pela Farmácia {razao_social} (CNPJ {cnpj_fmt}).')
    
    p_intro = doc.add_paragraph()
    _run(p_intro, 'A partir da metodologia desenvolvida pela CGU, sintetizada no item 4 da NT e consignada no Relatório de Auditoria nº 823121 constante de seu ANEXO I, foi identificada, para a Farmácia ', color='0F172A', size=11)
    _run(p_intro, razao_social, color='334155', size=11, bold=True)
    _run(p_intro, ', no período de ', color='0F172A', size=11)
    _run(p_intro, periodo_txt, color='334155', size=11, bold=True, underline=True)
    _run(p_intro, ', conforme item 6 e detalhamento contido em seu ANEXO II, ausência significativa de estoque compatível com as vendas (distribuições) de medicamentos realizadas à população, denominada pela CGU como “vendas sem comprovação”, o que sugere a possibilidade de fraudes cometidas pelo estabelecimento por meio do registro fictício de dispensações de medicamentos.', color='0F172A', size=11)
    
    snippets = [f'[Subitem 6.1] evolução atípica das transferências do Programa e das possíveis “vendas sem comprovação” realizadas pela Farmácia {razao_social}']
    criticidade_start = 1
    criticidade_items_intro = _iter_criticidade_items(
        criticos,
        razao_social,
        start_index=criticidade_start,
        ordered_keys=criticidade_order,
    )
    for _, num, full_title in criticidade_items_intro:
        snippets.append(f'[Subitem {num}] {full_title[:1].lower()}{full_title[1:]}')
    
    if len(snippets) <= 4:
        texto_snippets = ("; ".join(snippets[:-1]) + "; e " + snippets[-1]) if len(snippets) > 1 else snippets[0]
        doc.add_paragraph(f'Além disso, a presente NT revela criticidades que corroboram o achado principal de “vendas sem comprovação”, como: {texto_snippets}.')
    else:
        doc.add_paragraph('Além disso, a presente NT revela criticidades que corroboram o achado principal de “vendas sem comprovação”. Em síntese, destacam-se:')
        for snippet in snippets:
            p_snippet = doc.add_paragraph(style='List Bullet')
            p_snippet.paragraph_format.space_after = Pt(2)
            _run(p_snippet, f'{snippet}.', color='0F172A', size=9)
    if crm_evidencias_comp:
        if len(snippets) > 4:
            doc.add_paragraph()
        doc.add_paragraph(
            f'Adicionalmente, o ANEXO {anexo_crm_num} desta Nota Técnica traz evidências complementares relacionadas ao uso de CRMs no SAV, incluindo volume diário atípico de prescrições por CRM, volume de autorizações em horário anômalo, concentração temporal de autorizações vinculadas a um mesmo CRM, episódios de autorizações concentradas envolvendo múltiplos CRMs e CRMs de interesse com alertas operacionais associados ao estabelecimento, relevantes para a compreensão dos padrões de prescrição e dispensação observados no estabelecimento auditado.'
        )
    doc.add_paragraph('A NT traz ainda, em seu item 5, análise da empresa em relação aos seus sócios, capital social, porte, situação cadastral junto à Receita Federal do Brasil e ao PFPB, bem como da compatibilidade entre o número de empregados e o volume de recursos recebidos do MS.')

    fontes = ['Cadastro Nacional de Pessoas Jurídicas (CNPJ) e Cadastro de Pessoa Física (CPF) da Receita Federal do Brasil', 'Sistema de Escrituração Digital das Obrigações Fiscais, Previdenciárias e Trabalhistas (eSocial)', 'Sistema Integrado de Administração Financeira do Governo Federal (SIAFI)']
    if 'polimedicamento' in criticos or 'teto' in criticos: fontes.append('dados demográficos oficiais fornecidos pelo Instituto Brasileiro de Geografia e Estatística (IBGE)')
    if any(k in criticos for k in ['hhi_crm', 'crms_irregulares']): fontes.append('cadastros de médicos do Conselho Regional de Medicina (CRM)')
    fontes_txt = ("; ".join(fontes[:-1]) + "; e " + fontes[-1]) if len(fontes) > 1 else fontes[0]
    doc.add_paragraph(f'Os achados advindos das análises realizadas, consignados nos itens 5, 6 e 7 desta Nota Técnica, tomaram por base informações registradas pela Farmácia {razao_social} no Sistema Autorizador de Vendas (SAV) do Programa Farmácia Popular do Brasil e cópias de notas fiscais eletrônicas relativas a aquisições de medicamentos por ela realizadas, compartilhadas pela Receita Federal do Brasil. Além dessas informações, foram utilizados dados extraídos das seguintes fontes: {fontes_txt}.')

    nota_pfpb_2 = (
        'Consulta ao site https://www.gov.br/saude/pt-br/composicao/sectics/farmacia-popular, '
        f'em {date.today().strftime("%d/%m/%Y")}.'
    )
    nota_pfpb_3 = (
        'A lista dos medicamentos e produtos do PFPB, atualizada em 02.09.2025, pode ser obtida no endereço: '
        'https://www.gov.br/saude/pt-br/composicao/sectics/farmacia-popular/arquivos/elenco-de-medicamentos-e-insumos.pdf.'
    )
    nota_pfpb_4 = (
        'Após um intervalo sem exigência de renovação anual obrigatória do credenciamento desde 2018, conforme o artigo 15 '
        'do Anexo LXXVII da Portaria de Consolidação nº 5, de 28 de setembro de 2017, o Ministério da Saúde '
        'retomou essa exigência a partir de 17 de abril de 2025.'
    )
    nota_pfpb_5 = (
        'Cabe informar que existia também a modalidade de copagamento (em que o beneficiário arcava com uma parte '
        'do custo), que foi extinta após a edição da Portaria GM/MS nº 6.613, de 13.02.2025.'
    )
    nota_pfpb_6 = (
        'A Portaria GM/MS nº 111/2016, substituída pela Portaria GM/MS nº 2.898/2021, determinava, em seu art. 22, '
        'que o estabelecimento deveria manter a documentação comprobatória por cinco anos.'
    )

    # ── 7. Seção 4.1: Sobre o Programa ─────────────────────────────────────
    sec_41 = doc.add_section(WD_SECTION.CONTINUOUS)
    sec_41.footer.is_linked_to_previous = False
    sec_41.footer.paragraphs[0].text = ''
    sec_41.top_margin = Inches(0.5); sec_41.bottom_margin = Inches(0.5)
    sec_41.left_margin = Inches(0.7); sec_41.right_margin = Inches(0.7)

    # 4. SÍNTESE
    _format_main_heading(doc.add_heading('4. SÍNTESE DO PROGRAMA FARMÁCIA POPULAR DO BRASIL E DA METODOLOGIA DESENVOLVIDA PELA CGU PARA SEU MONITORAMENTO', level=1))
    doc.add_heading('4.1. Sobre o Programa Farmácia Popular do Brasil', level=2)
    p_intro_41 = doc.add_paragraph(
        'O Programa Farmácia Popular do Brasil, instituído em 2004 para ampliar o acesso a medicamentos essenciais, '
        'consolidou-se como um pilar da saúde pública brasileira. Segundo site do Ministério da Saúde'
    )
    _footnote_ref(doc, p_intro_41, 2, nota_pfpb_2)
    p_intro_41.add_run(', o Programa Farmácia Popular do Brasil – PFPB é:')
    # Bloco de Citação 1
    p_quote1 = doc.add_paragraph(style='Quote')
    p_quote1.paragraph_format.left_indent = Inches(1.57) # ~4cm
    p_quote1.paragraph_format.space_after = Pt(12)       # Adiciona espaço de uma linha entre os blocos
    run_quote1 = p_quote1.add_run(
        '... programa do Governo Federal que visa complementar a disponibilização de medicamentos utilizados na '
        'Atenção Primária à Saúde, por meio de parceria com farmácias da rede privada. Dessa forma, além das '
        'Unidades Básicas de Saúde e farmácias municipais, o cidadão pode obter medicamentos nas farmácias '
        'credenciadas ao Farmácia Popular.'
    )
    run_quote1.font.size = Pt(10)

    # Bloco de Citação 2 (Continuação na linha de baixo)
    p_quote2 = doc.add_paragraph(style='Quote')
    p_quote2.paragraph_format.left_indent = Inches(1.57) # ~4cm
    p_quote2.add_run('A partir de 14 de fevereiro de 2025, o Programa Farmácia Popular passou a disponibilizar gratuitamente 100% dos medicamentos e insumos de seu elenco à população brasileira. O programa atende 12 indicações, contemplando medicamentos para hipertensão, diabetes, asma, osteoporose, dislipidemia (colesterol alto), rinite, doença de Parkinson, glaucoma, diabetes mellitus associada a doenças cardiovasculares e anticoncepção. Além disso, oferece fraldas geriátricas para pessoas com incontinência e absorventes higiênicos para beneficiárias do Programa Dignidade Menstrual')
    _footnote_ref(doc, p_quote2, 3, nota_pfpb_3)
    p_quote2.add_run('.')
    for run in p_quote2.runs:
        if not run.font.superscript: run.font.size = Pt(10)

    p_op = doc.add_paragraph('A operacionalização do Programa ocorre com a participação de drogarias credenciadas pelo Ministério da Saúde (MS)')
    _footnote_ref(doc, p_op, 4, nota_pfpb_4)
    p_op.add_run(', que realizam a dispensação gratuita de medicamentos diretamente aos cidadãos. As drogarias são posteriormente ressarcidas pela União, de acordo com informações relativas às quantidades distribuídas de cada medicamento')
    _footnote_ref(doc, p_op, 5, nota_pfpb_5)
    p_op.add_run('.')

    p_sav = doc.add_paragraph('As informações sobre as dispensações são encaminhadas mensalmente pelas drogarias credenciadas ao MS por meio do Sistema Autorizador de Vendas (SAV), conforme disposto na Portaria de Consolidação GM/MS nº 5, de 28.09.2017, e normas anteriores. Por sua vez, o art. 22 da Portaria GM/MS nº 2.898, de 03.11.2021, dispõe que o estabelecimento deve manter por 10 (dez) anos')
    _footnote_ref(doc, p_sav, 6, nota_pfpb_6)
    p_sav.add_run(', em ordem cronológica de emissão, duas cópias mantidas em locais distintos, uma em meio físico e outra em arquivo digitalizado, dos cupons vinculados assinados, dos documentos fiscais, das prescrições, dos laudos ou atestados médicos e dos documentos de identidade oficial apresentados no ato da compra e, ainda, dos documentos fiscais de aquisição dos respectivos medicamentos e/ou fraldas geriátricas dispensados no âmbito do PFPB.')

    nota_cgu_7 = (
        'A CGU, em seu Relatório de Auditoria nº 823121, considerou "vendas sem comprovação" no âmbito do PFPB '
        'a diferença identificada por princípio ativo/insumo, após o batimento entre Notas Fiscais de entrada '
        '(compartilhadas pela Receita Federal do Brasil e relativas a aquisições de medicamentos do PFPB) e registro '
        'de saída no Sistema Autorizador de Vendas – SAV (onde as dispensações subsidiadas são informadas), tendo '
        'como elo os números que constam abaixo dos códigos de barra (Código GTIN).'
    )

    # ── 8. Seção 4.2: Metodologia CGU ───────────────────────────────────────
    sec_42 = doc.add_section(WD_SECTION.CONTINUOUS)
    sec_42.footer.is_linked_to_previous = False
    sec_42.footer.paragraphs[0].text = ''
    sec_42.top_margin = Inches(0.5); sec_42.bottom_margin = Inches(0.5)
    sec_42.left_margin = Inches(0.7); sec_42.right_margin = Inches(0.7)

    doc.add_heading('4.2. Sobre a metodologia desenvolvida pela CGU para apuração de possíveis “vendas sem comprovação”', level=2)
    doc.add_paragraph('O crescimento exponencial do PFPB, com gastos que saltaram de R$ 34,7 milhões em 2006 para patamares próximos a R$ 6 bilhões em 2025, impôs desafios complexos ao controle governamental, dada a capilaridade de mais de 30 mil estabelecimentos credenciados.')
    p_sent = doc.add_paragraph()
    _run(p_sent, 'Para enfrentar essa realidade, a CGU elaborou o ', color='0F172A', size=11)
    _run(p_sent, 'Relatório de Apuração nº 823121', color='334155', size=11, bold=True)
    _run(p_sent, ' (ANEXO I desta NT), fundamentado no desenvolvimento do ', color='0F172A', size=11)
    _run(p_sent, 'Sentinela', color='334155', size=11, bold=True)
    _run(p_sent, ', uma ferramenta de tecnologia da informação que automatiza o cruzamento de dados, em larga escala, do SAV com outras bases de informações.', color='0F172A', size=11)
    p_cgu = doc.add_paragraph('De forma sintética, a premissa central de controle adotada pela CGU, apresentada de forma detalhada no referido relatório, é de natureza lógica e contábil: um estabelecimento não pode dispensar medicamentos que não adquiriu formalmente. Caso isso ocorra, a farmácia estaria praticando uma “venda sem comprovação”')
    _footnote_ref(doc, p_cgu, 7, nota_cgu_7)
    p_cgu.add_run(', ou seja, uma distribuição de medicamentos para cidadãos, cobrada do Ministério da Saúde, sem comprovação de suas aquisições.')
    doc.add_paragraph('Para a aferição da regularidade das dispensações realizadas pelas farmácias, é necessário estimar um estoque inicial dos medicamentos para que seja possível, a partir dessa informação e de suas compras posteriores, verificar a compatibilidade de suas vendas no âmbito do PFPB. Dada a limitação do SAV, em razão da inexistência de informação disponível sobre o estoque inicial de medicamentos de cada drogaria credenciada pelo MS, a CGU desenvolveu metodologia que confronta as informações de vendas de medicamentos enviadas pelas farmácias ao Ministério da Saúde com as informações de suas compras contidas na base de Notas Fiscais Eletrônicas (NF-e) da Receita Federal do Brasil, utilizada tanto para estimar seus estoques iniciais quanto para aferir a compatibilidade destes e de suas compras posteriores com as vendas realizadas no âmbito do Programa.')
    p_cutoff = doc.add_paragraph('A metodologia técnica do Sentinela foi desenhada de forma conservadora para garantir a robustez dos achados. O sistema utiliza a técnica de ')
    p_cutoff.add_run('cut-off').italic = True
    p_cutoff.add_run(', estimando o estoque inicial como a soma das duas últimas compras anteriores à primeira venda registrada de cada medicamento. A partir desse ponto, o algoritmo realiza um balanço diário de entradas e saídas, considerando apenas as vendas do PFPB como débito no estoque e ignorando vendas privadas para o público geral, o que gera um saldo “virtual” favorável à farmácia. Em outras palavras, o conservadorismo da metodologia da CGU se ampara no fato de considerar, para os cálculos de estoque, que todos os medicamentos adquiridos pela farmácia que fazem parte do rol do PFPB foram vendidos somente para clientes que fizeram uso do Programa. Assim, a metodologia não leva em conta a possibilidade real de que parte desses medicamentos tenha sido vendida para clientes comuns, que desembolsaram recursos próprios para suas aquisições.')

    p_gtin = doc.add_paragraph()
    _run(p_gtin, 'Juridicamente, o controle sustenta-se na Portaria de Consolidação GM/MS nº 5/2017, que obriga a guarda das notas fiscais de aquisição por dez anos, e no Ajuste SINIEF nº 16/2010, que exige a identificação do produto pelo código ', color='0F172A', size=11)
    _run(p_gtin, 'GTIN/EAN', color='334155', size=11, bold=True)
    _run(p_gtin, '. Nesse sentido, reforça-se que a descrição textual do produto é insuficiente para a liquidação da despesa, sendo o código de barras a única chave capaz de vincular com precisão o medicamento comprado ao preço de referência pago pelo governo.', color='0F172A', size=11)
    doc.add_paragraph('Além do levantamento de valores de “vendas sem comprovação” para todas as empresas que operam no PFPB, o Sentinela extrai dos dados do Sistema Autorizador de Vendas (SAV) do Programa uma série de informações que permitem apontar para outras criticidades que corroboram a suspeita de possíveis registros fictícios de dispensações de medicamentos por parte dos estabelecimentos.')
    doc.add_paragraph(f'A seguir, são apresentadas informações sobre a Farmácia {razao_social} e o resultado das análises dos alertas extraídos para ela do Sentinela, tanto em relação a possíveis “vendas sem comprovação” quanto a outras criticidades que corroboram esse achado principal.')

    # ── Seção 5 intro (sem rodapé) ────────────────────────────────────────
    sec_5_intro = doc.add_section(WD_SECTION.CONTINUOUS)
    sec_5_intro.footer.is_linked_to_previous = False
    sec_5_intro.top_margin = Inches(0.5); sec_5_intro.bottom_margin = Inches(0.5)
    sec_5_intro.left_margin = Inches(0.7); sec_5_intro.right_margin = Inches(0.7)

    # 5. SOBRE A FARMACIA
    _format_main_heading(doc.add_heading(f'5. SOBRE A FARMÁCIA {razao_social} (CNPJ {cnpj_fmt})', level=1))
    doc.add_heading(f'5.1 Informações sobre a Farmácia {razao_social} (CNPJ {cnpj_fmt})', level=2)
    ultimo_mes_sav = _build_ultimo_mes_sav_context(cnpj, data_inicio, data_fim)
    situacao_pfpb = "ATIVA" if cnpj_data.get("is_conexao_ativa") else "INATIVA"
    p_sav_5 = doc.add_paragraph(
        f'Informações extraídas do SAV, relativas ao período de {periodo_txt}, apontam que a Farmácia {razao_social} '
        'se encontra “'
    )
    _run(p_sav_5, situacao_pfpb, bold=True)
    _run(p_sav_5, '” no Programa Farmácia Popular do Brasil, tendo realizado vendas totais de ')
    _run(p_sav_5, f'R$ {_format_decimal_pt(ultimo_mes_sav["total"], 2)}', underline=True)
    _run(p_sav_5, ' em ')
    _run(p_sav_5, ultimo_mes_sav["mes_formatado"], underline=True)
    _run(p_sav_5, ', último mês com movimentação disponível para a farmácia na base de dados.')

    # ── Seção de informações cadastrais ────────────────────────────────────
    sec_51 = doc.add_section(WD_SECTION.CONTINUOUS)
    sec_51.footer.is_linked_to_previous = False
    sec_51.footer.paragraphs[0].text = ''
    sec_51.top_margin = Inches(0.5); sec_51.bottom_margin = Inches(0.5)
    sec_51.left_margin = Inches(0.7); sec_51.right_margin = Inches(0.7)

    # Mapeamento do porte conforme padrões RFB/Filtros
    porte_raw = getattr(cnpj_data_obj, 'porte_empresa', 'ND') if cnpj_data_obj else "ND"
    porte_txt = "empresa"
    porte_lower = porte_raw.lower()
    if "microempresa" in porte_lower:
        porte_txt = "microempresa"
    elif "pequeno porte" in porte_lower:
        porte_txt = "empresa de pequeno porte"
    elif "médio" in porte_lower or "medio" in porte_lower:
        porte_txt = "empresa de médio porte"
    elif "grande" in porte_lower:
        porte_txt = "empresa de grande porte"
    elif "demais" in porte_lower:
        porte_txt = "empresa de médio/grande porte"
    
    situacao = getattr(cnpj_data_obj, 'situacao_rf', 'ATIVA') if cnpj_data_obj else "ATIVA"
    
    cap_social_val = Decimal(str(cadastro.get('capital_social') or 0.0))
    cap_social_txt = f"R$ {cap_social_val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

    p_intro_51 = doc.add_paragraph()
    _run(p_intro_51, f'De acordo com informações contidas no Cadastro Nacional de Pessoas Jurídicas da Receita Federal do Brasil (RFB), a seguir detalhadas, a Farmácia {razao_social}, localizada no município de {municipio}/{uf}, é uma {porte_txt}, com capital social de ')
    _run(p_intro_51, cap_social_txt, bold=True)
    _run(p_intro_51, ' e situação ')
    _run(p_intro_51, situacao.upper(), bold=True, underline=True)
    _run(p_intro_51, ':')

    # Adiciona o Quadro 01 (Informações Detalhadas)
    
    # Prepara dicionário para o quadro
    quadro_data = {
        'cnpj_fmt': cnpj_fmt,
        'razao_social': razao_social,
        'nome_fantasia': cadastro.get('nome_fantasia'),
        'natureza_juridica': cadastro.get('natureza_juridica'),
        'id_cnae_principal': cadastro.get('id_cnae_principal'),
        'cnae_principal': cadastro.get('cnae_principal'),
        'id_cnae_secundario': cadastro.get('id_cnae_secundario'),
        'cnae_secundario': cadastro.get('cnae_secundario'),
        'data_abertura': cadastro.get('data_abertura'),
        'situacao_rf': cnpj_data.get('situacao_rf') or 'ATIVA',
        'porte_empresa': cnpj_data.get('porte_empresa') or 'MICROEMPRESA',
        'endereco_completo': endereco,
        'bairro': bairro,
        'municipio': municipio,
        'uf': uf,
        'cep': cep,
        'telefone_1': cadastro.get('telefone_1'),
        'telefone_2': cadastro.get('telefone_2'),
        'email': cadastro.get('email'),
        'data_processamento': cadastro.get('data_processamento'),
        'total_mov': cnpj_data.get('totalMov') or 0.0,
        'socios_ativos': socios_ativos
    }
    # Inicia numeração de footnotes reais a partir de 8 (notas 1-7 estão nos rodapés de seção)
    _add_quadro_identificacao(doc, quadro_data, cap_social_val, periodo_txt)
    timing.mark("secao 5 identificacao e quadro cadastral")

    esocial_comp = _build_esocial_context(cnpj, data_inicio, data_fim)
    timing.mark("contexto esocial")
    _add_esocial_context_text(doc, razao_social, cnpj_fmt, esocial_comp)
    timing.mark("secao 5 contexto esocial")

    # ── 10. Seção 6 (rodapé limpo até o comparativo regional) ────────────────
    _start_section(doc)

    h6 = _format_main_heading(doc.add_heading(f'6. SOBRE “VENDAS SEM COMPROVAÇÃO” REALIZADAS PELA FARMÁCIA {razao_social}', level=1))
    _add_bookmark(h6, "secao6_percentual_nao_comprovacao")
    p_53 = doc.add_paragraph()
    _run(p_53, f'Em relação à Farmácia {razao_social}, verificou-se, conforme detalhamento contido no ANEXO {anexo_memoria_num} desta Nota Técnica, diferenças relevantes entre os estoques de medicamentos estimados e suas distribuições para os cidadãos subsidiadas pelo Programa Farmácia Popular do Brasil, ', color='0F172A', size=10)
    
    if data_inicio is not None and data_fim is not None:
        _run(p_53, 'no período de ', color='0F172A', size=10)
        _run(p_53, periodo_txt, color='0F172A', size=10, bold=True)
        _run(p_53, '. ', color='0F172A', size=10)
    else:
        _run(p_53, 'no período avaliado (', color='0F172A', size=10)
        _run(p_53, periodo_txt, color='0F172A', size=10, bold=True)
        _run(p_53, '). ', color='0F172A', size=10)
    _run(p_53, 'O quadro, a seguir, consolida os valores apurados para todas as dispensações de medicamentos realizadas pelo estabelecimento:', color='0F172A', size=10)
        
    _add_quadro_53(doc, razao_social, cnpj_fmt, cnpj_data, periodo_txt)
    
    p_conclusao_53 = doc.add_paragraph()
    _run(p_conclusao_53, f'Depreende-se do quadro anterior que a quantidade de dispensações de medicamentos informadas pela Farmácia {razao_social} no SAV não se encontra compatível com seus estoques, contabilizados de acordo com a metodologia adotada pela CGU, o que levou à estimativa de não comprovação de vendas no percentual de ', color='0F172A', size=10)
    
    perc_fmt = f"{cnpj_data.get('percValSemComp', 0):.2f}%".replace('.', ',')
    val_fmt = f"{cnpj_data.get('valSemComp', 0):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    _run(p_conclusao_53, perc_fmt, color='334155', size=10, bold=True)
    _run(p_conclusao_53, ', que corresponde a um potencial desvio de recursos públicos no montante estimado de ', color='0F172A', size=10)
    _run(p_conclusao_53, f'R$ {val_fmt}', color='334155', size=10, bold=True)
    _run(p_conclusao_53, '.', color='0F172A', size=10)
    timing.mark("secao 6 quadro geral de vendas sem comprovacao")

    regional_comp = _build_regional_comparison_context(cnpj_data, cadastro, data_inicio, data_fim)
    timing.mark("contexto comparativo regional")
    multiplicador_fmt = _format_decimal_pt(regional_comp["multiplicador"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(regional_comp["multiplicador_uf"], 2)
    multiplicador_brasil_fmt = _format_decimal_pt(regional_comp["multiplicador_brasil"], 2)
    multiplicador_unidade = _vez_ou_vezes(float(regional_comp["multiplicador"] or 0.0))
    multiplicador_uf_unidade = _vez_ou_vezes(float(regional_comp["multiplicador_uf"] or 0.0))
    multiplicador_brasil_unidade = _vez_ou_vezes(float(regional_comp["multiplicador_brasil"] or 0.0))
    qtd_farmacias = regional_comp["qtd_farmacias"]
    farmacia_txt = "farmácia" if qtd_farmacias == 1 else "farmácias"
    que_opera_txt = "que opera" if qtd_farmacias == 1 else "que operam"
    localizada_txt = "localizada" if qtd_farmacias == 1 else "localizadas"
    qtd_municipios = len(regional_comp["municipios"])
    municipio_txt = "município" if qtd_municipios == 1 else "municípios"
    municipios_txt = _format_list_pt(regional_comp["municipios"])

    p_regional_53 = doc.add_paragraph()
    _run(p_regional_53, 'Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p_regional_53, f'{multiplicador_fmt} {multiplicador_unidade}', color='334155', size=10, bold=True)
    _run(p_regional_53, ' a mediana dos percentuais de “vendas sem comprovação” das farmácias da sua região,', color='0F172A', size=10)
    _footnote_ref(
        doc,
        p_regional_53,
        11,
        'A região de saúde utilizada para os comparativos do Sentinela segue a mesma estabelecida pelo Sistema Único de Saúde (ver https://www.gov.br/saude/pt-br/se/dgip/regionalizacao), que, em resumo, a considera como um espaço geográfico contínuo, formado pelo agrupamento de municípios limítrofes, que compartilham características culturais, econômicas e sociais semelhantes.',
    )
    _run(p_regional_53, ' que contempla ', color='0F172A', size=10)
    _run(p_regional_53, f'{qtd_farmacias} {farmacia_txt}', color='334155', size=10, bold=True)
    _run(p_regional_53, f' {que_opera_txt} no PFPB, {localizada_txt} em {qtd_municipios} {municipio_txt} do Estado ({regional_comp["uf"]}): {municipios_txt}.', color='0F172A', size=10)

    p_geo_ampliado = doc.add_paragraph()
    _run(p_geo_ampliado, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p_geo_ampliado, f'{multiplicador_uf_fmt} {multiplicador_uf_unidade}', color='334155', size=10, bold=True)
    _run(p_geo_ampliado, f' a mediana dos percentuais de “vendas sem comprovação” das farmácias localizadas em seu Estado ({regional_comp["uf"]}) e a ', color='0F172A', size=10)
    _run(p_geo_ampliado, f'{multiplicador_brasil_fmt} {multiplicador_brasil_unidade}', color='334155', size=10, bold=True)
    _run(p_geo_ampliado, ' a mediana dos percentuais das farmácias de todo o Brasil.', color='0F172A', size=10)

    _add_quadro_comparativo_regional(doc, regional_comp, cnpj_data, periodo_txt)
    timing.mark("quadro comparativo regional")

    posicionamento_regional_comp = _build_posicionamento_regional_context(cnpj, cadastro, data_inicio, data_fim)
    timing.mark("contexto posicionamento regional")
    p_posicionamento_intro = doc.add_paragraph()
    _run(
        p_posicionamento_intro,
        'O posicionamento regional permite observar, simultaneamente, o valor total movimentado e o percentual de dispensações sem comprovação de cada estabelecimento da mesma Região de Saúde, destacando a posição relativa da farmácia analisada.',
        color='0F172A',
        size=10,
    )
    _add_figura_posicionamento_regional(doc, razao_social, cnpj_fmt, posicionamento_regional_comp, figure_number=1)
    timing.mark("figura posicionamento regional")

    percentil_risco_comp = _build_percentil_risco_context(cnpj_data, cadastro, data_inicio, data_fim)
    timing.mark("contexto percentil de risco")
    _add_figura_percentil_risco(
        doc,
        razao_social,
        cnpj_fmt,
        percentil_risco_comp,
        figure_number=2,
        show_title=False,
    )
    timing.mark("figura percentil de risco")

    tabela_num = 0

    gtin_comp = _build_gtin_sem_comprovacao_context(cnpj, data_inicio, data_fim)
    timing.mark("contexto GTIN sem comprovacao")
    p_gtin_intro = doc.add_paragraph()
    _run(p_gtin_intro, f'Do rol de medicamentos distribuídos pela Farmácia {razao_social} sem estoque amparado em notas fiscais de aquisição, constantes do levantamento apresentado no Quadro 01, destacam-se os seguintes:', color='0F172A', size=10)

    tabela_num += 1
    tabela_gtins_num = tabela_num
    _add_tabela_gtins_sem_comprovacao(doc, razao_social, cnpj_fmt, gtin_comp, periodo_txt, tabela_gtins_num)
    timing.mark("tabela GTIN sem comprovacao")

    p_gtin_conclusao = doc.add_paragraph()
    gtins_txt = "GTIN" if gtin_comp["total_gtins"] == 1 else "GTINs"
    representativos_txt = "GTIN" if gtin_comp["representativos_count"] == 1 else "GTINs"
    _run(p_gtin_conclusao, f'Conforme a Tabela {tabela_gtins_num}, as “vendas sem comprovação” estão distribuídas em ', color='0F172A', size=10)
    _run(p_gtin_conclusao, f'{gtin_comp["total_gtins"]} {gtins_txt}', color='334155', size=10, bold=True)
    _run(p_gtin_conclusao, ', que totalizam ', color='0F172A', size=10)
    _run(p_gtin_conclusao, f'R$ {_format_decimal_pt(gtin_comp["total_valor"], 2)}', color='334155', size=10, bold=True)
    _run(p_gtin_conclusao, '. Observa-se, contudo, concentração relevante em apenas ', color='0F172A', size=10)
    _run(p_gtin_conclusao, f'{gtin_comp["representativos_count"]} {representativos_txt}', color='334155', size=10, bold=True)
    _run(p_gtin_conclusao, ', responsáveis por ', color='0F172A', size=10)
    _run(p_gtin_conclusao, f'R$ {_format_decimal_pt(gtin_comp["representativos_valor"], 2)}', color='334155', size=10, bold=True)
    _run(p_gtin_conclusao, ', o equivalente a ', color='0F172A', size=10)
    _run(p_gtin_conclusao, f'{_format_decimal_pt(gtin_comp["representativos_pct"], 1)}%', color='334155', size=10, bold=True)
    _run(p_gtin_conclusao, ' do total analisado.', color='0F172A', size=10)
    
    doc.add_heading(f'6.1 Evolução das transferências do Programa Farmácia Popular do Brasil para a Farmácia {razao_social} e das possíveis “vendas sem comprovação” por ela realizadas', level=2)

    evolucao_comp = _build_evolucao_financeira_context(cnpj, data_inicio, data_fim)
    timing.mark("contexto evolucao financeira")
    socios_volume_atipico = _build_socios_volume_atipico_context(socios_ativos, evolucao_comp)
    timing.mark("contexto socios x volume atipico")

    semestres_atipicos = evolucao_comp["semestres_atipicos"]
    if semestres_atipicos:
        p_54_contexto = doc.add_paragraph()
        p_54_contexto.paragraph_format.space_before = Pt(6)
        _run(p_54_contexto, 'No âmbito do PFPB, espera-se que as distribuições de medicamentos para a população por parte das farmácias ocorram de forma orgânica, sem saltos repentinos e excessivos que sugiram práticas de faturamento fictício em lote.', color='0F172A', size=10)

        p_54_analise = doc.add_paragraph()
        _run(p_54_analise, f'A Farmácia {razao_social} recebeu recursos provenientes do Ministério da Saúde, referentes ao PFPB, no período de ', color='0F172A', size=10)
        _run(p_54_analise, evolucao_comp["periodo_meses"], color='0F172A', size=10)
        _run(p_54_analise, '. ', color='0F172A', size=10)
        crescimento_labels = _format_list_pt([
            (
                f'{row["semestre_fmt"]} (+{_format_decimal_pt(row["taxa_crescimento_pct"], 2)}%)'
                if row.get("taxa_crescimento_pct") is not None
                else row["semestre_fmt"]
            )
            for row in semestres_atipicos
        ])
        crescimento_prefixo = 'no ' if len(semestres_atipicos) == 1 else 'nos semestres '
        _run(p_54_analise, f'Nesse intervalo, chama a atenção o aumento expressivo das transferências {crescimento_prefixo}', color='0F172A', size=10)
        _run(p_54_analise, crescimento_labels, color='0F172A', size=10, bold=True)
        _run(p_54_analise, ', sempre em comparação ao semestre imediatamente anterior. ', color='0F172A', size=10)

        top_irregulares = evolucao_comp["top_irregulares"]
        if top_irregulares:
            top_labels = _format_list_pt([row["semestre_fmt"] for row in top_irregulares])
            top_irregular_valor = round(sum(row["irregular"] for row in top_irregulares), 2)
            top_prefixo = 'no ' if len(top_irregulares) == 1 else 'nos semestres '
            _run(p_54_analise, f'Também se verificam valores relevantes de “vendas sem comprovação” {top_prefixo}', color='0F172A', size=10)
            _run(p_54_analise, top_labels, color='0F172A', size=10)
            _run(p_54_analise, ', que somam ', color='0F172A', size=10)
            _run(p_54_analise, f'R$ {_format_decimal_pt(top_irregular_valor, 2)}', color='334155', size=10, bold=True)
            _run(p_54_analise, ', conforme tabela e figura a seguir.', color='0F172A', size=10)

    tabela_num += 1
    _add_quadro_evolucao_financeira(doc, razao_social, cnpj_fmt, evolucao_comp, tabela_num)
    medicamentos_aumento_atipico = evolucao_comp.get("medicamentos_aumento_atipico") or []
    if medicamentos_aumento_atipico:
        p_medicamentos_volume = doc.add_paragraph()
        p_medicamentos_volume.paragraph_format.space_before = Pt(6)
        _run(
            p_medicamentos_volume,
            'Para qualificar os semestres destacados como atípicos, a tabela complementar a seguir decompõe o aumento financeiro por GTIN, comparando o semestre atípico com o semestre anterior utilizado no cálculo do alerta.',
            color='0F172A',
            size=10,
        )
        tabela_num += 1
        _add_tabela_medicamentos_aumento_atipico(doc, medicamentos_aumento_atipico, tabela_num)
    if socios_volume_atipico:
        p_socios_volume = doc.add_paragraph()
        p_socios_volume.paragraph_format.space_before = Pt(6)
        _run(p_socios_volume, 'Também se observam ingressos societários próximos a semestres com aumento atípico das transferências, conforme detalhado no quadro a seguir.', color='0F172A', size=10)
    _add_quadro_socios_volume_atipico(doc, socios_volume_atipico)
    _add_figura_evolucao_financeira(doc, razao_social, cnpj_fmt, evolucao_comp, figure_number=3)
    timing.mark("quadros e figura evolucao financeira")

    # Seção 7 sem rodapé herdado da seção 6.
    _start_section(doc)
    _format_main_heading(doc.add_heading(f'7. SOBRE OUTRAS CRITICIDADES RELATIVAS À FARMÁCIA {razao_social}, NO ÂMBITO DO PFPB.', level=1))
    doc.add_paragraph(f'Analisando-se informações declaradas pela Farmácia {razao_social} no SAV e, em alguns casos, cruzando-as com outras bases de dados, foram identificadas criticidades que corroboram o achado principal de “vendas sem comprovação” apurado para ela. A tabela, a seguir, sintetiza os indicadores classificados como críticos na matriz de risco do Sistema Sentinela. Na sequência, são detalhadas as criticidades com evidências analíticas específicas para a presente Nota Técnica.')
    indicadores_criticos_quadro = _build_indicadores_criticos_quadro(cnpj, data_inicio, data_fim)
    if indicadores_criticos_quadro:
        tabela_num += 1
        _add_indicadores_criticos_quadro(doc, indicadores_criticos_quadro, tabela_num)
        timing.mark("secao 7 quadro indicadores criticos")
    resumos_criticidades: list[str] = []
    criticidade_start = 1

    def _add_enquadramento_regional_indicador(key: str) -> None:
        nonlocal tabela_num
        regional_context = _build_indicador_regional_context(cnpj, key, data_inicio, data_fim)
        tabela_num += 1
        _add_indicador_regional_table(doc, regional_context, tabela_num)
        timing.mark(f"secao 7 enquadramento regional {key}")

    criticidade_items = _iter_criticidade_items(
        criticos,
        razao_social,
        start_index=criticidade_start,
        ordered_keys=criticidade_order,
    )
    if criticidade_items:
        for key, num, full_title in criticidade_items:
            bookmark_name = f"secao7_{key}"
            if key == 'falecidos':
                if not falecidos_comp:
                    raise RuntimeError('Indicador falecidos classificado como critico, mas o contexto detalhado esta ausente na Nota Tecnica.')
                _add_falecidos_criticidade_text(
                    doc,
                    num,
                    razao_social,
                    falecidos_comp,
                    anexo_num=anexo_falecidos_num,
                    bookmark_name=bookmark_name,
                )
                _add_enquadramento_regional_indicador(key)
                resumos_criticidades.append(_build_resumo_falecidos(num, falecidos_comp))
                timing.mark(f"secao 7 criticidade {key}")
                continue
            if key == 'incompatibilidade_patologica':
                clinico_comp = _build_incompatibilidade_patologica_context(cnpj, data_inicio, data_fim)
                if clinico_comp:
                    if clinico_comp.get("unavailable"):
                        p_clinico_indisponivel = doc.add_paragraph()
                        _run(
                            p_clinico_indisponivel,
                            'O detalhamento clínico desta criticidade não foi incluído porque não há análise clínica materializada para o CNPJ no período selecionado.',
                            color='0F172A',
                            size=10,
                        )
                        timing.mark(f"secao 7 criticidade {key} indisponivel")
                        continue
                    tabela_inicio_clinica = tabela_num + 1
                    _add_incompatibilidade_patologica_text(
                        doc,
                        num,
                        razao_social,
                        clinico_comp,
                        tabela_inicio_clinica,
                        bookmark_name=bookmark_name,
                    )
                    tabela_num += _count_incompatibilidade_patologica_tables(clinico_comp)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, clinico_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'teto':
                teto_comp = _build_teto_context(cnpj, data_inicio, data_fim)
                if teto_comp:
                    _add_teto_text(doc, num, razao_social, teto_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, teto_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'polimedicamento' and full_title.startswith('Vendas de quatro ou mais itens'):
                polimedicamento_comp = _build_polimedicamento_context(cnpj, data_inicio, data_fim)
                if polimedicamento_comp:
                    _add_polimedicamento_text(doc, num, razao_social, polimedicamento_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, polimedicamento_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'ticket_medio':
                ticket_comp = _build_ticket_medio_context(cnpj, data_inicio, data_fim)
                if ticket_comp:
                    _add_ticket_medio_text(doc, num, razao_social, ticket_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, ticket_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'receita_paciente':
                receita_comp = _build_receita_paciente_context(cnpj, data_inicio, data_fim)
                if receita_comp:
                    _add_receita_paciente_text(doc, num, razao_social, receita_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, receita_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'per_capita':
                per_capita_comp = _build_per_capita_context(cnpj, data_inicio, data_fim)
                if per_capita_comp:
                    _add_per_capita_text(doc, num, razao_social, per_capita_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, per_capita_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'alto_custo':
                alto_custo_comp = _build_alto_custo_context(cnpj, data_inicio, data_fim)
                if alto_custo_comp:
                    _add_alto_custo_text(doc, num, razao_social, alto_custo_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, alto_custo_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'vendas_rapidas':
                vendas_rapidas_comp = _build_vendas_rapidas_context(cnpj, data_inicio, data_fim)
                if vendas_rapidas_comp:
                    _add_vendas_rapidas_text(doc, num, razao_social, vendas_rapidas_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, vendas_rapidas_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'recorrencia_sistemica':
                recorrencia_comp = _build_recorrencia_sistemica_context(cnpj, data_inicio, data_fim)
                if recorrencia_comp:
                    _add_recorrencia_sistemica_text(doc, num, razao_social, recorrencia_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, recorrencia_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'dias_pico':
                dias_pico_comp = _build_dias_pico_context(cnpj, data_inicio, data_fim)
                if dias_pico_comp:
                    _add_dias_pico_text(doc, num, razao_social, dias_pico_comp, bookmark_name=bookmark_name)
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, dias_pico_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'dispersao_geografica':
                dispersao_comp = _build_dispersao_geografica_context(cnpj, data_inicio, data_fim)
                if dispersao_comp:
                    tabela_num += 1
                    _add_dispersao_geografica_text(
                        doc,
                        num,
                        razao_social,
                        dispersao_comp,
                        tabela_num,
                        bookmark_name=bookmark_name,
                    )
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, dispersao_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'hhi_crm':
                hhi_crm_comp = _build_hhi_crm_context(
                    cnpj,
                    data_inicio,
                    data_fim,
                    cnpj_data.get('totalMov'),
                    crm_data=crm_data_comp,
                )
                if hhi_crm_comp:
                    tabela_num += 1
                    _add_hhi_crm_text(
                        doc,
                        num,
                        razao_social,
                        cnpj_fmt,
                        hhi_crm_comp,
                        tabela_num,
                        bookmark_name=bookmark_name,
                    )
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, hhi_crm_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                    timing.mark(f"secao 7 criticidade {key}")
                    continue
            if key == 'crms_irregulares':
                crms_irregulares_comp = _build_crms_irregulares_context(
                    cnpj,
                    data_inicio,
                    data_fim,
                    cnpj_data.get('totalMov'),
                    crm_data=crm_data_comp,
                )
                if crms_irregulares_comp:
                    tabela_num += 1
                    _add_crms_irregulares_text(
                        doc,
                        num,
                        razao_social,
                        cnpj_fmt,
                        crms_irregulares_comp,
                        tabela_num,
                        bookmark_name=bookmark_name,
                    )
                    _add_enquadramento_regional_indicador(key)
                    resumo = _build_resumo_criticidade(num, key, crms_irregulares_comp, float(cnpj_data.get('totalMov') or 0.0))
                    if resumo:
                        resumos_criticidades.append(resumo)
                timing.mark(f"secao 7 criticidade {key}")
                continue

            heading = doc.add_heading(f'{num} {full_title}', level=2)
            _add_bookmark(heading, bookmark_name)
            doc.add_paragraph(f'Foi detectado um alerta CRÍTICO para o indicador "{full_title}". Este comportamento indica uma distorção estatística severa (Modified Z-Score > 3.0) que exige verificação documental imediata.')
            _add_enquadramento_regional_indicador(key)
            timing.mark(f"secao 7 criticidade {key}")
        if crm_evidencias_comp:
            p_crm_anexo = doc.add_paragraph()
            _run(
                p_crm_anexo,
                f'Além dos indicadores críticos apresentados anteriormente, foram identificadas, no SAV, informações complementares relacionadas a CRMs vinculados a prescrições de medicamentos vendidos pela Farmácia. Essas evidências, apresentadas no ANEXO {anexo_crm_num} desta Nota Técnica, não correspondem, isoladamente, a indicadores de risco da matriz do Sentinela, mas auxiliam na contextualização da dinâmica dos lançamentos realizados pelo estabelecimento.',
                color='0F172A',
                size=10,
            )
    else:
        doc.add_paragraph('Não foram identificadas outras criticidades em nível crítico para detalhamento nesta seção, sem prejuízo do acompanhamento sistêmico dos demais indicadores do Sentinela.')
    timing.mark("secao 7 fechamento criticidades")

    # 8. CONCLUSÃO
    h8 = _format_main_heading(doc.add_heading('8. CONCLUSÃO E ENCAMINHAMENTO', level=1))
    total_mov_conclusao = float(cnpj_data.get('totalMov') or 0.0)
    val_sem_comp_conclusao = float(cnpj_data.get('valSemComp') or 0.0)
    perc_sem_comp_conclusao = float(cnpj_data.get('percValSemComp') or 0.0)
    periodo_conclusao_txt = periodo_txt.replace('/', '.')

    p_conclusao = doc.add_paragraph()
    _run(p_conclusao, 'Conforme detalhado no subitem 6 desta Nota Técnica, de um total de ', color='0F172A', size=10)
    _run(p_conclusao, f'R$ {_format_decimal_pt(total_mov_conclusao, 2)}', color='334155', size=10, bold=True)
    _run(p_conclusao, f' de medicamentos distribuídos pela Farmácia {razao_social} no âmbito do Programa Farmácia Popular do Brasil, no período de ', color='0F172A', size=10)
    _run(p_conclusao, periodo_conclusao_txt, color='334155', size=10, bold=True)
    _run(p_conclusao, ', foi identificado possível prejuízo ao erário público no valor de ', color='0F172A', size=10)
    _run(p_conclusao, f'R$ {_format_decimal_pt(val_sem_comp_conclusao, 2)}', color='334155', size=10, bold=True)
    _run(p_conclusao, ' (', color='0F172A', size=10)
    _run(p_conclusao, f'{_format_decimal_pt(perc_sem_comp_conclusao, 2)}%', color='334155', size=10, bold=True)
    _run(p_conclusao, ' daquele total), em virtude da prática de “vendas sem comprovação”, tipologia de fraude identificada pela CGU correspondente à dispensação de medicamentos sem quantitativo suficiente em estoque para suportá-la.', color='0F172A', size=10)

    _add_resumo_criticidades_conclusao(doc, resumos_criticidades)

    p_atencao_conclusao = doc.add_paragraph()
    _run(
        p_atencao_conclusao,
        'ATENÇÃO: Caso sejam identificadas criticidades em relação ao estabelecimento (sócio laranja, endereço inexistente, alteração de endereço para outro município, sócio com vínculo empregatício, etc.), trazer os achados de forma resumida para a conclusão.',
        color='DC2626',
        size=10,
        bold=True,
    )

    p_fontes_conclusao = doc.add_paragraph()
    _run(
        p_fontes_conclusao,
        f'De acordo com a “Introdução” desta Nota Técnica e os subitens de sua “Análise”, as suspeitas de “vendas sem comprovação” e as criticidades que as corroboram foram identificadas com base em informações autodeclaradas pela Farmácia {razao_social} no Sistema Autorizador de Vendas do PFPB, em cópias de notas fiscais eletrônicas relativas às suas aquisições de medicamentos, compartilhadas pela Receita Federal do Brasil, e em dados contidos em bases oficiais, que complementam as análises advindas daquele Sistema. Dada a relevância, gravidade e sensibilidade desses achados, que sugerem possível ação deliberada por parte do estabelecimento voltada ao desvio de recursos do PFPB, não foi realizada fiscalização in loco nem solicitada justificativa ao estabelecimento.',
        color='0F172A',
        size=10,
    )

    p_encaminhamento = doc.add_paragraph()
    _run(
        p_encaminhamento,
        f'Em virtude das constatações contidas na presente Nota Técnica, que indicam a necessidade de aprofundamento das investigações, propõe-se seu encaminhamento para a {unidade_pf}, para adoção das medidas que julgar pertinentes.',
        color='0F172A',
        size=10,
    )

    closing_block_paragraphs = []

    def _keep_closing_block_with_next(paragraph) -> None:
        paragraph.paragraph_format.keep_with_next = True
        closing_block_paragraphs.append(paragraph)

    def _add_signature_block(nome: str, cargo: str, *, space_before: float = 6) -> None:
        p_linha = doc.add_paragraph()
        p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_linha.paragraph_format.space_before = Pt(space_before)
        p_linha.paragraph_format.space_after = Pt(0)
        _keep_closing_block_with_next(p_linha)
        _run(p_linha, '________________________________________________________', color='0F172A', size=10)

        p_nome = doc.add_paragraph()
        p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nome.paragraph_format.space_after = Pt(0)
        _keep_closing_block_with_next(p_nome)
        _run(p_nome, nome, color='0F172A', size=10, bold=True)

        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cargo.paragraph_format.space_after = Pt(4)
        _keep_closing_block_with_next(p_cargo)
        _run(p_cargo, cargo, color='0F172A', size=10)

    for index, assinante in enumerate(assinantes_tecnicos_resolvidos):
        _add_signature_block(
            assinante["nome"],
            assinante["cargo"],
            space_before=22 if index == 0 else 14,
        )

    doc.add_paragraph()
    p_despacho_titulo = doc.add_paragraph()
    _keep_closing_block_with_next(p_despacho_titulo)
    _run(
        p_despacho_titulo,
        f'Despacho do(a) Superintendente da Controladoria-Regional da União no Estado de {regional_emissora["estado"]}',
        color='0F172A',
        size=10,
        bold=True,
    )
    p_despacho_data = doc.add_paragraph()
    _keep_closing_block_with_next(p_despacho_data)
    _run(
        p_despacho_data,
        f'{regional_emissora["cidade_uf"]}, {_format_full_date_long_pt(generated_at.date())}.',
        color='0F172A',
        size=10,
    )
    p_despacho = doc.add_paragraph()
    _keep_closing_block_with_next(p_despacho)
    _run(p_despacho, 'De acordo, encaminhe-se conforme proposto.', color='0F172A', size=10)
    _add_signature_block(
        regional_emissora["superintendente"],
        regional_emissora["cargo_superintendente"],
        space_before=22,
    )
    closing_block_paragraphs[-1].paragraph_format.keep_with_next = False
    timing.mark("conclusao e encaminhamento")

    if crm_evidencias_comp:
        tabela_num = _add_anexo_crm_evidencias(
            doc,
            razao_social,
            crm_evidencias_comp,
            tabela_num,
            numero_nota_tecnica,
            timing=timing,
            anexo_num=anexo_crm_num,
        )
        timing.mark(f"anexo {anexo_crm_num} fechamento")

    _add_anexo_ii_memoria_calculo(
        doc,
        razao_social,
        cnpj_fmt,
        periodo_txt,
        anexo_ii_comp,
        numero_nota_tecnica,
        timing=timing,
        anexo_num=anexo_memoria_num,
    )
    timing.mark(f"anexo {anexo_memoria_num} fechamento")

    if falecidos_comp:
        _add_anexo_falecidos(
            doc,
            razao_social,
            cnpj_fmt,
            falecidos_comp,
            numero_nota_tecnica,
            timing=timing,
            anexo_num=anexo_falecidos_num,
        )
        timing.mark(f"anexo {anexo_falecidos_num} fechamento")

    _apply_codigo_verificacao_footer(doc, codigo_verificacao)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    timing.mark("serializacao DOCX")
    if timing_log_enabled:
        timing.write()
    return file_stream
