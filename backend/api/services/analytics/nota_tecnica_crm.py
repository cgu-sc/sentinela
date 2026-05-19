import os
from datetime import date, datetime, timedelta
from typing import Any, Optional

import polars as pl
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from cache_files import CRM_RAIOX_TX_PARQUET
from data_cache import get_df_matriz_risco
from ._cache import _get_cnpj_cache_dir
from .crm import get_crm_data
from .nota_tecnica_docx_utils import (
    _cell_bg,
    _format_block_footnote,
    _format_block_title,
    _run,
    _set_table_fixed_widths,
    _write_cell,
)
from .nota_tecnica_formatters import _format_decimal_pt


def _as_float(value: Any) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def _as_optional_float(value: Any) -> float | None:
    if value is None:
        return None
    try:
        result = float(value)
        return None if result != result else result
    except (TypeError, ValueError):
        return None


def _format_optional_decimal_pt(value: Any, decimals: int = 1, empty: str = "n.c.") -> str:
    number = _as_optional_float(value)
    if number is None:
        return empty
    return _format_decimal_pt(number, decimals)


def _volume_horario_excesso_value(row: dict[str, Any]) -> float | None:
    multiplicador = _as_optional_float(row.get("multiplicador"))
    if multiplicador is not None:
        return multiplicador
    mediana = _as_optional_float(row.get("mediana_hora"))
    if mediana == 0:
        return float(_as_int(row.get("nu_prescricoes")))
    return None


def _format_volume_horario_excesso(row: dict[str, Any]) -> str:
    excesso = _volume_horario_excesso_value(row)
    if excesso is None:
        return "n.c."
    return f"{_format_decimal_pt(excesso, 1)}x"


def _as_int(value: Any) -> int:
    try:
        return int(float(value or 0))
    except (TypeError, ValueError):
        return 0


def _select_crm_concentracao_table_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [row for row in rows if _as_float(row.get("taxa_hora")) > 30]


def _crm_event_key(row: dict[str, Any]) -> tuple[str, int]:
    return (str(row.get("dt") or "")[:10], _as_int(row.get("hr")))


def _format_date_br(value: Any) -> str:
    if value is None:
        return "Não localizada"
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    text = str(value).strip()
    if not text:
        return "Não localizada"
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y"):
        try:
            return datetime.strptime(text[:19], fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return text


def _vez_ou_vezes(value: float) -> str:
    return "vez" if abs(value) <= 1 else "vezes"


def _crm_num_uf(id_medico: Any) -> tuple[str, str]:
    text = str(id_medico or "").strip()
    if "/" not in text:
        return text or "Não informado", ""
    numero, uf = text.split("/", 1)
    return numero.strip() or "Não informado", uf.strip()


def _format_competencia_crm(value: Any) -> str:
    text = str(value or "").strip()
    if len(text) == 6 and text.isdigit():
        return f"{text[4:6]}/{text[:4]}"
    return text or "—"


def _format_time_hour(value: Any) -> str:
    try:
        return f"{int(float(value or 0)):02d}:00"
    except (TypeError, ValueError):
        return "—"


def _format_hour_range(value: Any) -> str:
    try:
        hora = int(float(value or 0))
    except (TypeError, ValueError):
        return "—"
    return f"{hora % 24:02d}:00 até {(hora + 1) % 24:02d}:00"


def _format_datetime_br_minute(value: Any) -> str:
    dt = _parse_datetime_crm(value)
    if dt is None:
        return "N/d"
    return dt.strftime("%d/%m/%y %H:%M")


def _format_janela_minutos(value: Any) -> str:
    minutos = _as_int(value)
    if minutos <= 0:
        return "Mesmo instante"
    if minutos < 60:
        return f"{minutos} min"
    horas = minutos // 60
    resto = minutos % 60
    return f"{horas}h {resto}min" if resto else f"{horas}h"


def _format_crm_table_title(paragraph):
    _format_block_title(paragraph, space_before=16, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _format_crm_table_footnote(paragraph):
    _format_block_footnote(paragraph, space_before=5, space_after=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _format_crm_subheading(paragraph):
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    return paragraph


def _add_crm_subheading(doc, text: str):
    return _format_crm_subheading(doc.add_heading(text, level=3))


def _parse_datetime_crm(value: Any) -> datetime | None:
    if value is None:
        return None
    if hasattr(value, "to_pydatetime"):
        value = value.to_pydatetime()
    if isinstance(value, datetime):
        return value.replace(tzinfo=None)
    if isinstance(value, date):
        return datetime(value.year, value.month, value.day)

    text = str(value).strip()
    if not text:
        return None
    text = text.replace("T", " ").replace("Z", "")
    try:
        return datetime.fromisoformat(text).replace(tzinfo=None)
    except ValueError:
        pass

    for fmt in (
        "%Y-%m-%d %H:%M:%S.%f",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
    ):
        try:
            return datetime.strptime(text[:26], fmt)
        except ValueError:
            continue
    return None


def _datetime_in_alert_window(tx_dt: datetime, dt_ini: datetime, dt_fim: datetime) -> bool:
    if dt_ini <= tx_dt <= dt_fim:
        return True

    if tx_dt.second != 0 or tx_dt.microsecond != 0:
        return False

    def _floor_minute(value: datetime) -> datetime:
        return value.replace(second=0, microsecond=0)

    def _round_minute(value: datetime) -> datetime:
        floored = _floor_minute(value)
        return floored + timedelta(minutes=1) if value.second >= 30 or value.microsecond else floored

    # Algumas bases publicadas do Raio-X preservam apenas o minuto. Historicamente
    # isso apareceu tanto como truncamento quanto como arredondamento para o minuto
    # mais proximo, enquanto os alertas preservam segundos.
    tx_min = _floor_minute(tx_dt)
    ini_min = min(_floor_minute(dt_ini), _round_minute(dt_ini))
    fim_min = max(_floor_minute(dt_fim), _round_minute(dt_fim))
    return ini_min <= tx_min <= fim_min


def _datetime_windows_overlap(
    a_ini_raw: Any,
    a_fim_raw: Any,
    b_ini_raw: Any,
    b_fim_raw: Any,
) -> bool:
    a_ini = _parse_datetime_crm(a_ini_raw)
    a_fim = _parse_datetime_crm(a_fim_raw)
    b_ini = _parse_datetime_crm(b_ini_raw)
    b_fim = _parse_datetime_crm(b_fim_raw)
    if a_ini is None or a_fim is None or b_ini is None or b_fim is None:
        return False
    if a_fim < a_ini:
        a_ini, a_fim = a_fim, a_ini
    if b_fim < b_ini:
        b_ini, b_fim = b_fim, b_ini
    return a_ini <= b_fim and b_ini <= a_fim


def _enrich_crm_unico_valores(cnpj: str, rows: list[dict[str, Any]]) -> None:
    targets: list[tuple[dict[str, Any], str, str, datetime, datetime]] = []
    for row in rows:
        row["valor_alerta"] = 0.0
        row["valor_alerta_disponivel"] = False

        id_medico = str(row.get("id_medico") or "").strip()
        dt_ini = _parse_datetime_crm(row.get("dt_ini_hora"))
        dt_fim = _parse_datetime_crm(row.get("dt_fim_hora"))
        if not id_medico or dt_ini is None or dt_fim is None:
            continue
        if dt_fim < dt_ini:
            dt_ini, dt_fim = dt_fim, dt_ini
        targets.append((row, id_medico, dt_ini.date().isoformat(), dt_ini, dt_fim))

    if not targets:
        return

    parquet_path = os.path.join(_get_cnpj_cache_dir(cnpj), CRM_RAIOX_TX_PARQUET)
    if not os.path.exists(parquet_path):
        return

    id_medicos = sorted({target[1] for target in targets})
    datas = sorted({target[2] for target in targets})
    try:
        df_tx = (
            pl.scan_parquet(parquet_path)
            .select(["dt_janela", "data_hora", "id_medico", "valor_pago"])
            .with_columns([
                pl.col("dt_janela").cast(pl.Utf8).str.slice(0, 10).alias("_dt_janela"),
                pl.col("data_hora").cast(pl.Utf8),
                pl.col("id_medico").cast(pl.Utf8),
                pl.col("valor_pago").cast(pl.Float64),
            ])
            .filter(pl.col("id_medico").is_in(id_medicos) & pl.col("_dt_janela").is_in(datas))
            .collect()
        )
    except Exception as exc:
        print(f"[NOTA_TECNICA] Valores do Raio-X CRM indisponiveis para {cnpj}: {exc}")
        return

    tx_por_medico_data: dict[tuple[str, str], list[tuple[datetime, float]]] = {}
    for tx in df_tx.iter_rows(named=True):
        tx_dt = _parse_datetime_crm(tx.get("data_hora"))
        if tx_dt is None:
            continue
        key = (str(tx.get("id_medico") or ""), str(tx.get("_dt_janela") or ""))
        tx_por_medico_data.setdefault(key, []).append((tx_dt, _as_float(tx.get("valor_pago"))))

    for row, id_medico, data_key, dt_ini, dt_fim in targets:
        valores = [
            valor
            for tx_dt, valor in tx_por_medico_data.get((id_medico, data_key), [])
            if _datetime_in_alert_window(tx_dt, dt_ini, dt_fim)
        ]
        valor = sum(valores)
        row["valor_alerta"] = valor
        row["valor_alerta_disponivel"] = bool(valores)


def _enrich_crm_multiplo_valores(cnpj: str, rows: list[dict[str, Any]]) -> None:
    targets: list[tuple[dict[str, Any], str, datetime, datetime]] = []
    for row in rows:
        row["valor_alerta"] = 0.0
        row["valor_alerta_disponivel"] = False

        dt_ini = _parse_datetime_crm(row.get("dt_ini_hora"))
        dt_fim = _parse_datetime_crm(row.get("dt_fim_hora"))
        if dt_ini is None or dt_fim is None:
            continue
        if dt_fim < dt_ini:
            dt_ini, dt_fim = dt_fim, dt_ini
        targets.append((row, dt_ini.date().isoformat(), dt_ini, dt_fim))

    if not targets:
        return

    parquet_path = os.path.join(_get_cnpj_cache_dir(cnpj), CRM_RAIOX_TX_PARQUET)
    if not os.path.exists(parquet_path):
        return

    datas = sorted({target[1] for target in targets})
    try:
        df_tx = (
            pl.scan_parquet(parquet_path)
            .select(["dt_janela", "data_hora", "valor_pago"])
            .with_columns([
                pl.col("dt_janela").cast(pl.Utf8).str.slice(0, 10).alias("_dt_janela"),
                pl.col("data_hora").cast(pl.Utf8),
                pl.col("valor_pago").cast(pl.Float64),
            ])
            .filter(pl.col("_dt_janela").is_in(datas))
            .collect()
        )
    except Exception as exc:
        print(f"[NOTA_TECNICA] Valores do Raio-X CRM multiplo indisponiveis para {cnpj}: {exc}")
        return

    tx_por_data: dict[str, list[tuple[datetime, float]]] = {}
    for tx in df_tx.iter_rows(named=True):
        tx_dt = _parse_datetime_crm(tx.get("data_hora"))
        if tx_dt is None:
            continue
        tx_por_data.setdefault(str(tx.get("_dt_janela") or ""), []).append((tx_dt, _as_float(tx.get("valor_pago"))))

    for row, data_key, dt_ini, dt_fim in targets:
        valores = [
            valor
            for tx_dt, valor in tx_por_data.get(data_key, [])
            if _datetime_in_alert_window(tx_dt, dt_ini, dt_fim)
        ]
        valor = sum(valores)
        row["valor_alerta"] = valor
        row["valor_alerta_disponivel"] = bool(valores)


def _plural(value: int, singular: str, plural: str) -> str:
    return singular if value == 1 else plural


def _crm_table_header(table, headers: list[str], widths: list[Any]):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _cell_bg(cell, "E2E8F0")
        _write_cell(cell, header, size=7.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _crm_alertas_contexto_labels(row: dict[str, Any]) -> list[str]:
    labels: list[str] = []
    if _as_int(row.get("flag_crm_invalido")) > 0:
        labels.append("CRM inexistente")
    if _as_int(row.get("flag_prescricao_antes_registro")) > 0:
        labels.append("CRM irregular")
    if _as_int(row.get("flag_crm_exclusivo")) > 0:
        labels.append("CRM exclusivo")
    if row.get("alertas_crm_unico"):
        labels.append("Lançamentos sequenciais")
    if row.get("alerta5_geografico") or row.get("alertas_geograficos"):
        labels.append("Registro geográfico")

    prescricoes_dia_local = _as_float(row.get("nu_prescricoes_dia"))
    prescricoes_dia_brasil = _as_float(row.get("prescricoes_dia_total_brasil"))
    alerta_30_local = _as_int(row.get("flag_robo")) > 0 or prescricoes_dia_local > 30
    alerta_30_brasil = _as_int(row.get("flag_robo_oculto")) > 0 or prescricoes_dia_brasil > 30
    if alerta_30_local and alerta_30_brasil:
        labels.append(">30 presc./dia local/Brasil")
    elif alerta_30_local:
        labels.append(">30 presc./dia local")
    elif alerta_30_brasil:
        labels.append(">30 presc./dia Brasil")

    return labels


def _build_principais_crms_contexto_rows(crms: list[dict[str, Any]]) -> list[dict[str, Any]]:
    crms_ordenados = sorted(
        crms,
        key=lambda row: (
            _as_float(row.get("vl_total_prescricoes")),
            _as_int(row.get("nu_prescricoes")),
        ),
        reverse=True,
    )
    top10_ids = {str(row.get("id_medico") or "") for row in crms_ordenados[:10]}
    crms_selecionados = list(crms_ordenados[:10])
    crms_selecionados.extend(
        row
        for row in crms_ordenados[10:]
        if str(row.get("id_medico") or "") not in top10_ids
        and _crm_alertas_contexto_labels(row)
    )

    rows: list[dict[str, Any]] = []
    for row in crms_selecionados:
        contexto_row = {
            "id_medico": row.get("id_medico"),
            "no_medico": row.get("no_medico") or "Não localizado",
            "nu_prescricoes": _as_int(row.get("nu_prescricoes")),
            "vl_total_prescricoes": _as_float(row.get("vl_total_prescricoes")),
            "pct_participacao": _as_float(row.get("pct_participacao")),
            "pct_acumulado": _as_float(row.get("pct_acumulado")),
            "nu_prescricoes_dia": _as_float(row.get("nu_prescricoes_dia")),
            "prescricoes_dia_total_brasil": _as_float(row.get("prescricoes_dia_total_brasil")),
        }
        contexto_row["alertas_contexto"] = _crm_alertas_contexto_labels(row)
        rows.append(contexto_row)
    return rows


def _build_hhi_crm_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
    crm_data: Any = None,
) -> dict[str, Any] | None:
    """Monta o contexto do subitem de concentração atípica de CRM."""
    if crm_data is None:
        try:
            crm_data = get_crm_data(
                cnpj,
                data_inicio.isoformat() if data_inicio else None,
                data_fim.isoformat() if data_fim else None,
            )
        except Exception as exc:
            print(f"[NOTA_TECNICA] CRM indisponivel para {cnpj}: {exc}")
            return None

    crms = list(getattr(crm_data, "crms_interesse", None) or [])
    if not crms:
        return None

    total_medicos = len(crms)
    total_autorizacoes = sum(_as_int(row.get("nu_prescricoes")) for row in crms)
    valor_total = sum(_as_float(row.get("vl_total_prescricoes")) for row in crms)
    if total_medicos <= 0 or total_autorizacoes <= 0:
        return None

    crms_ordenados = sorted(
        crms,
        key=lambda row: (
            _as_float(row.get("vl_total_prescricoes")),
            _as_int(row.get("nu_prescricoes")),
        ),
        reverse=True,
    )
    top_crms: list[dict[str, Any]] = []
    autorizacoes_acumuladas = 0
    for row in crms_ordenados:
        if len(top_crms) >= 10:
            break
        top_crms.append(row)
        autorizacoes_acumuladas += _as_int(row.get("nu_prescricoes"))
        if len(top_crms) >= 5 and (autorizacoes_acumuladas / total_autorizacoes * 100) >= 80:
            break
    principal = top_crms[0]
    principal_autorizacoes = _as_int(principal.get("nu_prescricoes"))
    principal_valor = _as_float(principal.get("vl_total_prescricoes"))

    media_autorizacoes = total_autorizacoes / total_medicos
    media_valor = valor_total / total_medicos if valor_total else 0.0

    if data_inicio and data_fim:
        periodo_intervalo = f'de {data_inicio.strftime("%d.%m.%Y")} a {data_fim.strftime("%d.%m.%Y")}'
    elif data_inicio:
        periodo_intervalo = f'a partir de {data_inicio.strftime("%d.%m.%Y")}'
    elif data_fim:
        periodo_intervalo = f'até {data_fim.strftime("%d.%m.%Y")}'
    else:
        periodo_intervalo = "no período analisado"

    return {
        "periodo_intervalo": periodo_intervalo,
        "total_medicos": total_medicos,
        "total_autorizacoes": total_autorizacoes,
        "valor_total": valor_total,
        "media_autorizacoes": media_autorizacoes,
        "media_valor": media_valor,
        "top_crms": top_crms,
        "principal": principal,
        "principal_autorizacoes": principal_autorizacoes,
        "principal_valor": principal_valor,
        "pct_autorizacoes": principal_autorizacoes / total_autorizacoes * 100,
        "pct_valor": (principal_valor / valor_total * 100) if valor_total else 0.0,
        "mult_autorizacoes": principal_autorizacoes / media_autorizacoes,
        "mult_valor": (principal_valor / media_valor) if media_valor else 0.0,
    }


def _build_exclusividade_crm_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
    total_mov_quadro_02: Any = None,
    crm_data: Any = None,
) -> dict[str, Any] | None:
    """Monta o contexto do subitem de CRMs exclusivos."""
    matriz_row: dict[str, Any] = {}
    try:
        df_matriz = get_df_matriz_risco()
        df_matriz = df_matriz.rename({c: c.lower() for c in df_matriz.columns})
        rows = df_matriz.filter(df_matriz["cnpj"] == cnpj)
        if not rows.is_empty():
            matriz_row = rows.row(0, named=True)
    except Exception as exc:
        print(f"[NOTA_TECNICA] Matriz de risco indisponivel para exclusividade CRM {cnpj}: {exc}")

    def matriz_float(key: str) -> float:
        try:
            return float(matriz_row.get(key) or 0)
        except (TypeError, ValueError):
            return 0.0

    if crm_data is None:
        try:
            crm_data = get_crm_data(
                cnpj,
                data_inicio.isoformat() if data_inicio else None,
                data_fim.isoformat() if data_fim else None,
            )
        except Exception as exc:
            print(f"[NOTA_TECNICA] CRM indisponivel para {cnpj}: {exc}")
            crm_data = None

    crms = list(getattr(crm_data, "crms_interesse", None) or [])

    total_medicos = len(crms)
    total_autorizacoes = sum(_as_int(row.get("nu_prescricoes")) for row in crms)
    valor_total = sum(_as_float(row.get("vl_total_prescricoes")) for row in crms)

    def is_exclusivo(row: dict[str, Any]) -> bool:
        return (
            _as_int(row.get("flag_crm_exclusivo")) > 0
            or _as_int(row.get("nu_estabelecimentos")) == 1
            or _as_float(row.get("pct_volume_aqui_vs_total")) >= 99.99
        )

    crms_exclusivos = [row for row in crms if is_exclusivo(row)]

    exclusivos_ordenados = sorted(
        crms_exclusivos,
        key=lambda row: (
            _as_float(row.get("vl_total_prescricoes")),
            _as_int(row.get("nu_prescricoes")),
        ),
        reverse=True,
    )
    top_exclusivos = exclusivos_ordenados[:10]
    exclusivos_autorizacoes = sum(_as_int(row.get("nu_prescricoes")) for row in crms_exclusivos)
    exclusivos_valor = sum(_as_float(row.get("vl_total_prescricoes")) for row in crms_exclusivos)
    total_mov_quadro_02_float = _as_float(total_mov_quadro_02)
    total_financeiro_base = total_mov_quadro_02_float if total_mov_quadro_02_float > 0 else valor_total
    pct_medicos_exclusivos = (
        len(crms_exclusivos) / total_medicos * 100
        if total_medicos > 0
        else matriz_float("pct_exclusividade_crm")
    )
    pct_autorizacoes_exclusivas = (
        exclusivos_autorizacoes / total_autorizacoes * 100
        if total_autorizacoes > 0
        else 0.0
    )
    pct_valor_exclusivo = (
        exclusivos_valor / total_financeiro_base * 100
        if total_financeiro_base > 0 and exclusivos_valor > 0
        else 0.0
    )

    if not matriz_row and not crms_exclusivos:
        return None

    if data_inicio and data_fim:
        periodo_intervalo = f'de {data_inicio.strftime("%d.%m.%Y")} a {data_fim.strftime("%d.%m.%Y")}'
    elif data_inicio:
        periodo_intervalo = f'a partir de {data_inicio.strftime("%d.%m.%Y")}'
    elif data_fim:
        periodo_intervalo = f'até {data_fim.strftime("%d.%m.%Y")}'
    else:
        periodo_intervalo = "no período analisado"

    return {
        "periodo_intervalo": periodo_intervalo,
        "total_medicos": total_medicos,
        "total_autorizacoes": total_autorizacoes,
        "valor_total": valor_total,
        "total_financeiro_base": total_financeiro_base,
        "crms_exclusivos": crms_exclusivos,
        "top_exclusivos": top_exclusivos,
        "qtd_exclusivos": len(crms_exclusivos),
        "exclusivos_autorizacoes": exclusivos_autorizacoes,
        "exclusivos_valor": exclusivos_valor,
        "pct_medicos_exclusivos": pct_medicos_exclusivos,
        "pct_autorizacoes_exclusivas": pct_autorizacoes_exclusivas,
        "pct_valor_exclusivo": pct_valor_exclusivo,
        "matriz_pct_exclusividade": matriz_float("pct_exclusividade_crm"),
        "multiplicador_regiao": matriz_float("risco_exclusividade_crm_reg"),
        "multiplicador_uf": matriz_float("risco_exclusividade_crm_uf"),
        "multiplicador_brasil": matriz_float("risco_exclusividade_crm_br"),
    }


def _build_crms_irregulares_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
    total_mov_quadro_02: Any = None,
    crm_data: Any = None,
) -> dict[str, Any] | None:
    """Monta o contexto do subitem de CRMs irregulares ou invalidos."""
    matriz_row: dict[str, Any] = {}
    try:
        df_matriz = get_df_matriz_risco()
        df_matriz = df_matriz.rename({c: c.lower() for c in df_matriz.columns})
        rows = df_matriz.filter(df_matriz["cnpj"] == cnpj)
        if not rows.is_empty():
            matriz_row = rows.row(0, named=True)
    except Exception as exc:
        print(f"[NOTA_TECNICA] Matriz de risco indisponivel para CRMs irregulares {cnpj}: {exc}")

    def matriz_float(key: str) -> float:
        try:
            return float(matriz_row.get(key) or 0)
        except (TypeError, ValueError):
            return 0.0

    if crm_data is None:
        try:
            crm_data = get_crm_data(
                cnpj,
                data_inicio.isoformat() if data_inicio else None,
                data_fim.isoformat() if data_fim else None,
            )
        except Exception as exc:
            print(f"[NOTA_TECNICA] CRM indisponivel para {cnpj}: {exc}")
            crm_data = None

    crms = list(getattr(crm_data, "crms_interesse", None) or [])
    total_autorizacoes = sum(_as_int(row.get("nu_prescricoes")) for row in crms)
    valor_total_crm = sum(_as_float(row.get("vl_total_prescricoes")) for row in crms)
    total_mov_quadro_02_float = _as_float(total_mov_quadro_02)
    total_financeiro_base = total_mov_quadro_02_float if total_mov_quadro_02_float > 0 else valor_total_crm

    crms_irregulares = [
        row
        for row in crms
        if _as_int(row.get("flag_crm_invalido")) > 0
        or _as_int(row.get("flag_prescricao_antes_registro")) > 0
    ]
    irregulares_ordenados = sorted(
        crms_irregulares,
        key=lambda row: (
            _as_float(row.get("vl_total_prescricoes")),
            _as_int(row.get("nu_prescricoes")),
        ),
        reverse=True,
    )
    top_irregulares = irregulares_ordenados[:10]

    qtd_invalidos = sum(1 for row in crms if _as_int(row.get("flag_crm_invalido")) > 0)
    qtd_antes_registro = sum(1 for row in crms if _as_int(row.get("flag_prescricao_antes_registro")) > 0)
    valor_irregular = sum(_as_float(row.get("vl_total_prescricoes")) for row in crms_irregulares)
    pct_irregular = (
        valor_irregular / total_financeiro_base * 100
        if total_financeiro_base > 0 and valor_irregular > 0
        else matriz_float("pct_crms_irregulares")
    )
    if valor_irregular <= 0 and total_financeiro_base > 0 and pct_irregular > 0:
        valor_irregular = total_financeiro_base * pct_irregular / 100

    if not matriz_row and not crms_irregulares:
        return None

    if data_inicio and data_fim:
        periodo_intervalo = f'de {data_inicio.strftime("%d.%m.%Y")} a {data_fim.strftime("%d.%m.%Y")}'
    elif data_inicio:
        periodo_intervalo = f'a partir de {data_inicio.strftime("%d.%m.%Y")}'
    elif data_fim:
        periodo_intervalo = f'até {data_fim.strftime("%d.%m.%Y")}'
    else:
        periodo_intervalo = "no período analisado"

    return {
        "periodo_intervalo": periodo_intervalo,
        "total_autorizacoes": total_autorizacoes,
        "total_financeiro_base": total_financeiro_base,
        "top_irregulares": top_irregulares,
        "qtd_invalidos": qtd_invalidos,
        "qtd_antes_registro": qtd_antes_registro,
        "valor_irregular": valor_irregular,
        "pct_irregular": pct_irregular,
        "multiplicador_regiao": matriz_float("risco_crms_irregulares_reg"),
        "multiplicador_uf": matriz_float("risco_crms_irregulares_uf"),
        "multiplicador_brasil": matriz_float("risco_crms_irregulares_br"),
    }


def _build_crm_evidencias_complementares_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
    crm_data: Any = None,
) -> dict[str, Any] | None:
    """Detecta evidencias operacionais de CRM independentes da matriz de risco."""
    if crm_data is None:
        try:
            crm_data = get_crm_data(
                cnpj,
                data_inicio.isoformat() if data_inicio else None,
                data_fim.isoformat() if data_fim else None,
            )
        except Exception as exc:
            print(f"[NOTA_TECNICA] Evidencias complementares CRM indisponiveis para {cnpj}: {exc}")
            return None

    crms = list(getattr(crm_data, "crms_interesse", None) or [])
    cnpj_alerts = list(getattr(crm_data, "cnpj_alerts", None) or [])
    if not crms and not cnpj_alerts:
        return None

    principais_crms_contexto = _build_principais_crms_contexto_rows(crms)
    distancia_rows: list[dict[str, Any]] = []
    distancia_medicos: set[str] = set()
    intensiva_rows: list[dict[str, Any]] = []
    intensiva_medicos: set[str] = set()
    crm_unico_rows: list[dict[str, Any]] = []
    crm_unico_medicos: set[str] = set()
    crms_multiplos_rows: list[dict[str, Any]] = []
    crms_multiplos_medicos: set[str] = set()

    for row in crms:
        id_medico = str(row.get("id_medico") or "")

        prescricoes_dia_local = _as_float(row.get("nu_prescricoes_dia"))
        prescricoes_dia_brasil = _as_float(row.get("prescricoes_dia_total_brasil"))
        flag_intensiva_local = _as_int(row.get("flag_robo")) > 0 or prescricoes_dia_local > 30
        flag_intensiva_brasil = _as_int(row.get("flag_robo_oculto")) > 0 or prescricoes_dia_brasil > 30
        if flag_intensiva_local or flag_intensiva_brasil:
            if id_medico:
                intensiva_medicos.add(id_medico)
            if flag_intensiva_local and flag_intensiva_brasil:
                tipo_intensiva = "Local e Brasil"
            elif flag_intensiva_local:
                tipo_intensiva = "Local"
            else:
                tipo_intensiva = "Brasil"
            intensiva_rows.append({
                "id_medico": id_medico or "Não informado",
                "no_medico": row.get("no_medico") or "Não localizado",
                "tipo": tipo_intensiva,
                "nu_prescricoes": _as_int(row.get("nu_prescricoes")),
                "nu_prescricoes_total_brasil": _as_int(row.get("nu_prescricoes_total_brasil")),
                "nu_prescricoes_dia": prescricoes_dia_local,
                "prescricoes_dia_total_brasil": prescricoes_dia_brasil,
                "vl_total_prescricoes": _as_float(row.get("vl_total_prescricoes")),
            })

        for alerta in list(row.get("alertas_crm_unico") or []):
            nu_prescricoes = _as_int(alerta.get("nu_prescricoes"))
            if nu_prescricoes <= 0:
                continue
            nu_minutos_intervalo_raw = alerta.get("nu_minutos_intervalo")
            if id_medico:
                crm_unico_medicos.add(id_medico)
            crm_unico_rows.append({
                "id_medico": id_medico or "Não informado",
                "no_medico": row.get("no_medico") or "Não localizado",
                "dt": alerta.get("dt"),
                "nu_prescricoes": nu_prescricoes,
                "nu_minutos": _as_int(alerta.get("nu_minutos")),
                "nu_minutos_intervalo": (
                    _as_int(nu_minutos_intervalo_raw)
                    if nu_minutos_intervalo_raw is not None
                    else None
                ),
                "taxa_hora": _as_float(alerta.get("taxa_hora")),
                "id_severidade": _as_int(alerta.get("id_severidade")),
                "dt_ini_hora": alerta.get("dt_ini_hora"),
                "dt_fim_hora": alerta.get("dt_fim_hora"),
            })

        for alerta in list(row.get("alertas_crm_multiplos") or []):
            nu_presc_total = _as_int(alerta.get("nu_presc_total"))
            if nu_presc_total <= 0:
                continue
            if id_medico:
                crms_multiplos_medicos.add(id_medico)
            crms_multiplos_rows.append({
                "id_medico": id_medico or "Não informado",
                "no_medico": row.get("no_medico") or "Não localizado",
                "dt": alerta.get("dt"),
                "hr": alerta.get("hr"),
                "nu_presc_crm": _as_int(alerta.get("nu_presc_crm")),
                "nu_presc_total": nu_presc_total,
                "nu_crms_total": _as_int(alerta.get("nu_crms_total")),
                "nu_minutos": _as_int(alerta.get("nu_minutos")),
                "taxa_hora": _as_float(alerta.get("taxa_hora")),
                "dt_ini_hora": alerta.get("dt_ini_hora"),
                "dt_fim_hora": alerta.get("dt_fim_hora"),
                "id_severidade": _as_int(alerta.get("id_severidade")),
                "severidade": alerta.get("severidade") or "ALERTA",
                "descricao": alerta.get("descricao"),
            })

        alertas_geograficos = list(row.get("alertas_geograficos") or [])
        for alerta in alertas_geograficos:
            distancia_km = _as_float(alerta.get("distancia_km"))
            if distancia_km <= 400:
                continue
            if id_medico:
                distancia_medicos.add(id_medico)
            distancia_rows.append({
                "id_medico": id_medico or "Não informado",
                "no_medico": row.get("no_medico") or "Não localizado",
                "competencia": alerta.get("competencia"),
                "cnpj_a": alerta.get("cnpj_a"),
                "municipio_a": alerta.get("municipio_a"),
                "uf_a": alerta.get("uf_a"),
                "dt_ini_a": alerta.get("dt_ini_a"),
                "dt_fim_a": alerta.get("dt_fim_a"),
                "nu_presc_a": _as_int(alerta.get("nu_presc_a")),
                "vl_autorizacoes_a": _as_float(alerta.get("vl_autorizacoes_a")),
                "cnpj_b": alerta.get("cnpj_b"),
                "municipio_b": alerta.get("municipio_b"),
                "uf_b": alerta.get("uf_b"),
                "dt_ini_b": alerta.get("dt_ini_b"),
                "dt_fim_b": alerta.get("dt_fim_b"),
                "nu_presc_b": _as_int(alerta.get("nu_presc_b")),
                "vl_autorizacoes_b": _as_float(alerta.get("vl_autorizacoes_b")),
                "vl_autorizacoes_total": _as_float(alerta.get("vl_autorizacoes_total")),
                "distancia_km": distancia_km,
            })

    intensiva_rows.sort(
        key=lambda item: (
            max(item["nu_prescricoes_dia"], item["prescricoes_dia_total_brasil"]),
            item["vl_total_prescricoes"],
        ),
        reverse=True,
    )
    crm_unico_rows.sort(
        key=lambda item: (item["taxa_hora"], item["nu_prescricoes"], item["id_severidade"]),
        reverse=True,
    )
    crms_multiplos_rows.sort(
        key=lambda item: (item["nu_presc_total"], item["nu_crms_total"]),
        reverse=True,
    )
    distancia_rows.sort(key=lambda item: item["distancia_km"], reverse=True)

    qtd_intensiva_local = sum(1 for row in intensiva_rows if row["nu_prescricoes_dia"] > 30)
    qtd_intensiva_brasil = sum(1 for row in intensiva_rows if row["prescricoes_dia_total_brasil"] > 30)
    qtd_distancia = len(distancia_medicos)
    qtd_crm_unico = len(crm_unico_medicos)
    qtd_crms_multiplos = len(crms_multiplos_medicos)
    intensiva_local_rows = [row for row in intensiva_rows if row["nu_prescricoes_dia"] > 30]
    volume_horario_rows = [
        {
            "dt": alerta.get("dt"),
            "hr": alerta.get("hr"),
            "nu_prescricoes": _as_int(alerta.get("nu_prescricoes")),
            "nu_crms": _as_int(alerta.get("nu_crms")),
            "mediana_hora": _as_optional_float(alerta.get("mediana_hora")),
            "multiplicador": _as_optional_float(alerta.get("multiplicador")),
        }
        for alerta in cnpj_alerts
        if alerta.get("tipo") == "VOLUME"
        and _as_int(alerta.get("nu_prescricoes")) > 0
    ]
    volume_horario_rows.sort(
        key=lambda item: (
            _volume_horario_excesso_value(item) is not None,
            _volume_horario_excesso_value(item) or 0.0,
            item["nu_prescricoes"],
        ),
        reverse=True,
    )
    volume_horario_top_rows = volume_horario_rows[:30]
    surtos_multiplos = [
        alerta
        for alerta in cnpj_alerts
        if alerta.get("tipo") == "MULTIPLO" and _as_int(alerta.get("nu_prescricoes")) > 0
    ]
    principal_crm_por_evento: dict[tuple[str, int], dict[str, Any]] = {}
    for row in crms_multiplos_rows:
        key = (str(row.get("dt") or ""), _as_int(row.get("hr")))
        atual = principal_crm_por_evento.get(key)
        if atual is None or (
            _as_int(row.get("nu_presc_crm")),
            str(row.get("id_medico") or ""),
        ) > (
            _as_int(atual.get("nu_presc_crm")),
            str(atual.get("id_medico") or ""),
        ):
            principal_crm_por_evento[key] = {
                "id_medico": row.get("id_medico"),
                "nu_presc_crm": _as_int(row.get("nu_presc_crm")),
            }
    for alerta in surtos_multiplos:
        key = (str(alerta.get("dt") or "")[:10], _as_int(alerta.get("hr")))
        principal_crm = principal_crm_por_evento.get(key)
        if principal_crm:
            alerta["crm_principal"] = principal_crm.get("id_medico")
            alerta["crm_principal_autorizacoes"] = _as_int(principal_crm.get("nu_presc_crm"))
    if not surtos_multiplos and crms_multiplos_rows:
        eventos_derivados: dict[tuple[str, int], dict[str, Any]] = {}
        for row in crms_multiplos_rows:
            key = (str(row.get("dt") or ""), _as_int(row.get("hr")))
            atual = eventos_derivados.get(key)
            candidato = {
                "tipo": "MULTIPLO",
                "dt": row.get("dt"),
                "hr": row.get("hr"),
                "nu_prescricoes": _as_int(row.get("nu_presc_total")),
                "nu_crms": _as_int(row.get("nu_crms_total")),
                "nu_minutos": _as_int(row.get("nu_minutos")),
                "taxa_hora": _as_float(row.get("taxa_hora")),
                "dt_ini_hora": row.get("dt_ini_hora"),
                "dt_fim_hora": row.get("dt_fim_hora"),
                "crm_principal": row.get("id_medico"),
                "crm_principal_autorizacoes": _as_int(row.get("nu_presc_crm")),
                "id_severidade": _as_int(row.get("id_severidade")),
            }
            if atual is None or (
                _as_float(candidato.get("taxa_hora")),
                _as_int(candidato.get("nu_prescricoes")),
                _as_int(candidato.get("nu_crms")),
            ) > (
                _as_float(atual.get("taxa_hora")),
                _as_int(atual.get("nu_prescricoes")),
                _as_int(atual.get("nu_crms")),
            ):
                eventos_derivados[key] = candidato
        surtos_multiplos = list(eventos_derivados.values())
    surtos_multiplos.sort(
        key=lambda item: (
            _as_float(item.get("taxa_hora")),
            _as_int(item.get("nu_prescricoes")),
            _as_int(item.get("nu_crms")),
        ),
        reverse=True,
    )
    qtd_surtos_multiplos = len(surtos_multiplos)
    crm_unico_por_medico: dict[str, list[dict[str, Any]]] = {}
    for row in crm_unico_rows:
        crm_unico_por_medico.setdefault(str(row.get("id_medico") or ""), []).append(row)
    surtos_multiplos = [
        alerta
        for alerta in surtos_multiplos
        if not any(
            _datetime_windows_overlap(
                alerta.get("dt_ini_hora"),
                alerta.get("dt_fim_hora"),
                unico.get("dt_ini_hora"),
                unico.get("dt_fim_hora"),
            )
            for unico in crm_unico_por_medico.get(str(alerta.get("crm_principal") or ""), [])
        )
    ]
    qtd_surtos_multiplos = len(surtos_multiplos)
    surtos_multiplos_top = _select_crm_concentracao_table_rows(surtos_multiplos)
    if surtos_multiplos_top:
        _enrich_crm_multiplo_valores(cnpj, surtos_multiplos_top)
    crm_unico_top_rows = _select_crm_concentracao_table_rows(crm_unico_rows)
    if crm_unico_top_rows:
        _enrich_crm_unico_valores(cnpj, crm_unico_top_rows)
    crm_unico_top_medicos = {
        str(row.get("id_medico") or "")
        for row in crm_unico_top_rows
        if str(row.get("id_medico") or "")
    }
    surtos_multiplos_top_keys = {_crm_event_key(row) for row in surtos_multiplos_top}
    crms_multiplos_top_rows = [
        row for row in crms_multiplos_rows if _crm_event_key(row) in surtos_multiplos_top_keys
    ]
    crms_multiplos_top_medicos = {
        str(row.get("id_medico") or "")
        for row in crms_multiplos_top_rows
        if str(row.get("id_medico") or "")
    }
    if not crms_multiplos_top_medicos and surtos_multiplos_top:
        qtd_crms_multiplos_top = max((_as_int(row.get("nu_crms")) for row in surtos_multiplos_top), default=0)
    else:
        qtd_crms_multiplos_top = len(crms_multiplos_top_medicos)
    qtd_principais_contexto_alertas = sum(
        1 for row in principais_crms_contexto if row.get("alertas_contexto")
    )

    total_evidencias = (
        qtd_intensiva_local
        + len(volume_horario_top_rows)
        + qtd_distancia
        + len(crm_unico_top_medicos)
        + qtd_crms_multiplos_top
        + len(surtos_multiplos_top)
    )
    if total_evidencias <= 0 and qtd_principais_contexto_alertas <= 0:
        return None

    return {
        "principais_crms_contexto": {
            "rows": principais_crms_contexto,
            "qtd_com_alerta": qtd_principais_contexto_alertas,
        } if principais_crms_contexto else None,
        "qtd_intensiva_local": qtd_intensiva_local,
        "qtd_intensiva_brasil": qtd_intensiva_brasil,
        "qtd_distancia": qtd_distancia,
        "qtd_crm_unico": qtd_crm_unico,
        "qtd_crms_multiplos": qtd_crms_multiplos,
        "qtd_surtos_multiplos": qtd_surtos_multiplos,
        "volume_horario": {
            "qtd_alertas": len(volume_horario_top_rows),
            "maior_multiplicador": next(
                (
                    _volume_horario_excesso_value(row)
                    for row in volume_horario_top_rows
                    if _volume_horario_excesso_value(row) is not None
                ),
                None,
            ),
            "maior_qtd": max((row["nu_prescricoes"] for row in volume_horario_top_rows), default=0),
            "rows": volume_horario_top_rows,
        } if volume_horario_top_rows else None,
        "distancia": {
            "qtd_medicos": qtd_distancia,
            "qtd_alertas": len(distancia_rows),
            "maior_distancia_km": distancia_rows[0]["distancia_km"] if distancia_rows else 0.0,
            "rows": distancia_rows[:10],
        } if qtd_distancia > 0 else None,
        "intensiva": {
            "qtd_medicos": len({row["id_medico"] for row in intensiva_local_rows}),
            "qtd_local": qtd_intensiva_local,
            "qtd_brasil": qtd_intensiva_brasil,
            "maior_media_local": max((row["nu_prescricoes_dia"] for row in intensiva_local_rows), default=0.0),
            "maior_media_brasil": max((row["prescricoes_dia_total_brasil"] for row in intensiva_local_rows), default=0.0),
            "rows": intensiva_local_rows[:10],
        } if qtd_intensiva_local > 0 else None,
        "crm_unico": {
            "qtd_medicos": len(crm_unico_top_medicos),
            "qtd_alertas": len(crm_unico_top_rows),
            "maior_qtd": max((row["nu_prescricoes"] for row in crm_unico_top_rows), default=0),
            "rows": crm_unico_top_rows,
        } if crm_unico_top_rows else None,
        "crms_multiplos": {
            "qtd_medicos": qtd_crms_multiplos_top,
            "qtd_surtos": len(surtos_multiplos_top),
            "qtd_participacoes": len(crms_multiplos_top_rows),
            "maior_qtd": max((_as_int(row.get("nu_prescricoes")) for row in surtos_multiplos_top), default=0),
            "eventos": surtos_multiplos_top,
        } if surtos_multiplos_top else None,
    }


def _add_crm_distancia_complementar_text(
    doc,
    letra: str,
    razao_social: str,
    distancia_comp: dict[str, Any],
):
    """Adiciona o bloco de evidencias de CRMs com distancia geografica superior a 400 km."""
    qtd_medicos = _as_int(distancia_comp.get("qtd_medicos"))
    qtd_alertas = _as_int(distancia_comp.get("qtd_alertas"))
    maior_distancia = _as_float(distancia_comp.get("maior_distancia_km"))
    rows = list(distancia_comp.get("rows") or [])
    distancias_exibidas = {
        int(round(_as_float(row.get("distancia_km"))))
        for row in rows
        if _as_float(row.get("distancia_km")) > 0
    }
    identificado_txt = "foi identificado" if qtd_medicos == 1 else "foram identificados"
    crm_txt = "CRM" if qtd_medicos == 1 else "CRMs"
    evidencia_txt = "evidência geográfica" if qtd_alertas == 1 else "evidências geográficas"

    _add_crm_subheading(doc, f"{letra}) Distância superior a 400 km entre estabelecimentos vinculados ao mesmo CRM")

    p_dist = doc.add_paragraph()
    _run(
        p_dist,
        f"Em relação à Farmácia {razao_social}, {identificado_txt} ",
        color="0F172A",
        size=10,
    )
    _run(p_dist, f"{qtd_medicos}", color="334155", size=10, bold=True)
    _run(
        p_dist,
        f" {crm_txt} com registros de dispensações associados a estabelecimentos situados a mais de 400 km de distância entre si, totalizando ",
        color="0F172A",
        size=10,
    )
    _run(p_dist, f"{qtd_alertas}", color="334155", size=10, bold=True)
    if len(distancias_exibidas) == 1 and qtd_alertas > 1:
        _run(p_dist, f" {evidencia_txt}. A distância observada nas evidências foi de ", color="0F172A", size=10)
    else:
        _run(p_dist, f" {evidencia_txt}. A maior distância observada foi de ", color="0F172A", size=10)
    _run(p_dist, f"{_format_decimal_pt(maior_distancia, 0)} km", color="334155", size=10, bold=True)
    _run(
        p_dist,
        ", situação que sugere a necessidade de verificação quanto à plausibilidade operacional das prescrições vinculadas ao mesmo registro médico.",
        color="0F172A",
        size=10,
    )

    if not rows:
        return

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(
        title,
        "Principais evidências de distância geográfica associadas a CRMs informados no SAV.",
        color="0F172A",
        size=9,
        bold=True,
    )

    headers = ["CRM/UF", "Nome", "Estabelecimento A", "Estabelecimento B", "Valor total", "Distância", "Compet."]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [Inches(0.68), Inches(1.35), Inches(1.65), Inches(1.65), Inches(0.85), Inches(0.52), Inches(0.4)]
    _crm_table_header(table, headers, widths)

    for row in rows:
        cells = table.add_row().cells
        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        local_a = (
            f'{row.get("municipio_a") or "Não informado"}/{row.get("uf_a") or "--"}\n'
            f'CNPJ {row.get("cnpj_a") or "Não informado"} - {row.get("nu_presc_a") or 0} presc.'
        )
        local_b = (
            f'{row.get("municipio_b") or "Não informado"}/{row.get("uf_b") or "--"}\n'
            f'CNPJ {row.get("cnpj_b") or "Não informado"} - {row.get("nu_presc_b") or 0} presc.'
        )
        values = [
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            local_a,
            local_b,
            f'R$ {_format_decimal_pt(_as_float(row.get("vl_autorizacoes_total")), 2)}',
            f'{_format_decimal_pt(_as_float(row.get("distancia_km")), 0)} km',
            _format_competencia_crm(row.get("competencia")),
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx == 4 else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 5, 6) else None
            _write_cell(cells[idx], value, size=6.8, align=align)


def _add_crm_intensiva_complementar_text(
    doc,
    letra: str,
    razao_social: str,
    intensiva_comp: dict[str, Any],
):
    """Adiciona o bloco de prescritores com mais de 30 prescricoes por dia."""
    qtd_medicos = _as_int(intensiva_comp.get("qtd_medicos"))
    qtd_local = _as_int(intensiva_comp.get("qtd_local"))
    qtd_brasil = _as_int(intensiva_comp.get("qtd_brasil"))
    maior_media_local = _as_float(intensiva_comp.get("maior_media_local"))
    maior_media_brasil = _as_float(intensiva_comp.get("maior_media_brasil"))
    rows = list(intensiva_comp.get("rows") or [])

    _add_crm_subheading(doc, f"{letra}) Médicos com mais de 30 prescrições por dia")

    p = doc.add_paragraph()
    _run(
        p,
        f"Na Farmácia {razao_social}, {_plural(qtd_medicos, 'foi identificado', 'foram identificados')} ",
        color="0F172A",
        size=10,
    )
    _run(p, f"{qtd_medicos}", color="334155", size=10, bold=True)
    _run(p, f" {_plural(qtd_medicos, 'CRM com média', 'CRMs com média')} superior a ", color="0F172A", size=10)
    _run(p, "30 prescrições por dia", color="334155", size=10, bold=True)
    _run(
        p,
        " no próprio estabelecimento, indicando uso intensivo do mesmo registro médico nas autorizações do SAV. ",
        color="0F172A",
        size=10,
    )
    if maior_media_local > 0:
        _run(p, "A maior média local observada foi de ", color="0F172A", size=10)
        _run(p, f"{_format_decimal_pt(maior_media_local, 2)} prescrições/dia", color="334155", size=10, bold=True)
        if qtd_brasil > 0 and maior_media_brasil > 0:
            _run(
                p,
                f"; considerando todas as autorizações associadas {_plural(qtd_medicos, 'a esse CRM', 'a esses CRMs')} no Brasil, a maior média chegou a ",
                color="0F172A",
                size=10,
            )
            _run(p, f"{_format_decimal_pt(maior_media_brasil, 2)} prescrições/dia", color="334155", size=10, bold=True)
        _run(p, ".", color="0F172A", size=10)

    if not rows:
        return

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(title, "Principais médicos com volume médio diário superior a 30 prescrições.", color="0F172A", size=9, bold=True)

    headers = ["CRM/UF", "Nome", "Tipo", "Presc./dia local", "Presc./dia Brasil", "Autorizações", "Valor"]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [Inches(0.75), Inches(1.75), Inches(0.75), Inches(0.85), Inches(0.85), Inches(0.9), Inches(1.25)]
    _crm_table_header(table, headers, widths)

    for row in rows:
        cells = table.add_row().cells
        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        values = [
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            str(row.get("tipo") or "—"),
            _format_decimal_pt(_as_float(row.get("nu_prescricoes_dia")), 2),
            _format_decimal_pt(_as_float(row.get("prescricoes_dia_total_brasil")), 2),
            str(_as_int(row.get("nu_prescricoes"))),
            f'R$ {_format_decimal_pt(_as_float(row.get("vl_total_prescricoes")), 2)}',
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else None
            _write_cell(cells[idx], value, size=6.8, align=align)


def _add_crm_unico_complementar_text(
    doc,
    letra: str,
    razao_social: str,
    crm_unico_comp: dict[str, Any],
):
    """Adiciona o bloco de concentracao temporal em um unico CRM."""
    qtd_medicos = _as_int(crm_unico_comp.get("qtd_medicos"))
    qtd_alertas = _as_int(crm_unico_comp.get("qtd_alertas"))
    maior_qtd = _as_int(crm_unico_comp.get("maior_qtd"))
    rows = list(crm_unico_comp.get("rows") or [])
    principal = rows[0] if rows else {}
    principal_intervalo = None
    if principal:
        principal_intervalo = principal.get("nu_minutos_intervalo")
        if principal_intervalo is None:
            principal_intervalo = principal.get("nu_minutos")

    _add_crm_subheading(doc, f"{letra}) Autorizações concentradas para um único CRM")

    p = doc.add_paragraph()
    _run(p, f"{_plural(qtd_alertas, 'Foi identificado', 'Foram identificados')} ", color="0F172A", size=10)
    _run(p, f"{qtd_alertas}", color="334155", size=10, bold=True)
    _run(
        p,
        f" {_plural(qtd_alertas, 'episódio', 'episódios')} em que autorizações vinculadas a um mesmo CRM foram registradas em curto intervalo de tempo, envolvendo ",
        color="0F172A",
        size=10,
    )
    _run(p, f"{qtd_medicos}", color="334155", size=10, bold=True)
    _run(
        p,
        f" {_plural(qtd_medicos, 'médico', 'médicos')}. Esse padrão indica concentração incomum de lançamentos no SAV e pode sugerir uso sequencial {_plural(qtd_medicos, 'desse registro profissional', 'desses registros profissionais')} em operações de balcão. ",
        color="0F172A",
        size=10,
    )
    _run(p, "O episódio de maior ritmo observado concentrou ", color="0F172A", size=10)
    _run(p, f"{_as_int(principal.get('nu_prescricoes')) or maior_qtd}", color="334155", size=10, bold=True)
    ritmo_principal = _as_float(principal.get("taxa_hora"))
    _run(
        p,
        (
            f" autorizações em {_format_janela_minutos(principal_intervalo) if principal_intervalo is not None else 'intervalo reduzido'}, "
            f"o que corresponde a um ritmo de {_format_decimal_pt(ritmo_principal, 1)} autorizações por hora."
        ),
        color="0F172A",
        size=10,
    )

    if not rows:
        return

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(title, "Principais episódios de autorizações concentradas para um único CRM.", color="0F172A", size=9, bold=True)

    headers = ["Início", "Fim", "CRM/UF", "Nome", "Autorizações", "Intervalo", "Taxa/hora", "Valor"]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [Inches(1.08), Inches(1.08), Inches(0.65), Inches(2.37), Inches(0.42), Inches(0.45), Inches(0.45), Inches(0.6)]
    _crm_table_header(table, headers, widths)

    for row in rows:
        cells = table.add_row().cells
        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        intervalo = row.get("nu_minutos_intervalo")
        values = [
            _format_datetime_br_minute(row.get("dt_ini_hora") or row.get("dt")),
            _format_datetime_br_minute(row.get("dt_fim_hora") or row.get("dt")),
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            str(_as_int(row.get("nu_prescricoes"))),
            _format_janela_minutos(intervalo if intervalo is not None else row.get("nu_minutos")),
            f'{_format_decimal_pt(_as_float(row.get("taxa_hora")), 1)}/h',
            f'R$ {_format_decimal_pt(_as_float(row.get("valor_alerta")), 2)}'
            if row.get("valor_alerta_disponivel")
            else "N/d",
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (4, 6, 7) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 2, 5) else None
            _write_cell(cells[idx], value, size=6.8, align=align)


def _add_crms_multiplos_complementar_text(
    doc,
    letra: str,
    razao_social: str,
    crms_multiplos_comp: dict[str, Any],
):
    """Adiciona o bloco de concentracao temporal envolvendo multiplos CRMs."""
    qtd_medicos = _as_int(crms_multiplos_comp.get("qtd_medicos"))
    qtd_surtos = _as_int(crms_multiplos_comp.get("qtd_surtos"))
    maior_qtd = _as_int(crms_multiplos_comp.get("maior_qtd"))
    eventos = list(crms_multiplos_comp.get("eventos") or [])
    principal = eventos[0] if eventos else {}

    _add_crm_subheading(doc, f"{letra}) Autorizações concentradas envolvendo múltiplos CRMs")

    p = doc.add_paragraph()
    _run(
        p,
        f"No nível do estabelecimento, {_plural(qtd_surtos, 'foi identificado', 'foram identificados')} ",
        color="0F172A",
        size=10,
    )
    _run(p, f"{qtd_surtos}", color="334155", size=10, bold=True)
    _run(
        p,
        f" {_plural(qtd_surtos, 'episódio', 'episódios')} em que autorizações associadas a diferentes CRMs foram registradas em curto intervalo, com participação de ",
        color="0F172A",
        size=10,
    )
    _run(p, f"{qtd_medicos}", color="334155", size=10, bold=True)
    _run(
        p,
        f" {_plural(qtd_medicos, 'médico', 'médicos')} entre os principais prescritores analisados. Esse comportamento sugere concentração operacional de lançamentos e pode indicar processamento sequencial de autorizações com alternância de registros profissionais. ",
        color="0F172A",
        size=10,
    )
    if maior_qtd > 0:
        _run(p, "O principal episódio reuniu ", color="0F172A", size=10)
        _run(p, f"{_as_int(principal.get('nu_prescricoes')) or maior_qtd}", color="334155", size=10, bold=True)
        _run(p, " autorizações", color="0F172A", size=10)
        if principal:
            _run(p, ", envolvendo ", color="0F172A", size=10)
            _run(p, f"{_as_int(principal.get('nu_crms'))}", color="334155", size=10, bold=True)
            _run(
                p,
                f" CRMs distintos em {_format_janela_minutos(principal.get('nu_minutos'))}.",
                color="0F172A",
                size=10,
            )
        else:
            _run(p, ".", color="0F172A", size=10)

    if eventos:
        title = doc.add_paragraph()
        _format_crm_table_title(title)
        _run(title, "Principais episódios de autorizações concentradas envolvendo múltiplos CRMs.", color="0F172A", size=9, bold=True)

        headers = ["Início", "Fim", "CRMs", "CRM mais usado", "Autorizações", "Intervalo", "Taxa/hora", "Valor"]
        table = doc.add_table(rows=1, cols=len(headers))
        widths = [Inches(1.0), Inches(1.0), Inches(0.48), Inches(1.0), Inches(0.68), Inches(0.66), Inches(0.9), Inches(1.38)]
        _crm_table_header(table, headers, widths)
        for evento in eventos:
            cells = table.add_row().cells
            crm_principal = str(evento.get("crm_principal") or "N/d")
            crm_principal_aut = _as_int(evento.get("crm_principal_autorizacoes"))
            crm_principal_txt = (
                f"{crm_principal} ({crm_principal_aut})"
                if crm_principal_aut > 0 and crm_principal != "N/d"
                else crm_principal
            )
            values = [
                _format_datetime_br_minute(evento.get("dt_ini_hora") or evento.get("dt")),
                _format_datetime_br_minute(evento.get("dt_fim_hora") or evento.get("dt")),
                str(_as_int(evento.get("nu_crms"))),
                crm_principal_txt,
                str(_as_int(evento.get("nu_prescricoes"))),
                _format_janela_minutos(evento.get("nu_minutos")),
                f'{_format_decimal_pt(_as_float(evento.get("taxa_hora")), 1)}/h',
                f'R$ {_format_decimal_pt(_as_float(evento.get("valor_alerta")), 2)}'
                if evento.get("valor_alerta_disponivel")
                else "N/d",
            ]
            for idx, value in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (2, 4, 6, 7) else WD_ALIGN_PARAGRAPH.CENTER
                _write_cell(cells[idx], value, size=6.8, align=align)


def _add_principais_crms_contexto_text(
    doc,
    razao_social: str,
    principais_comp: dict[str, Any],
):
    """Adiciona uma tabela-resumo dos principais CRMs por valor autorizado."""
    rows = list(principais_comp.get("rows") or [])
    if not rows:
        return

    qtd_com_alerta = _as_int(principais_comp.get("qtd_com_alerta"))

    p = doc.add_paragraph()
    _run(
        p,
        f"Para contextualizar os achados, o quadro a seguir apresenta os 10 principais CRMs utilizados pela Farmácia {razao_social}, ordenados pelo valor autorizado no SAV.",
        color="0F172A",
        size=10,
    )
    if qtd_com_alerta > 0:
        _run(
            p,
            " CRMs fora do top 10 também são exibidos quando possuem alertas ou anomalias relevantes para a análise; nesses casos, as linhas aparecem destacadas.",
            color="0F172A",
            size=10,
        )

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(title, "Principais CRMs do estabelecimento por valor autorizado.", color="0F172A", size=9, bold=True)

    headers = [
        "CRM/UF",
        "Nome",
        "Alertas/anomalias",
        "Autorizações",
        "Valor autorizado",
        "Part.",
        "Presc./dia",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [
        Inches(0.7),
        Inches(1.45),
        Inches(1.35),
        Inches(0.68),
        Inches(1.0),
        Inches(0.62),
        Inches(1.3),
    ]
    _crm_table_header(table, headers, widths)

    for row in rows:
        cells = table.add_row().cells
        alertas = list(row.get("alertas_contexto") or [])
        if alertas:
            for cell in cells:
                _cell_bg(cell, "FFF7ED")

        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        presc_dia = (
            f'{_format_decimal_pt(_as_float(row.get("nu_prescricoes_dia")), 2)} local\n'
            f'{_format_decimal_pt(_as_float(row.get("prescricoes_dia_total_brasil")), 2)} Brasil'
        )
        values = [
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            "; ".join(alertas) if alertas else "Sem alerta",
            str(_as_int(row.get("nu_prescricoes"))),
            f'R$ {_format_decimal_pt(_as_float(row.get("vl_total_prescricoes")), 2)}',
            f'{_format_decimal_pt(_as_float(row.get("pct_participacao")), 2)}%',
            presc_dia,
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (3, 4, 5) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 6) else None
            _write_cell(cells[idx], value, size=6.6, align=align)


def _add_crm_volume_horario_complementar_text(
    doc,
    letra: str,
    razao_social: str,
    volume_comp: dict[str, Any],
):
    """Adiciona o bloco de volume horario anomalo de autorizacoes."""
    qtd_alertas = _as_int(volume_comp.get("qtd_alertas"))
    rows = list(volume_comp.get("rows") or [])
    principal = rows[0] if rows else {}

    _add_crm_subheading(doc, f"{letra}) Volume horário anômalo de autorizações")

    p = doc.add_paragraph()
    _run(
        p,
        f"Na Farmácia {razao_social}, {_plural(qtd_alertas, 'foi identificado', 'foram identificados')} ",
        color="0F172A",
        size=10,
    )
    _run(p, f"{qtd_alertas}", color="334155", size=10, bold=True)
    _run(
        p,
        f" {_plural(qtd_alertas, 'horário', 'horários')} com volume de autorizações acima do padrão histórico do próprio estabelecimento para a mesma faixa horária. A tabela exibe ",
        color="0F172A",
        size=10,
    )
    qtd_exibida = len(rows)
    _run(
        p,
        f"{qtd_exibida} {_plural(qtd_exibida, 'principal horário', 'principais horários')}",
        color="334155",
        size=10,
        bold=True,
    )
    _run(
        p,
        ", ordenados pelo afastamento em relação à mediana. ",
        color="0F172A",
        size=10,
    )
    if principal:
        _run(p, "O principal horário concentrou ", color="0F172A", size=10)
        _run(p, f"{_as_int(principal.get('nu_prescricoes'))}", color="334155", size=10, bold=True)
        _run(
            p,
            f" autorizações em {_format_date_br(principal.get('dt'))}, às {_format_time_hour(principal.get('hr'))}, volume ",
            color="0F172A",
            size=10,
        )
        multiplicador_principal = _as_optional_float(principal.get("multiplicador"))
        mediana_principal = _as_optional_float(principal.get("mediana_hora"))
        excesso_principal = _volume_horario_excesso_value(principal)
        if multiplicador_principal is not None:
            _run(p, f"{_format_decimal_pt(multiplicador_principal, 1)} vezes", color="334155", size=10, bold=True)
        elif excesso_principal is not None:
            _run(p, f"{_format_decimal_pt(excesso_principal, 1)} vezes", color="334155", size=10, bold=True)
        else:
            _run(p, "sem comparação calculável", color="334155", size=10, bold=True)
        _run(p, ", com mediana histórica de ", color="0F172A", size=10)
        mediana_principal = _as_optional_float(principal.get("mediana_hora"))
        _run(
            p,
            f"{_format_optional_decimal_pt(mediana_principal, 1)} {_plural(int(round(mediana_principal or 0.0)), 'autorização', 'autorizações')}",
            color="334155",
            size=10,
            bold=True,
        )
        _run(p, " para esse horário.", color="0F172A", size=10)

    if not rows:
        return

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(title, "Principais horários com volume anômalo de autorizações.", color="0F172A", size=9, bold=True)

    headers = ["Data", "Hora", "Autorizações", "CRMs", "Mediana", "Excesso"]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [Inches(1.0), Inches(0.75), Inches(1.25), Inches(0.75), Inches(1.35), Inches(2.0)]
    _crm_table_header(table, headers, widths)

    for row in rows:
        cells = table.add_row().cells
        values = [
            _format_date_br(row.get("dt")),
            _format_hour_range(row.get("hr")),
            str(_as_int(row.get("nu_prescricoes"))),
            str(_as_int(row.get("nu_crms"))),
            _format_optional_decimal_pt(row.get("mediana_hora"), 1),
            _format_volume_horario_excesso(row),
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (2, 3, 4, 5) else WD_ALIGN_PARAGRAPH.CENTER
            _write_cell(cells[idx], value, size=6.8, align=align)


def _add_crm_evidencias_complementares_text(
    doc,
    num: str,
    razao_social: str,
    evidencias_comp: dict[str, Any],
):
    """Adiciona o esqueleto do subitem de evidencias complementares de CRM."""
    doc.add_heading(f"{num} Evidências complementares relacionadas ao uso de CRMs no SAV", level=2)

    p_intro = doc.add_paragraph()
    _run(
        p_intro,
        "Além dos indicadores críticos apresentados nos subitens anteriores, foram identificadas evidências operacionais complementares relacionadas ao uso de CRMs no SAV. Tais evidências não correspondem, isoladamente, a indicadores da matriz de risco, mas ajudam a contextualizar a dinâmica dos lançamentos realizados pelo estabelecimento.",
        color="0F172A",
        size=10,
    )

    principais_crms_comp = evidencias_comp.get("principais_crms_contexto")
    if principais_crms_comp:
        _add_principais_crms_contexto_text(doc, razao_social, principais_crms_comp)

    letras = iter("abcdefghijklmnopqrstuvwxyz")

    volume_horario_comp = evidencias_comp.get("volume_horario")
    if volume_horario_comp:
        _add_crm_volume_horario_complementar_text(doc, next(letras), razao_social, volume_horario_comp)
    distancia_comp = evidencias_comp.get("distancia")
    if distancia_comp:
        _add_crm_distancia_complementar_text(doc, next(letras), razao_social, distancia_comp)
    intensiva_comp = evidencias_comp.get("intensiva")
    if intensiva_comp:
        _add_crm_intensiva_complementar_text(doc, next(letras), razao_social, intensiva_comp)
    crm_unico_comp = evidencias_comp.get("crm_unico")
    if crm_unico_comp:
        _add_crm_unico_complementar_text(doc, next(letras), razao_social, crm_unico_comp)
    crms_multiplos_comp = evidencias_comp.get("crms_multiplos")
    if crms_multiplos_comp:
        _add_crms_multiplos_complementar_text(doc, next(letras), razao_social, crms_multiplos_comp)


def _add_hhi_crm_text(doc, num: str, razao_social: str, cnpj_fmt: str, hhi_crm_comp: dict[str, Any]):
    """Adiciona o subitem de concentração atípica de registros do mesmo CRM."""
    periodo_intervalo = hhi_crm_comp["periodo_intervalo"]
    total_medicos = hhi_crm_comp["total_medicos"]
    total_autorizacoes = hhi_crm_comp["total_autorizacoes"]
    valor_total = hhi_crm_comp["valor_total"]
    media_autorizacoes = hhi_crm_comp["media_autorizacoes"]
    media_valor = hhi_crm_comp["media_valor"]
    principal = hhi_crm_comp["principal"]
    principal_autorizacoes = hhi_crm_comp["principal_autorizacoes"]
    principal_valor = hhi_crm_comp["principal_valor"]
    top_autorizacoes = sum(_as_int(row.get("nu_prescricoes")) for row in hhi_crm_comp["top_crms"])
    top_pct_autorizacoes = (top_autorizacoes / total_autorizacoes * 100) if total_autorizacoes else 0.0

    crm_num, crm_uf = _crm_num_uf(principal.get("id_medico"))
    crm_ident = f"{crm_num}/{crm_uf}" if crm_uf else crm_num
    nome_medico = str(principal.get("no_medico") or "Não localizado")

    doc.add_heading(
        f"{num} Concentração atípica de registros do mesmo médico (CRM) no Sistema Autorizador de Vendas do PFPB",
        level=2,
    )

    p1 = doc.add_paragraph()
    _run(
        p1,
        "Dentre os dados lançados pelas farmácias credenciadas no PFPB no Sistema Autorizador de Vendas (SAV) está o número de inscrição do médico no Conselho Regional de Medicina (CRM) e sua respectiva unidade federativa, a fim de respaldar as dispensações de medicamentos por meio das prescrições (receitas). ",
        color="0F172A",
        size=10,
    )
    _run(
        p1,
        "O comportamento esperado para os estabelecimentos é de que diversos pacientes apresentem receitas de médicos distintos. Concentração excessiva pode indicar a ocorrência de acordos irregulares relacionados à prescrição de receitas, atuação direcionada de prescritores junto ao estabelecimento e/ou uso indevido de CRMs por parte da farmácia.",
        color="0F172A",
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f"No período {periodo_intervalo}, foram identificados ", color="0F172A", size=10)
    _run(p2, f"{total_medicos}", color="334155", size=10, bold=True)
    _run(
        p2,
        f" médicos lançados pela Farmácia {razao_social} como responsáveis pelas receitas prescritas de medicamentos supostamente retirados no estabelecimento. O quadro a seguir apresenta os principais CRMs por valor pago, com indicação da participação individual e acumulada de cada um na produção total da farmácia, observado o mínimo de 5 e o máximo de 10 médicos:",
        color="0F172A",
        size=10,
    )

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(
        title,
        f"Quadro 07 - Médicos/CRMs com maiores valores pagos pelo PFPB em vendas lançadas pela Farmácia {razao_social} (CNPJ {cnpj_fmt}) no Sistema Autorizador de Vendas, no período {periodo_intervalo}.",
        color="0F172A",
        size=9,
        bold=True,
    )

    headers = [
        "CRM/UF",
        "Nome",
        "Data da inscrição no CFM",
        "Número de autorizações vinculadas ao CRM",
        "% sobre a produção total da farmácia",
        "% acumulado da produção total",
        "Valor total pago pelo PFPB tendo como base o CRM",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.8), Inches(1.8), Inches(0.75), Inches(0.95), Inches(0.85), Inches(0.85), Inches(1.1)]
    _set_table_fixed_widths(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _cell_bg(cell, "E2E8F0")
        _write_cell(cell, header, size=7.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    autorizacoes_acumuladas_tabela = 0
    for row in hhi_crm_comp["top_crms"]:
        cells = table.add_row().cells
        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        row_autorizacoes = _as_int(row.get("nu_prescricoes"))
        autorizacoes_acumuladas_tabela += row_autorizacoes
        pct_producao_total = (row_autorizacoes / total_autorizacoes * 100) if total_autorizacoes else 0.0
        pct_acumulado = (autorizacoes_acumuladas_tabela / total_autorizacoes * 100) if total_autorizacoes else 0.0
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        values = [
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            _format_date_br(row.get("dt_inscricao_crm")),
            str(row_autorizacoes),
            f'{_format_decimal_pt(pct_producao_total, 2)}%',
            f'{_format_decimal_pt(pct_acumulado, 2)}%',
            f'R$ {_format_decimal_pt(_as_float(row.get("vl_total_prescricoes")), 2)}',
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else None
            _write_cell(cells[idx], value, size=7.0, align=align)

    fonte = doc.add_paragraph()
    _format_crm_table_footnote(fonte)
    _run(
        fonte,
        "Fonte: Consulta ao CFM (https://portal.cfm.org.br/busca-medicos) e Sistema Autorizador de Vendas (SAV).",
        color="475569",
        size=8,
        italic=True,
    )

    p3 = doc.add_paragraph()
    _run(p3, "Conforme o Quadro 07, observa-se concentração relevante das dispensações no médico ", color="0F172A", size=10)
    _run(p3, nome_medico, color="334155", size=10, bold=True)
    _run(p3, ", CRM ", color="0F172A", size=10)
    _run(p3, crm_ident, color="334155", size=10, bold=True)
    _run(p3, ". No período analisado, esse CRM concentrou ", color="0F172A", size=10)
    _run(p3, f"{principal_autorizacoes}", color="334155", size=10, bold=True)
    _run(p3, f" das {total_autorizacoes} autorizações verificadas, correspondendo a ", color="0F172A", size=10)
    _run(p3, f'{_format_decimal_pt(hhi_crm_comp["pct_autorizacoes"], 2)}%', color="334155", size=10, bold=True)
    _run(
        p3,
        f" da produção da farmácia. Embora tenham sido identificados {total_medicos} médicos no período, com média de {_format_decimal_pt(media_autorizacoes, 2)} autorizações por CRM, o volume associado ao CRM {crm_ident} foi ",
        color="0F172A",
        size=10,
    )
    mult_autorizacoes = _as_float(hhi_crm_comp["mult_autorizacoes"])
    mult_autorizacoes_fmt = _format_decimal_pt(mult_autorizacoes, 2)
    _run(p3, f'{mult_autorizacoes_fmt} {_vez_ou_vezes(mult_autorizacoes)} essa média', color="334155", size=10, bold=True)
    _run(p3, ".", color="0F172A", size=10)

    p4 = doc.add_paragraph()
    _run(p4, f"Em termos financeiros, as vendas vinculadas ao referido CRM somaram ", color="0F172A", size=10)
    _run(p4, f"R$ {_format_decimal_pt(principal_valor, 2)}", color="334155", size=10, bold=True)
    _run(p4, ", equivalentes a ", color="0F172A", size=10)
    _run(p4, f'{_format_decimal_pt(hhi_crm_comp["pct_valor"], 2)}%', color="334155", size=10, bold=True)
    _run(
        p4,
        f" dos R$ {_format_decimal_pt(valor_total, 2)} analisados. Esse montante também se mostra destoante da distribuição média por médico, uma vez que supera em ",
        color="0F172A",
        size=10,
    )
    _run(p4, f'{_format_decimal_pt(hhi_crm_comp["mult_valor"], 2)} vezes', color="334155", size=10, bold=True)
    _run(p4, f" a média de R$ {_format_decimal_pt(media_valor, 2)} por CRM. A coincidência entre concentração de autorizações e concentração de valores reforça o caráter atípico do padrão observado.", color="0F172A", size=10)

    if top_pct_autorizacoes >= 80:
        p5 = doc.add_paragraph()
        _run(p5, "Ademais, os CRMs listados no Quadro 07 concentram conjuntamente ", color="0F172A", size=10)
        _run(p5, f"{_format_decimal_pt(top_pct_autorizacoes, 2)}%", color="334155", size=10, bold=True)
        _run(p5, " da produção total da farmácia, alcançando o patamar de concentração definido para a seleção do quadro e indicando que a dispersão esperada entre prescritores não se verificou de forma regular no período analisado.", color="0F172A", size=10)


def _add_crms_irregulares_text(doc, num: str, razao_social: str, cnpj_fmt: str, irregulares_comp: dict[str, Any]):
    """Adiciona o subitem de vendas vinculadas a CRMs irregulares."""
    periodo_intervalo = irregulares_comp["periodo_intervalo"]
    total_autorizacoes = irregulares_comp["total_autorizacoes"]
    total_financeiro_base = irregulares_comp["total_financeiro_base"]
    qtd_invalidos = irregulares_comp["qtd_invalidos"]
    qtd_antes_registro = irregulares_comp["qtd_antes_registro"]
    valor_irregular = irregulares_comp["valor_irregular"]
    pct_irregular = irregulares_comp["pct_irregular"]
    has_detalhe_irregulares = bool(irregulares_comp["top_irregulares"])
    multiplicador_reg_fmt = _format_decimal_pt(irregulares_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(irregulares_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(irregulares_comp["multiplicador_brasil"], 2)
    multiplicador_reg_unidade = _vez_ou_vezes(_as_float(irregulares_comp["multiplicador_regiao"]))
    multiplicador_uf_unidade = _vez_ou_vezes(_as_float(irregulares_comp["multiplicador_uf"]))
    multiplicador_br_unidade = _vez_ou_vezes(_as_float(irregulares_comp["multiplicador_brasil"]))

    doc.add_heading(f"{num} Vendas de medicamentos prescritos por médicos com irregularidade em seus CRMs", level=2)

    p1 = doc.add_paragraph()
    _run(
        p1,
        "No âmbito do PFPB, as dispensações devem estar respaldadas por prescrições emitidas por médicos com registro ativo e regular no Conselho Regional de Medicina (CRM). Para este indicador, foram consideradas duas situações de irregularidade: CRMs inválidos ou não localizados na base do Conselho Federal de Medicina (CFM), e prescrições com data anterior à inscrição do médico no respectivo conselho. A ocorrência de qualquer dessas situações aponta para o processamento de dispensações com prescrição médica incompatível com os requisitos legais do Programa.",
        color="0F172A",
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f"Em relação à Farmácia {razao_social}, verificou-se que, do total de ", color="0F172A", size=10)
    _run(p2, f"R$ {_format_decimal_pt(total_financeiro_base, 2)}", color="334155", size=10, bold=True)
    _run(p2, f" em vendas de medicamentos efetivadas no âmbito do PFPB no período {periodo_intervalo}, ", color="0F172A", size=10)
    _run(p2, f"{_format_decimal_pt(pct_irregular, 2)}%", color="334155", size=10, bold=True)
    _run(p2, " (", color="0F172A", size=10)
    _run(p2, f"R$ {_format_decimal_pt(valor_irregular, 2)}", color="334155", size=10, bold=True)
    _run(p2, ") foram realizadas com receitas prescritas por médicos com CRMs irregulares ou inválidos", color="0F172A", size=10)
    if has_detalhe_irregulares:
        _run(p2, ", sendo ", color="0F172A", size=10)
        _run(p2, f"{qtd_invalidos}", color="334155", size=10, bold=True)
        _run(p2, " com números inválidos e ", color="0F172A", size=10)
        _run(p2, f"{qtd_antes_registro}", color="334155", size=10, bold=True)
        _run(p2, " com prescrição médica emitida antes do registro do CRM", color="0F172A", size=10)
    _run(p2, ". Tal percentual corresponde a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_reg_fmt} {multiplicador_reg_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o percentual mediano de vendas com essa mesma criticidade entre as farmácias de sua região. Ampliando-se o comparativo geográfico, o percentual equivale a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_uf_fmt} {multiplicador_uf_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o percentual mediano das farmácias de seu Estado e a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_br_fmt} {multiplicador_br_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o das farmácias de todo o Brasil.", color="0F172A", size=10)

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(
        title,
        f"Quadro 08 - Médicos com CRM irregular ou inválido vinculados a vendas lançadas pela Farmácia {razao_social} (CNPJ {cnpj_fmt}) no Sistema Autorizador de Vendas, no período {periodo_intervalo}.",
        color="0F172A",
        size=9,
        bold=True,
    )

    headers = [
        "CRM/UF",
        "Nome",
        "Data da inscrição no CFM",
        "Irregularidade identificada",
        "Número de autorizações vinculadas ao CRM",
        "% sobre a produção total da farmácia",
        "Valor total pago pelo PFPB tendo como base o CRM",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.75), Inches(1.65), Inches(0.75), Inches(1.1), Inches(0.9), Inches(0.8), Inches(1.15)]
    _set_table_fixed_widths(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _cell_bg(cell, "E2E8F0")
        _write_cell(cell, header, size=7.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    linhas_irregulares = irregulares_comp["top_irregulares"] if has_detalhe_irregulares else [None]
    for row in linhas_irregulares:
        cells = table.add_row().cells
        if row is None:
            values = [
                "Consolidado",
                "Detalhamento por CRM não disponível",
                "Não localizada",
                "CRMs irregulares ou inválidos",
                "Não disponível",
                f"{_format_decimal_pt(pct_irregular, 2)}%",
                f"R$ {_format_decimal_pt(valor_irregular, 2)}",
            ]
        else:
            crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
            crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
            row_autorizacoes = _as_int(row.get("nu_prescricoes"))
            row_valor = _as_float(row.get("vl_total_prescricoes"))
            pct_producao_total = (row_autorizacoes / total_autorizacoes * 100) if total_autorizacoes else 0.0
            motivos = []
            if _as_int(row.get("flag_crm_invalido")) > 0:
                motivos.append("CRM inválido")
            if _as_int(row.get("flag_prescricao_antes_registro")) > 0:
                motivos.append("Prescrição antes do registro")
            values = [
                crm_uf,
                str(row.get("no_medico") or "Não localizado"),
                _format_date_br(row.get("dt_inscricao_crm")),
                "; ".join(motivos) or "Irregularidade CRM",
                str(row_autorizacoes),
                f"{_format_decimal_pt(pct_producao_total, 2)}%",
                f"R$ {_format_decimal_pt(row_valor, 2)}",
            ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (4, 5, 6) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else None
            _write_cell(cells[idx], value, size=7.0, align=align)

    fonte = doc.add_paragraph()
    _format_crm_table_footnote(fonte)
    _run(
        fonte,
        "Fonte: Consulta ao CFM (https://portal.cfm.org.br/busca-medicos) e Sistema Autorizador de Vendas (SAV).",
        color="475569",
        size=8,
        italic=True,
    )


def _add_exclusividade_crm_text(doc, num: str, razao_social: str, cnpj_fmt: str, exclusividade_comp: dict[str, Any]):
    """Adiciona o subitem de CRMs exclusivos."""
    periodo_intervalo = exclusividade_comp["periodo_intervalo"]
    total_medicos = exclusividade_comp["total_medicos"]
    total_autorizacoes = exclusividade_comp["total_autorizacoes"]
    total_financeiro_base = exclusividade_comp["total_financeiro_base"]
    qtd_exclusivos = exclusividade_comp["qtd_exclusivos"]
    exclusivos_valor = exclusividade_comp["exclusivos_valor"]
    pct_medicos_exclusivos = exclusividade_comp["pct_medicos_exclusivos"]
    pct_valor_exclusivo = exclusividade_comp["pct_valor_exclusivo"]
    has_detalhe_exclusivos = bool(exclusividade_comp["top_exclusivos"])
    multiplicador_reg_fmt = _format_decimal_pt(exclusividade_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(exclusividade_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(exclusividade_comp["multiplicador_brasil"], 2)
    multiplicador_reg_unidade = _vez_ou_vezes(_as_float(exclusividade_comp["multiplicador_regiao"]))
    multiplicador_uf_unidade = _vez_ou_vezes(_as_float(exclusividade_comp["multiplicador_uf"]))
    multiplicador_br_unidade = _vez_ou_vezes(_as_float(exclusividade_comp["multiplicador_brasil"]))
    crm_exclusivo_desc = "CRM de médico" if qtd_exclusivos == 1 else "CRMs de médicos"
    crm_exclusivo_quant = "CRM exclusivo" if qtd_exclusivos == 1 else "CRMs exclusivos"
    identificados_intro = "Foi identificado" if qtd_exclusivos == 1 else "Foram identificados"
    equivalente_txt = "equivalente" if qtd_exclusivos == 1 else "equivalentes"

    doc.add_heading(
        f"{num} Vendas de medicamentos vinculadas a CRMs registrados exclusivamente pela Farmácia {razao_social} no Sistema Autorizador de Vendas do PFPB",
        level=2,
    )

    p1 = doc.add_paragraph()
    _run(
        p1,
        "No âmbito do PFPB, espera-se que médicos prescritores tenham dispensações registradas em mais de uma farmácia, em razão da diversidade esperada de pacientes e estabelecimentos. A identificação de retiradas de medicamentos associadas a um médico em apenas um estabelecimento farmacêutico pode indicar atuação direcionada junto à farmácia ou uso indevido de CRM pelo estabelecimento.",
        color="0F172A",
        size=10,
    )

    p2 = doc.add_paragraph()
    if has_detalhe_exclusivos:
        _run(p2, f"Em relação à Farmácia {razao_social}, verificou-se que, do total de ", color="0F172A", size=10)
        _run(p2, f"R$ {_format_decimal_pt(total_financeiro_base, 2)}", color="334155", size=10, bold=True)
        _run(p2, f" em vendas de medicamentos efetivadas no âmbito do PFPB no período {periodo_intervalo}, ", color="0F172A", size=10)
        _run(p2, f"R$ {_format_decimal_pt(exclusivos_valor, 2)}", color="334155", size=10, bold=True)
        _run(p2, ", equivalente a ", color="0F172A", size=10)
        _run(p2, f"{_format_decimal_pt(pct_valor_exclusivo, 2)}%", color="334155", size=10, bold=True)
        _run(p2, f", esteve associado a {crm_exclusivo_desc} cujos clientes retiraram seus medicamentos exclusivamente nesse estabelecimento. {identificados_intro} ", color="0F172A", size=10)
        _run(p2, f"{qtd_exclusivos}", color="334155", size=10, bold=True)
        _run(p2, f" {crm_exclusivo_quant}, {equivalente_txt} a ", color="0F172A", size=10)
    else:
        _run(p2, f"Em relação à Farmácia {razao_social}, o indicador de exclusividade de CRMs foi classificado como crítico na matriz de risco do Sistema Sentinela. O percentual de CRMs exclusivos observado foi de ", color="0F172A", size=10)
    _run(p2, f"{_format_decimal_pt(pct_medicos_exclusivos, 2)}%", color="334155", size=10, bold=True)
    if has_detalhe_exclusivos:
        _run(p2, f" dos {total_medicos} médicos observados. Esse percentual corresponde a ", color="0F172A", size=10)
    else:
        _run(p2, ". Esse percentual corresponde a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_reg_fmt} {multiplicador_reg_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o percentual mediano de CRMs exclusivos das farmácias de sua região. Ampliando-se o comparativo geográfico, o percentual equivale a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_uf_fmt} {multiplicador_uf_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o percentual mediano das farmácias de seu Estado e a ", color="0F172A", size=10)
    _run(p2, f"{multiplicador_br_fmt} {multiplicador_br_unidade}", color="334155", size=10, bold=True)
    _run(p2, " o das farmácias de todo o Brasil.", color="0F172A", size=10)

    if not has_detalhe_exclusivos:
        return

    title = doc.add_paragraph()
    _format_crm_table_title(title)
    _run(
        title,
        f"Quadro 09 - Médicos com CRMs registrados exclusivamente pela Farmácia {razao_social} (CNPJ {cnpj_fmt}) no Sistema Autorizador de Vendas, no período {periodo_intervalo}.",
        color="0F172A",
        size=9,
        bold=True,
    )

    headers = [
        "CRM/UF",
        "Nome",
        "Inscrição no CFM",
        "Núm. autorizações",
        "Valor total",
        "% sobre o valor total",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.8), Inches(2.25), Inches(0.9), Inches(1.0), Inches(1.25), Inches(0.9)]
    _set_table_fixed_widths(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _cell_bg(cell, "E2E8F0")
        _write_cell(cell, header, size=7.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in exclusividade_comp["top_exclusivos"]:
        cells = table.add_row().cells
        crm_row, uf_row = _crm_num_uf(row.get("id_medico"))
        crm_uf = f"{crm_row}/{uf_row}" if uf_row else crm_row
        row_autorizacoes = _as_int(row.get("nu_prescricoes"))
        row_valor = _as_float(row.get("vl_total_prescricoes"))
        pct_valor_total = (row_valor / total_financeiro_base * 100) if total_financeiro_base else 0.0
        values = [
            crm_uf,
            str(row.get("no_medico") or "Não localizado"),
            _format_date_br(row.get("dt_inscricao_crm")),
            str(row_autorizacoes),
            f"R$ {_format_decimal_pt(row_valor, 2)}",
            f"{_format_decimal_pt(pct_valor_total, 2)}%",
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if idx in (3, 4, 5) else WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else None
            _write_cell(cells[idx], value, size=7.0, align=align)

    fonte = doc.add_paragraph()
    _format_crm_table_footnote(fonte)
    _run(
        fonte,
        "Fonte: Consulta ao CFM (https://portal.cfm.org.br/busca-medicos) e Sistema Autorizador de Vendas (SAV).",
        color="475569",
        size=8,
        italic=True,
    )
