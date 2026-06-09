from datetime import date
from typing import Any, Optional
import unicodedata

import polars as pl
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from data_cache import (
    get_df_dados_ibge_demografia,
    get_df_dados_farmacia,
    get_df_perfil_estabelecimento,
    scan_analise_gtin_inconsistencia_clinica,
    scan_geografico_origem_uf,
)
from .falecidos import get_falecidos_data
from .matriz_risco_dinamica import (
    INDICATOR_MAPPING,
    _INDICATOR_FLAGS,
    build_dynamic_matriz_risco as _build_dynamic_matriz_risco,
)
from .nota_tecnica_docx_utils import (
    _add_bookmark,
    _add_internal_hyperlink,
    _cell_bg,
    _footnote_ref,
    _keep_small_table_together,
    _run,
    _set_table_fixed_widths,
)
from .nota_tecnica_charts import (
    _add_mapa_geografico_origem_uf,
    _add_figura_parkinson_comparacao,
    _add_figura_parkinson_faixas_etarias,
)
from .nota_tecnica_formatters import _format_date_month_year_long_pt, _format_decimal_pt, _title_case_pt

# ── Mapeamento da Seção 5 ──────────────────────────────────────────────────

_SECAO5_MAP = [
    ('falecidos',                    '5.5',  'Vendas de medicamentos para pessoas falecidas'),
    ('incompatibilidade_patologica', '5.6',  'Vendas de medicamentos com incompatibilidade patológica'),
    ('teto',                         '5.7',  'Vendas no “teto máximo” para clientes da Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('polimedicamento',              '5.8',  'Vendas de quatro ou mais itens de medicamentos por cupom realizadas pela Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('ticket_medio',                 '5.10', 'Valor do “ticket médio” dos medicamentos vendidos pela Farmácia {farmacia}, muito superior ao dos estabelecimentos de sua região'),
    ('receita_paciente',             '5.11', 'Faturamento médio mensal por cliente, obtido pela Farmácia {farmacia}, muito superior ao dos estabelecimentos de sua região'),
    ('per_capita',                   '5.12', 'Faturamento mensal per capita, obtido pela Farmácia {farmacia}, muito superior ao dos estabelecimentos de sua região'),
    ('alto_custo',                   '5.13', 'Vendas de medicamentos de alto custo realizadas pela Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('vendas_rapidas',               '5.14', 'Vendas de medicamentos em tempo inferior a 60 segundos'),
    ('recorrencia_sistemica',        '5.15', 'Vendas de medicamentos com precisão absoluta de 30 dias realizadas pela Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('dias_pico',                    '5.16', 'Vendas de medicamentos em dias de pico realizadas pela Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('dispersao_geografica',         '5.17', 'Vendas para pessoas residentes em outros Estados realizadas pela Farmácia {farmacia} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região'),
    ('hhi_crm',                      '5.19', 'Concentração atípica de registros de CRMs nas autorizações'),
    ('crms_irregulares',             '5.21', 'Vendas de medicamentos prescritos por médicos com irregularidade em seus CRMs'),
]
_SECAO5_ORDER = {key: idx for idx, (key, _, _) in enumerate(_SECAO5_MAP)}
_FORCAR_TODOS_CRITICOS_NOTA_TECNICA = False
_PARKINSON_PREVALENCIA_50_MAIS = 0.0086
_IBGE_ANO_CENSO_DEMOGRAFIA = 2022
_CLINICA_VALOR_MINIMO_DETALHAMENTO = 1000.0
_ZERO_BASELINE_CRITICAL_INDICATORS_NOTA = {
    "falecidos",
    "incompatibilidade_patologica",
    "volume_atipico",
}

_INDICADOR_QUADRO_META = {
    "percentual_nao_comprovacao": ("Percentual não comprovação", "pct"),
    "falecidos": ("Vendas para falecidos", "pct3"),
    "incompatibilidade_patologica": ("Incompatibilidade patológica", "pct"),
    "teto": ("Dispensação em teto máximo", "pct"),
    "polimedicamento": ("4+ itens por autorização", "pct"),
    "ticket_medio": ("Ticket médio", "val"),
    "receita_paciente": ("Receita por paciente", "val"),
    "per_capita": ("Venda per capita mensal", "val"),
    "alto_custo": ("Medicamentos de alto custo", "pct"),
    "vendas_rapidas": ("Vendas rápidas (<60s)", "pct"),
    "volume_atipico": ("Volume atípico", "dec"),
    "recorrencia_sistemica": ("Recorrência sistêmica", "pct"),
    "dias_pico": ("Concentração em dias de pico", "pct"),
    "dispersao_geografica": ("Dispersão interestadual", "pct"),
    "hhi_crm": ("Concentração de CRMs (HHI)", "dec"),
    "crms_irregulares": ("CRMs irregulares", "pct"),
}

_CLINICA_PATOLOGIA_META = {
    ("DOENCA DE PARKINSON", "IDADE_MENOR_50"): {
        "titulo": "Doença de Parkinson",
        "objeto": "Parkinson",
        "criterio": "beneficiários com menos de 50 anos",
        "descricao": "medicamentos destinados ao tratamento da doença de Parkinson, cuja ocorrência é incomum em pessoas com menos de 50 anos",
    },
    ("OSTEOPOROSE", "SEXO_MASCULINO"): {
        "titulo": "Osteoporose",
        "objeto": "osteoporose",
        "criterio": "beneficiários do sexo masculino",
        "descricao": "medicamentos destinados ao tratamento de osteoporose, cuja utilização é incomum em homens biológicos",
    },
    ("DIABETES", "IDADE_MENOR_20"): {
        "titulo": "Diabetes",
        "objeto": "diabetes",
        "criterio": "beneficiários com menos de 20 anos",
        "descricao": "medicamentos destinados ao tratamento de diabetes, cuja utilização é incomum em pessoas abaixo de 20 anos no recorte monitorado",
    },
    ("HIPERTENSAO", "IDADE_MENOR_20"): {
        "titulo": "Hipertensão",
        "objeto": "hipertensão",
        "criterio": "beneficiários com menos de 20 anos",
        "descricao": "medicamentos destinados ao tratamento de hipertensão, cuja utilização é incomum em pessoas abaixo de 20 anos no recorte monitorado",
    },
}

_CLINICA_COLUNAS_OBRIGATORIAS = {
    "id_cnpj",
    "patologia",
    "regra_clinica",
    "ano_base",
    "qtd_cpfs_distintos",
    "qtd_cpfs_incompativeis",
    "qtd_autorizacoes",
    "qtd_autorizacoes_incompativeis",
    "valor_total_pago",
    "valor_incompativel_pago",
    "percentual_cpfs_incompativeis",
    "rank_regional_qtd_cpfs_incompativeis",
    "percentil_regional_qtd_cpfs_incompativeis",
    "participacao_cpfs_incompativeis_regiao",
    "percentual_regional_cpfs_incompativeis",
    "razao_percentual_vs_regiao",
    "cpfs_incompativeis_esperados_regiao",
    "excesso_cpfs_incompativeis_vs_regiao",
}


def _require_columns(df: pl.DataFrame, required: set[str], fonte: str) -> None:
    missing = sorted(required - set(df.columns))
    if missing:
        raise RuntimeError(f"{fonte} sem colunas obrigatorias: {', '.join(missing)}")


def _normalize_ascii_upper(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = text.encode("ascii", "ignore").decode("ascii")
    return " ".join(text.replace("_", " ").split()).upper()


def _clinica_meta_key(patologia: Any, regra_clinica: Any) -> tuple[str, str]:
    patologia_key = _normalize_ascii_upper(patologia)
    regra_key = _normalize_ascii_upper(regra_clinica).replace(" ", "_")
    return patologia_key, regra_key


def _cnpj_digits(cnpj: str) -> str:
    digits = "".join(ch for ch in str(cnpj or "") if ch.isdigit())
    if len(digits) != 14:
        raise RuntimeError(f"CNPJ invalido para Nota Tecnica: {cnpj}")
    return digits


def _format_cnpj_pt(value: Any) -> str:
    digits = _cnpj_digits(str(value))
    return f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}"


def _format_int_pt(value: Any) -> str:
    try:
        return f"{int(round(float(value))):,}".replace(",", ".")
    except (TypeError, ValueError):
        raise RuntimeError(f"Valor inteiro obrigatorio invalido na Nota Tecnica: {value}")


def _format_brl_pt(value: Any) -> str:
    try:
        valor = float(value)
    except (TypeError, ValueError):
        raise RuntimeError(f"Valor financeiro obrigatorio invalido na Nota Tecnica: {value}")
    return f"R$ {_format_decimal_pt(valor, 2)}"


def _format_optional_decimal(value: Any, casas: int = 2) -> str:
    if value is None:
        return "não calculado"
    try:
        return _format_decimal_pt(float(value), casas)
    except (TypeError, ValueError):
        raise RuntimeError(f"Valor decimal opcional invalido na Nota Tecnica: {value}")


def _format_optional_percent(value: Any, casas: int = 2) -> str:
    formatted = _format_optional_decimal(value, casas)
    if formatted == "não calculado":
        return formatted
    return f"{formatted}%"


def _format_optional_ratio_percent(value: Any, casas: int = 2) -> str:
    if value is None:
        return "não calculado"
    try:
        return f"{_format_decimal_pt(float(value) * 100.0, casas)}%"
    except (TypeError, ValueError):
        raise RuntimeError(f"Valor proporcional opcional invalido na Nota Tecnica: {value}")


def _ratio(numerator: Any, denominator: Any) -> float | None:
    try:
        den = float(denominator)
        if den == 0:
            return None
        return float(numerator) / den
    except (TypeError, ValueError):
        raise RuntimeError(
            f"Razao obrigatoria invalida na Nota Tecnica: numerador={numerator}; denominador={denominator}"
        )


def _format_periodo_anos_item(item: dict[str, Any]) -> str:
    try:
        ano_inicio = int(item["ano_inicio"])
        ano_fim = int(item["ano_fim"])
    except (KeyError, TypeError, ValueError):
        raise RuntimeError("Periodo anual obrigatorio ausente para tabela clinica da Nota Tecnica.")
    if ano_fim < ano_inicio:
        raise RuntimeError(f"Periodo anual invalido para tabela clinica da Nota Tecnica: {ano_inicio}-{ano_fim}.")
    if ano_inicio == ano_fim:
        return f"em {ano_inicio}"
    return f"no período de {ano_inicio} a {ano_fim}"


def _sum_numeric(rows: list[dict[str, Any]], key: str) -> float:
    total = 0.0
    for row in rows:
        try:
            total += float(row.get(key) or 0)
        except (TypeError, ValueError):
            raise RuntimeError(f"Valor numerico obrigatorio invalido na Nota Tecnica: campo={key}; valor={row.get(key)}")
    return total


def _build_clinica_municipal_context(
    item: dict[str, Any],
    clinica_municipio: pl.DataFrame,
    perfil_estabelecimento: pl.DataFrame,
    id_cnpj_alvo: int,
) -> dict[str, Any]:
    key = _clinica_meta_key(item["patologia"], item["regra_clinica"])
    clinica_recorte = clinica_municipio.filter(
        pl.struct(["patologia", "regra_clinica"]).map_elements(
            lambda row: _clinica_meta_key(row["patologia"], row["regra_clinica"]) == key,
            return_dtype=pl.Boolean,
        )
    )
    if clinica_recorte.is_empty():
        raise RuntimeError(
            "Comparativo municipal clinico sem linhas para o recorte da Nota Tecnica: "
            f"patologia={item['patologia']}; regra_clinica={item['regra_clinica']}."
        )

    agregado_cnpj = (
        clinica_recorte
        .group_by("id_cnpj")
        .agg(
            pl.sum("qtd_cpfs_distintos").alias("qtd_cpfs_distintos"),
            pl.sum("qtd_cpfs_incompativeis").alias("qtd_cpfs_incompativeis"),
            pl.sum("qtd_autorizacoes_incompativeis").alias("qtd_autorizacoes_incompativeis"),
            pl.sum("valor_incompativel_pago").alias("valor_incompativel_pago"),
        )
        .with_columns(pl.col("id_cnpj").cast(pl.Int64))
    )
    alvo_rows = agregado_cnpj.filter(pl.col("id_cnpj") == id_cnpj_alvo)
    if alvo_rows.is_empty():
        raise RuntimeError(f"Comparativo municipal clinico sem CNPJ alvo id_cnpj={id_cnpj_alvo}.")

    all_rows = agregado_cnpj.to_dicts()
    alvo = alvo_rows.row(0, named=True)
    total = {
        "qtd_farmacias": len(all_rows),
        "qtd_cpfs_distintos": _sum_numeric(all_rows, "qtd_cpfs_distintos"),
        "qtd_cpfs_incompativeis": _sum_numeric(all_rows, "qtd_cpfs_incompativeis"),
        "qtd_autorizacoes_incompativeis": _sum_numeric(all_rows, "qtd_autorizacoes_incompativeis"),
        "valor_incompativel_pago": _sum_numeric(all_rows, "valor_incompativel_pago"),
    }
    demais = {
        "qtd_farmacias": max(int(total["qtd_farmacias"]) - 1, 0),
        "qtd_cpfs_distintos": float(total["qtd_cpfs_distintos"]) - float(alvo["qtd_cpfs_distintos"]),
        "qtd_cpfs_incompativeis": float(total["qtd_cpfs_incompativeis"]) - float(alvo["qtd_cpfs_incompativeis"]),
        "qtd_autorizacoes_incompativeis": (
            float(total["qtd_autorizacoes_incompativeis"]) - float(alvo["qtd_autorizacoes_incompativeis"])
        ),
        "valor_incompativel_pago": float(total["valor_incompativel_pago"]) - float(alvo["valor_incompativel_pago"]),
    }
    resumo = [
        {"grupo": "Farmácia analisada", "qtd_farmacias": 1, **alvo},
        {"grupo": "Demais farmácias do município", **demais},
        {"grupo": "Total do município", **total},
    ]

    perfil_required = {"id_cnpj", "cnpj", "razao_social"}
    _require_columns(perfil_estabelecimento, perfil_required, "Cache de perfil dos estabelecimentos")
    perfil = (
        perfil_estabelecimento
        .select(["id_cnpj", "cnpj", "razao_social"])
        .with_columns([
            pl.col("id_cnpj").cast(pl.Int64),
            pl.col("cnpj").cast(pl.Utf8).str.replace_all(r"\D", "").str.zfill(14),
            pl.col("razao_social").cast(pl.Utf8),
        ])
        .unique(subset=["id_cnpj"], keep="first")
    )

    top = (
        agregado_cnpj
        .sort(
            ["valor_incompativel_pago", "qtd_cpfs_incompativeis", "qtd_cpfs_distintos"],
            descending=[True, True, True],
        )
        .head(10)
        .join(perfil, on="id_cnpj", how="left")
        .with_row_index("posicao", offset=1)
    )
    if top.filter(pl.col("cnpj").is_null() | pl.col("razao_social").is_null()).height > 0:
        ids = top.filter(pl.col("cnpj").is_null() | pl.col("razao_social").is_null()).get_column("id_cnpj").to_list()
        raise RuntimeError(f"Perfil de estabelecimento ausente para Top 10 municipal clinico: id_cnpj={ids}.")

    total_valor_incompativel = float(total["valor_incompativel_pago"])
    top_rows = []
    for row in top.to_dicts():
        top_rows.append({
            **row,
            "is_alvo": int(row["id_cnpj"]) == id_cnpj_alvo,
            "percentual_cpfs_incompativeis": _ratio(row["qtd_cpfs_incompativeis"], row["qtd_cpfs_distintos"]),
            "participacao_municipal": _ratio(row["valor_incompativel_pago"], total_valor_incompativel),
        })

    return {
        "resumo": resumo,
        "top20": top_rows,
        "qtd_farmacias_municipio": int(total["qtd_farmacias"]),
        "qtd_cpfs_incompativeis_municipio": float(total["qtd_cpfs_incompativeis"]),
        "valor_incompativel_municipio": total_valor_incompativel,
    }


def _build_parkinson_demografia_context(
    farmacia_row: dict[str, Any],
    evolucao_anual: list[dict[str, Any]],
) -> dict[str, Any]:
    if not evolucao_anual:
        raise RuntimeError("Evolucao anual de Parkinson obrigatoria ausente para comparacao demografica.")

    id_ibge7 = str(farmacia_row.get("id_ibge7") or "").strip()
    municipio = str(farmacia_row.get("municipio") or "").strip()
    uf = str(farmacia_row.get("uf") or "").strip().upper()
    if not id_ibge7 or not municipio or not uf:
        raise RuntimeError("Cache de farmacias sem municipio/UF/id_ibge7 obrigatorios para comparacao demografica de Parkinson.")

    observacao = max(
        evolucao_anual,
        key=lambda row: (
            float(row["qtd_cpfs_distintos"]),
            int(row["ano_base"]),
        ),
    )
    qtd_cpfs_observado = int(observacao["qtd_cpfs_distintos"])
    ano_observado = int(observacao["ano_base"])
    if qtd_cpfs_observado <= 0:
        raise RuntimeError("Quantidade de CPFs distintos de Parkinson invalida para comparacao demografica.")

    demografia = get_df_dados_ibge_demografia()
    _require_columns(
        demografia,
        {"id_ibge7", "ano_censo", "idade_min", "nu_populacao"},
        "Cache demografico IBGE",
    )
    demografia = demografia.with_columns([
        pl.col("id_ibge7").cast(pl.Utf8),
        pl.col("ano_censo").cast(pl.Int16, strict=False),
        pl.col("idade_min").cast(pl.Int16, strict=False),
        pl.col("nu_populacao").cast(pl.Int64, strict=False),
    ])
    demo_municipio = demografia.filter(
        (pl.col("id_ibge7") == id_ibge7)
        & (pl.col("ano_censo") == _IBGE_ANO_CENSO_DEMOGRAFIA)
    )
    if demo_municipio.is_empty():
        raise RuntimeError(
            f"Demografia IBGE {_IBGE_ANO_CENSO_DEMOGRAFIA} obrigatoria ausente para id_ibge7={id_ibge7}."
        )

    pop_total = demo_municipio.select(pl.sum("nu_populacao")).item()
    pop_50_mais = (
        demo_municipio
        .filter(pl.col("idade_min") >= 50)
        .select(pl.sum("nu_populacao"))
        .item()
    )
    if pop_total is None or int(pop_total) <= 0:
        raise RuntimeError(f"Demografia IBGE sem populacao total valida para id_ibge7={id_ibge7}.")
    if pop_50_mais is None or int(pop_50_mais) <= 0:
        raise RuntimeError(f"Demografia IBGE sem populacao 50+ valida para id_ibge7={id_ibge7}.")
    if demo_municipio.filter(pl.col("idade_min").is_null() | pl.col("nu_populacao").is_null()).height > 0:
        raise RuntimeError(f"Demografia IBGE com idade/populacao nula para id_ibge7={id_ibge7}.")

    pop_total_int = int(pop_total)
    pop_50_int = int(pop_50_mais)
    faixas_dict: dict[int, int] = {}
    for row in demo_municipio.select(["idade_min", "nu_populacao"]).to_dicts():
        idade_min = int(row["idade_min"])
        populacao = int(row["nu_populacao"])
        if idade_min < 0 or populacao < 0:
            raise RuntimeError(f"Demografia IBGE com idade/populacao invalida para id_ibge7={id_ibge7}.")
        faixa_inicio = 80 if idade_min >= 80 else (idade_min // 10) * 10
        faixas_dict[faixa_inicio] = faixas_dict.get(faixa_inicio, 0) + populacao
    faixas_etarias = [
        {
            "faixa": "80+",
            "faixa_inicio": faixa_inicio,
            "populacao": populacao,
            "percentual": populacao / pop_total_int,
            "destacar_50_mais": True,
        }
        if faixa_inicio >= 80
        else {
            "faixa": f"{faixa_inicio} a {faixa_inicio + 9}",
            "faixa_inicio": faixa_inicio,
            "populacao": populacao,
            "percentual": populacao / pop_total_int,
            "destacar_50_mais": faixa_inicio >= 50,
        }
        for faixa_inicio, populacao in sorted(faixas_dict.items())
    ]
    if sum(int(item["populacao"]) for item in faixas_etarias) != pop_total_int:
        raise RuntimeError(f"Faixas etarias IBGE nao reconciliam com populacao total para id_ibge7={id_ibge7}.")

    casos_esperados = pop_50_int * _PARKINSON_PREVALENCIA_50_MAIS
    if casos_esperados <= 0:
        raise RuntimeError(f"Casos esperados de Parkinson invalidos para id_ibge7={id_ibge7}.")

    razao_observado_esperado = qtd_cpfs_observado / casos_esperados

    return {
        "id_ibge7": id_ibge7,
        "municipio": _title_case_pt(municipio),
        "uf": uf,
        "ano_censo": _IBGE_ANO_CENSO_DEMOGRAFIA,
        "populacao_total": pop_total_int,
        "populacao_50_mais": pop_50_int,
        "percentual_50_mais": pop_50_int / pop_total_int,
        "faixas_etarias": faixas_etarias,
        "prevalencia_referencia": _PARKINSON_PREVALENCIA_50_MAIS,
        "casos_esperados": casos_esperados,
        "ano_observado": ano_observado,
        "qtd_cpfs_distintos_observado": qtd_cpfs_observado,
        "razao_observado_esperado": razao_observado_esperado,
        "percentual_superior": (razao_observado_esperado - 1) * 100,
    }


def _vez_ou_vezes(valor_formatado: str) -> str:
    try:
        valor = float(valor_formatado.replace(".", "").replace(",", "."))
    except ValueError:
        return "vezes"
    return "vez" if abs(valor) <= 1 else "vezes"


def _get_matriz_dinamica_nota(
    cnpj: str,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
) -> tuple[pl.DataFrame, dict[str, Any]]:
    cnpj_norm = _cnpj_digits(cnpj)
    perfil_df = get_df_perfil_estabelecimento()
    df = _build_dynamic_matriz_risco(
        data_inicio=data_inicio,
        data_fim=data_fim,
        perfil_df=perfil_df,
    )
    if df.is_empty():
        raise RuntimeError("Matriz dinamica de risco sem linhas para a Nota Tecnica.")

    rows = df.with_columns(
        pl.col("cnpj").cast(pl.Utf8).str.replace_all(r"\D", "").str.zfill(14).alias("_cnpj_norm")
    ).filter(pl.col("_cnpj_norm") == cnpj_norm)
    if rows.height == 0:
        raise RuntimeError(f"CNPJ {cnpj_norm} nao encontrado na matriz dinamica de risco.")
    if rows.height > 1:
        raise RuntimeError(f"CNPJ {cnpj_norm} possui mais de uma linha na matriz dinamica de risco.")
    return df, rows.row(0, named=True)


def _get_matriz_row_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
    descricao: str,
) -> dict[str, Any]:
    try:
        _df, row = _get_matriz_dinamica_nota(cnpj, data_inicio, data_fim)
        return row
    except Exception as exc:
        raise RuntimeError(
            f"Matriz de risco dinamica indisponivel para {descricao} {cnpj}: {exc}"
        ) from exc


def _get_criticos(
    cnpj: str,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
) -> set[str]:
    """Identifica quais indicadores estão em nível CRÍTICO para o CNPJ."""
    if _FORCAR_TODOS_CRITICOS_NOTA_TECNICA:
        return {key for key, _, _ in _SECAO5_MAP}

    try:
        _df, row = _get_matriz_dinamica_nota(cnpj, data_inicio, data_fim)
        return {
            key for key, (_, flag_c) in _INDICATOR_FLAGS.items()
            if row.get(flag_c.lower()) == 1
        }
    except Exception as exc:
        raise RuntimeError(f"Falha ao identificar indicadores criticos para {cnpj}: {exc}") from exc


def _get_criticos_ordenados_por_risco(
    cnpj: str,
    criticos: set[str] | None = None,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
) -> list[str]:
    """Ordena indicadores criticos pela severidade regional usada na Nota Tecnica."""
    criticos = set(criticos if criticos is not None else _get_criticos(cnpj, data_inicio, data_fim))
    if not criticos:
        return []

    if _FORCAR_TODOS_CRITICOS_NOTA_TECNICA:
        return [key for key, _, _ in _SECAO5_MAP if key in criticos]

    cnpj_norm = _cnpj_digits(cnpj)
    _df, matriz_row = _get_matriz_dinamica_nota(cnpj_norm, data_inicio, data_fim)

    required_columns = {"cnpj"}
    for key in criticos:
        if key not in INDICATOR_MAPPING:
            raise RuntimeError(f"Indicador critico sem mapeamento de colunas para ordenacao da Nota Tecnica: {key}")
        c_val, c_med_reg, _c_med_uf, _c_med_br, c_risco_reg, _c_risco_uf, _c_risco_br = INDICATOR_MAPPING[key]
        required_columns.update({c_val, c_med_reg, c_risco_reg})

    missing_columns = sorted(required_columns - set(matriz_row.keys()))
    if missing_columns:
        raise RuntimeError(
            "Matriz de risco sem colunas obrigatorias para ordenar indicadores criticos da Nota Tecnica: "
            + ", ".join(missing_columns)
        )
    risco_por_key: dict[str, float] = {}
    for key in criticos:
        c_val, c_med_reg, _c_med_uf, _c_med_br, c_risco_reg, _c_risco_uf, _c_risco_br = INDICATOR_MAPPING[key]
        valor = matriz_row.get(c_val)
        mediana_reg = matriz_row.get(c_med_reg)
        risco_reg = matriz_row.get(c_risco_reg)
        risco_por_key[key] = _risco_regional_sort_value(key, valor, mediana_reg, risco_reg)

    return sorted(
        criticos,
        key=lambda key: (
            0 if key == "percentual_nao_comprovacao" else 1,
            -risco_por_key[key],
            _SECAO5_ORDER.get(key, len(_SECAO5_ORDER)),
            key,
        ),
    )


def _format_indicador_quadro_value(value: Any, formato: str) -> str:
    if value is None:
        raise RuntimeError("Indicador critico sem valor calculado para o quadro da Nota Tecnica.")
    try:
        numeric_value = float(value)
    except (TypeError, ValueError):
        raise RuntimeError(f"Valor de indicador critico invalido para a Nota Tecnica: {value}")
    if formato == "val":
        return _format_brl_pt(numeric_value)
    if formato == "pct3":
        return f"{_format_decimal_pt(numeric_value, 3)}%"
    if formato == "pct":
        return f"{_format_decimal_pt(numeric_value, 2)}%"
    if formato == "dec":
        return _format_decimal_pt(numeric_value, 2)
    raise RuntimeError(f"Formato de indicador critico nao mapeado para a Nota Tecnica: {formato}")


def _format_indicador_quadro_optional(value: Any, formato: str) -> str:
    if value is None:
        return "—"
    return _format_indicador_quadro_value(value, formato)


def _optional_float(value: Any) -> float | None:
    if value is None:
        return None
    try:
        numeric_value = float(value)
    except (TypeError, ValueError):
        return None
    if numeric_value != numeric_value:
        return None
    return numeric_value


def _is_zero_baseline_critical(indicador_key: str, valor: Any, mediana: Any, risco: Any) -> bool:
    valor_float = _optional_float(valor)
    mediana_float = _optional_float(mediana)
    risco_float = _optional_float(risco)
    return (
        indicador_key in _ZERO_BASELINE_CRITICAL_INDICATORS_NOTA
        and valor_float is not None
        and valor_float > 0
        and risco_float is None
        and (mediana_float is None or mediana_float <= 0)
    )


def _risco_regional_sort_value(indicador_key: str, valor: Any, mediana: Any, risco: Any) -> float:
    risco_float = _optional_float(risco)
    if risco_float is not None:
        return risco_float
    if _is_zero_baseline_critical(indicador_key, valor, mediana, risco):
        return float("inf")
    raise RuntimeError(f"Indicador critico {indicador_key} sem risco regional na matriz.")


def _format_risco_regional(value: Any) -> str:
    risco_float = _optional_float(value)
    if risco_float is None:
        return "—"
    return f'{_format_decimal_pt(risco_float, 1)}x'


def _indicador_valor_farmacia_header(formato: str) -> str:
    if formato == "val":
        return "Valor farmácia"
    if formato in {"pct", "pct3"}:
        return "Percentual farmácia"
    if formato == "dec":
        return "Índice farmácia"
    raise RuntimeError(f"Formato de indicador critico nao mapeado para cabecalho da Nota Tecnica: {formato}")


def _build_indicadores_criticos_quadro(
    cnpj: str,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
) -> list[dict[str, Any]]:
    """Monta linhas do quadro-resumo de indicadores críticos a partir da matriz de risco."""
    cnpj_norm = _cnpj_digits(cnpj)
    _df, matriz_row = _get_matriz_dinamica_nota(cnpj_norm, data_inicio, data_fim)
    required_columns = {"cnpj"}
    for key in _INDICATOR_FLAGS:
        if key not in INDICATOR_MAPPING:
            raise RuntimeError(f"Indicador sem mapeamento de colunas para a Nota Tecnica: {key}")
        if key not in _INDICADOR_QUADRO_META:
            raise RuntimeError(f"Indicador sem metadados de quadro para a Nota Tecnica: {key}")
        c_val, c_med_reg, _c_med_uf, _c_med_br, c_risco_reg, _c_risco_uf, _c_risco_br = INDICATOR_MAPPING[key]
        _c_atencao, c_critico = _INDICATOR_FLAGS[key]
        required_columns.update({c_val, c_med_reg, c_risco_reg, c_critico})

    missing_columns = sorted(required_columns - set(matriz_row.keys()))
    if missing_columns:
        raise RuntimeError(
            "Matriz de risco sem colunas obrigatorias para o quadro de indicadores criticos da Nota Tecnica: "
            + ", ".join(missing_columns)
        )
    quadro_rows: list[dict[str, Any]] = []
    for key, (_c_atencao, c_critico) in _INDICATOR_FLAGS.items():
        if int(matriz_row.get(c_critico) or 0) != 1:
            continue
        c_val, c_med_reg, _c_med_uf, _c_med_br, c_risco_reg, _c_risco_uf, _c_risco_br = INDICATOR_MAPPING[key]
        label, formato = _INDICADOR_QUADRO_META[key]
        valor = matriz_row.get(c_val)
        mediana_reg = matriz_row.get(c_med_reg)
        risco_reg = matriz_row.get(c_risco_reg)
        if valor is None:
            raise RuntimeError(f"Indicador critico {key} sem valor na matriz.")
        if mediana_reg is None and not _is_zero_baseline_critical(key, valor, mediana_reg, risco_reg):
            raise RuntimeError(f"Indicador critico {key} sem mediana regional na matriz.")
        if risco_reg is None and not _is_zero_baseline_critical(key, valor, mediana_reg, risco_reg):
            raise RuntimeError(f"Indicador critico {key} sem valor, mediana regional ou risco regional na matriz.")
        quadro_rows.append({
            "key": key,
            "indicador": label,
            "bookmark": (
                "secao6_percentual_nao_comprovacao"
                if key == "percentual_nao_comprovacao"
                else f"secao7_{key}" if key in _SECAO5_ORDER else None
            ),
            "valor": _format_indicador_quadro_value(valor, formato),
            "mediana_regional": _format_indicador_quadro_optional(mediana_reg, formato),
            "risco_regional": _optional_float(risco_reg),
            "status": "CRÍTICO",
        })

    ordem = _get_criticos_ordenados_por_risco(cnpj, {row["key"] for row in quadro_rows}, data_inicio, data_fim)
    ordem_por_key = {key: idx for idx, key in enumerate(ordem)}
    return sorted(quadro_rows, key=lambda item: ordem_por_key[item["key"]])


def _build_indicador_regional_context(
    cnpj: str,
    indicador_key: str,
    data_inicio: Optional[date] = None,
    data_fim: Optional[date] = None,
) -> dict[str, Any]:
    """Monta o enquadramento da farmácia auditada na região de saúde do indicador."""
    if indicador_key not in INDICATOR_MAPPING:
        raise RuntimeError(f"Indicador sem mapeamento para enquadramento regional da Nota Tecnica: {indicador_key}")
    if indicador_key not in _INDICADOR_QUADRO_META:
        raise RuntimeError(f"Indicador sem metadados para enquadramento regional da Nota Tecnica: {indicador_key}")

    cnpj_norm = _cnpj_digits(cnpj)
    label, formato = _INDICADOR_QUADRO_META[indicador_key]
    c_val, c_med_reg, _c_med_uf, _c_med_br, c_risco_reg, _c_risco_uf, _c_risco_br = INDICATOR_MAPPING[indicador_key]

    df, _target_row = _get_matriz_dinamica_nota(cnpj_norm, data_inicio, data_fim)
    required_columns = {
        "cnpj",
        "id_regiao_saude",
        c_val,
        c_med_reg,
        c_risco_reg,
    }
    missing_columns = sorted(required_columns - set(df.columns))
    if missing_columns:
        raise RuntimeError(
            "Matriz de risco sem colunas obrigatorias para enquadramento regional da Nota Tecnica: "
            + ", ".join(missing_columns)
        )

    df = df.with_columns([
        pl.col("cnpj").cast(pl.Utf8).str.replace_all(r"\D", "").str.zfill(14).alias("_cnpj_norm"),
        pl.col("id_regiao_saude").cast(pl.Utf8).str.strip_chars().alias("_id_regiao_norm"),
    ])
    target_rows = df.filter(pl.col("_cnpj_norm") == cnpj_norm)
    if target_rows.height == 0:
        raise RuntimeError(f"CNPJ {cnpj_norm} nao encontrado na matriz de risco para enquadramento regional.")
    if target_rows.height > 1:
        raise RuntimeError(f"CNPJ {cnpj_norm} possui mais de uma linha na matriz de risco para enquadramento regional.")

    target = target_rows.row(0, named=True)
    id_regiao = str(target["_id_regiao_norm"] or "").strip()
    if not id_regiao:
        raise RuntimeError(f"CNPJ {cnpj_norm} sem id_regiao_saude na matriz de risco para enquadramento regional.")

    regional = (
        df
        .filter(
            (pl.col("_id_regiao_norm") == id_regiao)
            & pl.col(c_val).is_not_null()
        )
        .with_columns([
            pl.col(c_val).cast(pl.Float64).alias("_valor_rank"),
            (
                pl.when(
                    (pl.lit(indicador_key in _ZERO_BASELINE_CRITICAL_INDICATORS_NOTA))
                    & pl.col(c_val).is_not_null()
                    & (pl.col(c_val) > 0)
                    & (pl.col(c_med_reg).is_null() | (pl.col(c_med_reg) <= 0))
                    & pl.col(c_risco_reg).is_null()
                )
                .then(pl.lit(float("inf")))
                .otherwise(pl.col(c_risco_reg).cast(pl.Float64))
                .alias("_risco_rank")
            ),
        ])
        .sort(["_valor_rank", "_risco_rank", "_cnpj_norm"], descending=[True, True, False])
        .with_row_index("posicao_regional", offset=1)
    )
    if regional.is_empty():
        raise RuntimeError(f"Enquadramento regional sem linhas para indicador {indicador_key} e id_regiao_saude={id_regiao}.")

    target_rank_rows = regional.filter(pl.col("_cnpj_norm") == cnpj_norm)
    if target_rank_rows.is_empty():
        raise RuntimeError(f"CNPJ {cnpj_norm} sem valor/risco calculado no indicador {indicador_key} para enquadramento regional.")
    regional_row = target_rank_rows.row(0, named=True)

    valor = regional_row.get(c_val)
    mediana_reg = regional_row.get(c_med_reg)
    risco_reg = regional_row.get(c_risco_reg)
    if valor is None:
        raise RuntimeError(f"CNPJ {cnpj_norm} sem valor para indicador {indicador_key}.")
    if mediana_reg is None and not _is_zero_baseline_critical(indicador_key, valor, mediana_reg, risco_reg):
        raise RuntimeError(f"CNPJ {cnpj_norm} sem mediana regional para indicador {indicador_key}.")
    if risco_reg is None and not _is_zero_baseline_critical(indicador_key, valor, mediana_reg, risco_reg):
        raise RuntimeError(f"CNPJ {cnpj_norm} sem risco regional para indicador {indicador_key}.")

    total_regional = regional.height
    if total_regional <= 0:
        raise RuntimeError(f"Universo regional invalido para indicador {indicador_key} e id_regiao_saude={id_regiao}.")
    if total_regional == 1:
        percentil_regional = 100.0
    else:
        percentil_regional = (total_regional - int(regional_row["posicao_regional"])) / (total_regional - 1) * 100.0

    return {
        "indicador_key": indicador_key,
        "indicador_label": label,
        "formato": formato,
        "id_regiao_saude": id_regiao,
        "valor": _format_indicador_quadro_value(valor, formato),
        "mediana_regional": _format_indicador_quadro_optional(mediana_reg, formato),
        "risco_regional": _optional_float(risco_reg),
        "posicao_regional": int(regional_row["posicao_regional"]),
        "total_regional": total_regional,
        "percentil_regional": percentil_regional,
    }


def _add_indicador_regional_table(doc, context: dict[str, Any], tabela_num: int) -> None:
    if not context:
        raise RuntimeError("Contexto regional de indicador ausente para a Nota Tecnica.")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Enquadramento regional da farmácia auditada no indicador {context["indicador_label"]}',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    _set_table_fixed_widths(
        table,
        [Inches(1.65), Inches(1.05), Inches(1.15), Inches(0.85), Inches(0.75), Inches(0.95), Inches(0.60)],
    )

    headers = [
        "Indicador",
        _indicador_valor_farmacia_header(context["formato"]),
        "Mediana região",
        "Risco vs. região",
        "Posição região",
        "Núm. Farmácias Reg.",
        "Percentil reg.",
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=7, bold=True)
        _cell_bg(cell, 'E2E8F0')

    values = [
        context["indicador_label"],
        context["valor"],
        context["mediana_regional"],
        _format_risco_regional(context["risco_regional"]),
        _format_int_pt(context["posicao_regional"]),
        _format_int_pt(context["total_regional"]),
        f'{_format_decimal_pt(context["percentil_regional"], 1)}%',
    ]
    for col_idx, value in enumerate(values):
        cell = table.rows[1].cells[col_idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        _cell_bg(cell, 'F8FAFC')
        _run(para, value, color='0F172A', size=7)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        f'Fonte: Sentinela, a partir da matriz de risco consolidada. Posição e percentil calculados pelo valor do indicador entre as farmácias com dado válido na região de saúde ID {context["id_regiao_saude"]}.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_indicadores_criticos_quadro(doc, rows: list[dict[str, Any]], tabela_num: int) -> None:
    """Adiciona tabela sintética com os indicadores críticos da matriz de risco."""
    if not rows:
        return

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Indicadores classificados como críticos na matriz de risco do Sentinela',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = 'Table Grid'
    _set_table_fixed_widths(table, [Inches(2.65), Inches(1.18), Inches(1.18), Inches(0.94), Inches(1.05)])

    headers = ["Indicador", "Farmácia", "Mediana região", "Risco região", "Status"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='FFFFFF', size=7, bold=True)
        _cell_bg(cell, '1E293B')

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        values = [
            row["indicador"],
            row["valor"],
            row["mediana_regional"],
            _format_risco_regional(row["risco_regional"]),
            row["status"],
        ]
        fill = 'F8FAFC' if row_idx % 2 else 'FFFFFF'
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 4:
                _cell_bg(cell, 'FEE2E2')
                _run(para, value, color='991B1B', size=7, bold=True)
            elif col_idx == 3:
                _cell_bg(cell, 'FEF2F2')
                _run(para, value, color='991B1B', size=7, bold=True)
            else:
                _cell_bg(cell, fill)
                if col_idx == 0 and row.get("bookmark"):
                    _add_internal_hyperlink(
                        para,
                        value,
                        row["bookmark"],
                        color='1D4ED8',
                        size=7,
                        bold=True,
                        underline=False,
                    )
                else:
                    _run(para, value, color='0F172A', size=7, bold=col_idx == 1)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir da matriz de risco consolidada. O risco regional corresponde ao multiplicador entre o indicador da farmácia e a mediana dos estabelecimentos de sua região de saúde.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _build_falecidos_context(
    cnpj: str,
    uf_farmacia: str | None,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Agrega transacoes de falecidos para a Nota Tecnica quando houver dados."""
    try:
        falecidos = get_falecidos_data(cnpj, data_inicio, data_fim)
    except Exception as exc:
        print(f"[NOTA_TECNICA] Falecidos indisponivel para {cnpj}: {exc}")
        return None

    transacoes = getattr(falecidos, "transacoes", None) or []
    if not transacoes:
        return None

    summary = getattr(falecidos, "summary", None)
    total_autorizacoes = int(getattr(summary, "total_autorizacoes", len(transacoes)) or 0)
    cpfs_distintos = int(getattr(summary, "cpfs_distintos", 0) or 0)
    valor_total = float(getattr(summary, "valor_total", 0.0) or 0.0)

    cpfs = {
        getattr(t, "cpf", None)
        for t in transacoes
        if getattr(t, "cpf", None)
    }
    if cpfs:
        cpfs_distintos = len(cpfs)

    datas: list[date] = []
    for t in transacoes:
        data_autorizacao = getattr(t, "data_autorizacao", None)
        if isinstance(data_autorizacao, date):
            datas.append(data_autorizacao)
    inicio_ref = data_inicio or (min(datas) if datas else None)
    fim_ref = data_fim or (max(datas) if datas else None)
    if inicio_ref and fim_ref:
        inicio_txt = _format_date_month_year_long_pt(inicio_ref)
        fim_txt = _format_date_month_year_long_pt(fim_ref)
        periodo_desc = (
            f'no mês de {inicio_txt}'
            if inicio_txt == fim_txt
            else f'no período de {inicio_txt} a {fim_txt}'
        )
    elif inicio_ref:
        periodo_desc = f'a partir de {_format_date_month_year_long_pt(inicio_ref)}'
    elif fim_ref:
        periodo_desc = f'até {_format_date_month_year_long_pt(fim_ref)}'
    else:
        periodo_desc = 'no período analisado'

    return {
        "total_autorizacoes": total_autorizacoes,
        "cpfs_distintos": cpfs_distintos,
        "valor_total": valor_total,
        "periodo_desc": periodo_desc,
        "transacoes": transacoes,
    }


def _add_falecidos_criticidade_text(
    doc,
    num: str,
    razao_social: str,
    falecidos_comp: dict[str, Any],
    anexo_num: str = 'III',
    bookmark_name: str | None = None,
):
    """Adiciona texto analitico de vendas para pessoas falecidas."""
    total_autorizacoes = falecidos_comp["total_autorizacoes"]
    cpfs_distintos = falecidos_comp["cpfs_distintos"]
    valor_total = falecidos_comp["valor_total"]
    periodo_desc = falecidos_comp["periodo_desc"]

    heading = doc.add_heading(f'{num} Vendas de medicamentos para pessoas falecidas', level=2)
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)
    p1 = doc.add_paragraph()
    _run(p1, f'Em análise às informações lançadas pela Farmácia {razao_social} no Sistema Autorizador de Vendas (SAV) do PFPB, {periodo_desc}, foram identificados registros de vendas (distribuição) de medicamentos para pessoas na data de seus óbitos e/ou posteriormente a essa data, com confirmação nas bases de dados do SIRC', color='0F172A', size=10)
    _footnote_ref(
        doc,
        p1,
        14,
        'O Sistema Nacional de Informações de Registro Civil (Sirc) é uma base de governo federal que tem por finalidade captar, processar, arquivar e disponibilizar dados relativos a registros de nascimento, casamento, óbito e natimorto, produzidos pelos cartórios de registro civil das pessoas naturais. Sistema disponível em: https://sirc.gov.br/o-que-e/.',
    )
    _run(p1, ' e SISOBI', color='0F172A', size=10)
    _footnote_ref(
        doc,
        p1,
        15,
        'O Sistema de Controle de Óbitos (Sisobi) é ainda utilizado pelos cartórios para tratamento de certidões anteriores à implantação do Sistema Nacional de Informações de Registro Civil (Sirc). Sistema disponível em: https://www.dataprev.gov.br/sisobi/.',
    )
    _run(p1, '. Foram realizadas ', color='0F172A', size=10)
    _run(p1, f'{total_autorizacoes:,}'.replace(',', '.'), color='334155', size=10, bold=True)
    _run(p1, ' vendas em data posterior ao registro de morte de ', color='0F172A', size=10)
    _run(p1, f'{cpfs_distintos:,}'.replace(',', '.'), color='334155', size=10, bold=True)
    _run(p1, ' beneficiários. Estas vendas representaram um valor total de ', color='0F172A', size=10)
    _run(p1, f'R$ {_format_decimal_pt(valor_total, 2)}', color='334155', size=10, bold=True)
    _run(p1, ', no referido período.', color='0F172A', size=10)

    p2 = doc.add_paragraph()
    _run(
        p2,
        f'O ANEXO {anexo_num} desta Nota Técnica traz o detalhamento de todas as vendas realizadas pela Farmácia {razao_social}, '
        f'{periodo_desc}, na data do óbito da pessoa e/ou posteriormente a essa data.',
        color='0F172A',
        size=10,
    )


def _build_incompatibilidade_patologica_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca matriz e, quando houver detalhamento, ranking anual de patologias."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "incompatibilidade patologica")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    cnpj_limpo = _cnpj_digits(cnpj)
    farmacias = get_df_dados_farmacia()
    _require_columns(farmacias, {"id_cnpj", "cnpj", "id_ibge7", "municipio", "uf"}, "Cache de farmacias")
    farmacias = farmacias.with_columns(
        pl.col("cnpj").cast(pl.Utf8).str.replace_all(r"\D", "").alias("_cnpj_limpo")
    )
    farmacia_rows = farmacias.filter(pl.col("_cnpj_limpo") == cnpj_limpo)
    if farmacia_rows.is_empty():
        raise RuntimeError(f"CNPJ {cnpj_limpo} sem id_cnpj no cache de farmacias.")
    farmacia_row = farmacia_rows.row(0, named=True)
    id_cnpj = int(farmacia_row["id_cnpj"])
    id_ibge7 = str(farmacia_row["id_ibge7"]).strip()
    if not id_ibge7:
        raise RuntimeError(f"CNPJ {cnpj_limpo} sem id_ibge7 no cache de farmacias.")

    clinica = (
        scan_analise_gtin_inconsistencia_clinica()
        .filter(pl.col("id_cnpj") == id_cnpj)
        .collect()
    )
    _require_columns(clinica, _CLINICA_COLUNAS_OBRIGATORIAS, "Cache clinico por CNPJ")
    clinica_municipio = (
        scan_analise_gtin_inconsistencia_clinica()
        .filter(pl.col("id_ibge7").cast(pl.Utf8) == id_ibge7)
        .collect()
    )
    _require_columns(clinica_municipio, _CLINICA_COLUNAS_OBRIGATORIAS, "Cache clinico municipal")
    if data_inicio:
        clinica = clinica.filter(pl.col("ano_base") >= data_inicio.year)
        clinica_municipio = clinica_municipio.filter(pl.col("ano_base") >= data_inicio.year)
    if data_fim:
        clinica = clinica.filter(pl.col("ano_base") <= data_fim.year)
        clinica_municipio = clinica_municipio.filter(pl.col("ano_base") <= data_fim.year)
    perfil_estabelecimento = get_df_perfil_estabelecimento()
    ranking = []
    clinica_anual = clinica
    if not clinica_anual.is_empty():
        clinica_incompativel = clinica_anual.filter(pl.col("qtd_cpfs_incompativeis") > 0)
        if not clinica_incompativel.is_empty():
            ranking_df = (
                clinica_incompativel
                .group_by(["patologia", "regra_clinica"])
                .agg(
                    pl.sum("qtd_cpfs_distintos").alias("qtd_cpfs_distintos"),
                    pl.sum("qtd_cpfs_incompativeis").alias("qtd_cpfs_incompativeis"),
                    pl.sum("qtd_autorizacoes").alias("qtd_autorizacoes"),
                    pl.sum("qtd_autorizacoes_incompativeis").alias("qtd_autorizacoes_incompativeis"),
                    pl.sum("valor_total_pago").alias("valor_total_pago"),
                    pl.sum("valor_incompativel_pago").alias("valor_incompativel_pago"),
                    pl.sum("cpfs_incompativeis_esperados_regiao").alias("cpfs_incompativeis_esperados_regiao"),
                    pl.sum("excesso_cpfs_incompativeis_vs_regiao").alias("excesso_cpfs_incompativeis_vs_regiao"),
                    pl.mean("percentual_cpfs_incompativeis").alias("percentual_medio_cpfs_incompativeis"),
                    pl.mean("percentual_regional_cpfs_incompativeis").alias("percentual_medio_regional_cpfs_incompativeis"),
                    pl.mean("razao_percentual_vs_regiao").alias("razao_media_percentual_vs_regiao"),
                    pl.min("rank_regional_qtd_cpfs_incompativeis").alias("melhor_rank_regional_qtd_cpfs_incompativeis"),
                    pl.max("percentil_regional_qtd_cpfs_incompativeis").alias("maior_percentil_regional_qtd_cpfs_incompativeis"),
                    pl.max("participacao_cpfs_incompativeis_regiao").alias("maior_participacao_cpfs_incompativeis_regiao"),
                    pl.min("ano_base").alias("ano_inicio"),
                    pl.max("ano_base").alias("ano_fim"),
                    pl.len().alias("qtd_linhas_anuais"),
                )
                .filter(
                    (pl.col("valor_incompativel_pago") >= _CLINICA_VALOR_MINIMO_DETALHAMENTO)
                )
            )

            for item in ranking_df.iter_rows(named=True):
                key = _clinica_meta_key(item["patologia"], item["regra_clinica"])
                meta = _CLINICA_PATOLOGIA_META.get(key)
                if meta is None:
                    raise RuntimeError(
                        "Recorte clinico sem mapeamento textual para a Nota Tecnica: "
                        f"patologia={item['patologia']}; regra_clinica={item['regra_clinica']}."
                    )
                evolucao_anual = sorted(
                    [
                        row
                        for row in clinica_anual.iter_rows(named=True)
                        if _clinica_meta_key(row["patologia"], row["regra_clinica"]) == key
                    ],
                    key=lambda row: int(row["ano_base"]),
                )
                item_context = {
                    **item,
                    "titulo": meta["titulo"],
                    "objeto": meta["objeto"],
                    "criterio": meta["criterio"],
                    "descricao": meta["descricao"],
                    "evolucao_anual": evolucao_anual,
                    "municipal": _build_clinica_municipal_context(
                        item,
                        clinica_municipio,
                        perfil_estabelecimento,
                        id_cnpj,
                    ),
                }
                if key == ("DOENCA DE PARKINSON", "IDADE_MENOR_50"):
                    item_context["demografia_parkinson"] = _build_parkinson_demografia_context(
                        farmacia_row,
                        evolucao_anual,
                    )
                ranking.append(item_context)

    ranking.sort(
        key=lambda item: (
            float(item["excesso_cpfs_incompativeis_vs_regiao"])
            if item["excesso_cpfs_incompativeis_vs_regiao"] is not None
            else float("-inf"),
            float(item["qtd_cpfs_incompativeis"]),
            float(item["valor_incompativel_pago"]),
            float(item["percentual_medio_cpfs_incompativeis"])
            if item["percentual_medio_cpfs_incompativeis"] is not None
            else float("-inf"),
        ),
        reverse=True,
    )

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_clinico"),
        "mediana_regiao": as_float("med_clinico_reg"),
        "mediana_uf": as_float("med_clinico_uf"),
        "mediana_brasil": as_float("med_clinico_br"),
        "multiplicador_regiao": as_float("risco_clinico_reg"),
        "multiplicador_uf": as_float("risco_clinico_uf"),
        "multiplicador_brasil": as_float("risco_clinico_br"),
        "id_cnpj": id_cnpj,
        "ranking_patologias": ranking,
    }


def _add_clinica_evolucao_anual_table(doc, item: dict[str, Any], tabela_num: int):
    evolucao_anual = item.get("evolucao_anual") or []
    if not evolucao_anual:
        raise RuntimeError("Evolucao anual clinica obrigatoria ausente para a Nota Tecnica.")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Dispensa\u00e7\u00f5es anuais de medicamentos para {item["objeto"]} no per\u00edodo selecionado',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=len(evolucao_anual) + 1, cols=8)
    table.style = 'Table Grid'
    _set_table_fixed_widths(
        table,
        [Inches(0.55), Inches(0.75), Inches(0.82), Inches(0.70), Inches(0.82), Inches(0.92), Inches(0.75), Inches(1.69)],
    )

    headers = [
        'Ano',
        'CPFs total',
        'CPFs incompat.',
        '% CPFs',
        'Aut. total',
        'Aut. incompat.',
        '% aut.',
        'Valor incompatível',
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=7, bold=True)
        _cell_bg(cell, 'E2E8F0')

    for row_idx, row in enumerate(evolucao_anual, start=1):
        pct_cpfs = row.get("percentual_cpfs_incompativeis")
        if pct_cpfs is None:
            pct_cpfs = _ratio(row.get("qtd_cpfs_incompativeis"), row.get("qtd_cpfs_distintos"))
        pct_aut = _ratio(row.get("qtd_autorizacoes_incompativeis"), row.get("qtd_autorizacoes"))
        values = [
            str(int(row["ano_base"])),
            _format_int_pt(row["qtd_cpfs_distintos"]),
            _format_int_pt(row["qtd_cpfs_incompativeis"]),
            _format_optional_ratio_percent(pct_cpfs, 2),
            _format_int_pt(row["qtd_autorizacoes"]),
            _format_int_pt(row["qtd_autorizacoes_incompativeis"]),
            _format_optional_ratio_percent(pct_aut, 2),
            _format_brl_pt(row["valor_incompativel_pago"]),
        ]
        for col_idx, value in enumerate(values):
            para = table.rows[row_idx].cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(para, value, color='0F172A', size=7, bold=col_idx in (2, 5, 7))

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir dos registros do SAV/PFPB e das regras clínicas do indicador de incompatibilidade patológica.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_clinica_municipio_resumo_table(doc, item: dict[str, Any], tabela_num: int) -> None:
    municipal = item.get("municipal")
    if not municipal:
        raise RuntimeError("Comparativo municipal clinico obrigatorio ausente para a Nota Tecnica.")
    resumo = municipal.get("resumo") or []
    if len(resumo) != 3:
        raise RuntimeError("Resumo municipal clinico deve conter farmacia analisada, demais farmacias e total municipal.")
    total_valor_incompativel = resumo[-1]["valor_incompativel_pago"]

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Comparativo municipal de valores incompatíveis no recorte de {item["titulo"]}, {_format_periodo_anos_item(item)}',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=len(resumo) + 1, cols=4)
    table.style = 'Table Grid'
    _set_table_fixed_widths(
        table,
        [Inches(2.90), Inches(0.75), Inches(1.80), Inches(1.55)],
    )

    headers = ["Grupo", "Farm.", "Valor incompatível", "Part. valor mun."]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=6, bold=True)
        _cell_bg(cell, 'E2E8F0')

    for row_idx, row in enumerate(resumo, start=1):
        values = [
            row["grupo"],
            _format_int_pt(row["qtd_farmacias"]),
            _format_brl_pt(row["valor_incompativel_pago"]),
            _format_optional_ratio_percent(_ratio(row["valor_incompativel_pago"], total_valor_incompativel), 2),
        ]
        for col_idx, value in enumerate(values):
            para = table.rows[row_idx].cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _run(para, value, color='0F172A', size=7, bold=row_idx == 1 or col_idx in (2, 3))

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir dos registros do SAV/PFPB e das regras clínicas do indicador de incompatibilidade patológica. Valores consolidados para o período selecionado.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_clinica_municipio_top20_table(doc, item: dict[str, Any], tabela_num: int) -> None:
    municipal = item.get("municipal")
    if not municipal:
        raise RuntimeError("Ranking municipal clinico obrigatorio ausente para a Nota Tecnica.")
    top20 = municipal.get("top20") or []
    if not top20:
        raise RuntimeError("Top 10 municipal clinico sem farmacias para a Nota Tecnica.")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Farmácias do município com maior valor incompatível no recorte de {item["titulo"]}, {_format_periodo_anos_item(item)}',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=len(top20) + 1, cols=5)
    table.style = 'Table Grid'
    _set_table_fixed_widths(
        table,
        [Inches(0.45), Inches(1.20), Inches(2.65), Inches(1.45), Inches(1.25)],
    )

    headers = ["Pos.", "CNPJ", "Razão social", "Valor incompatível", "Part. valor mun."]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=6, bold=True)
        _cell_bg(cell, 'E2E8F0')

    for row_idx, row in enumerate(top20, start=1):
        values = [
            _format_int_pt(row["posicao"]),
            _format_cnpj_pt(row["cnpj"]),
            str(row["razao_social"]),
            _format_brl_pt(row["valor_incompativel_pago"]),
            _format_optional_ratio_percent(row["participacao_municipal"], 2),
        ]
        for col_idx, value in enumerate(values):
            para = table.rows[row_idx].cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 2 else WD_ALIGN_PARAGRAPH.CENTER
            _run(para, value, color='0F172A', size=6.5, bold=bool(row["is_alvo"]) or col_idx in (3, 4))

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir dos registros do SAV/PFPB e das regras clínicas do indicador de incompatibilidade patológica. Ranking ordenado pelo valor incompatível consolidado no período selecionado.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _is_parkinson_menor_50_item(item: dict[str, Any]) -> bool:
    return _clinica_meta_key(item["patologia"], item["regra_clinica"]) == ("DOENCA DE PARKINSON", "IDADE_MENOR_50")


def _count_incompatibilidade_patologica_tables(clinico_comp: dict[str, Any]) -> int:
    ranking_patologias = clinico_comp.get("ranking_patologias") or []
    total = len(ranking_patologias) * 3
    total += sum(1 for item in ranking_patologias if _is_parkinson_menor_50_item(item))
    return total


def _add_parkinson_demografia_table(doc, demografia: dict[str, Any], tabela_num: int) -> None:
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Memória de cálculo da comparação epidemiológica de doença de Parkinson',
        color='334155',
        size=8,
        bold=True,
    )

    rows = [
        (
            "População municipal total",
            _format_int_pt(demografia["populacao_total"]),
            'IBGE/Censo',
        ),
        (
            "População municipal com 50 anos ou mais",
            _format_int_pt(demografia["populacao_50_mais"]),
            "Denominador usado na estimativa epidemiológica",
        ),
        (
            "Prevalência de referência para Parkinson em 50+",
            _format_optional_ratio_percent(demografia["prevalencia_referencia"], 2),
            "Parâmetro aplicado à população municipal com 50 anos ou mais",
        ),
        (
            "Casos esperados no município",
            _format_optional_decimal(demografia["casos_esperados"], 0),
            "População 50+ multiplicada pela prevalência de referência",
        ),
        (
            f'CPFs observados na farmácia em {demografia["ano_observado"]}',
            _format_int_pt(demografia["qtd_cpfs_distintos_observado"]),
            f'CPFs com dispensação de medicamentos para Parkinson em {demografia["ano_observado"]}',
        ),
        (
            "Razão observado/esperado",
            f'{_format_optional_decimal(demografia["razao_observado_esperado"], 2)}x',
            "CPFs observados divididos pelos casos esperados no município",
        ),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = 'Table Grid'
    _set_table_fixed_widths(table, [Inches(2.45), Inches(1.25), Inches(3.35)])

    headers = ["Métrica", "Valor", "Observação"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=7, bold=True)
        _cell_bg(cell, 'E2E8F0')

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            para = table.rows[row_idx].cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            _run(para, value, color='0F172A', size=7, bold=col_idx == 1)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, IBGE/Censo e prevalência nacional ajustada divulgada pelo Hospital de Clínicas de Porto Alegre com base na coorte ELSI-Brasil.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_parkinson_demografia_text(doc, item: dict[str, Any], tabela_num: int) -> None:
    demografia = item.get("demografia_parkinson")
    if not demografia:
        raise RuntimeError("Comparacao demografica de Parkinson obrigatoria ausente para a Nota Tecnica.")

    p_demo = doc.add_paragraph()
    _run(
        p_demo,
        'Segundo os dados demogr\u00e1ficos do IBGE/Censo, a popula\u00e7\u00e3o de '
        f'{demografia["municipio"]}/{demografia["uf"]} é de ',
        color='0F172A',
        size=10,
    )
    _run(p_demo, _format_int_pt(demografia["populacao_total"]), color='334155', size=10, bold=True)
    _run(p_demo, ' habitantes, dos quais ', color='0F172A', size=10)
    _run(p_demo, _format_int_pt(demografia["populacao_50_mais"]), color='334155', size=10, bold=True)
    _run(p_demo, ' possuem idade igual ou superior a 50 anos, correspondendo a ', color='0F172A', size=10)
    _run(p_demo, _format_optional_ratio_percent(demografia["percentual_50_mais"], 2), color='334155', size=10, bold=True)
    _run(
        p_demo,
        ' da popula\u00e7\u00e3o municipal. Estudo divulgado pelo Hospital de Cl\u00ednicas de Porto Alegre, com base na coorte ELSI-Brasil, indica preval\u00eancia nacional ajustada de aproximadamente ',
        color='0F172A',
        size=10,
    )
    _run(p_demo, _format_optional_ratio_percent(demografia["prevalencia_referencia"], 2), color='334155', size=10, bold=True)
    _run(
        p_demo,
        ' de doen\u00e7a de Parkinson entre pessoas com 50 anos ou mais. Aplicando-se essa preval\u00eancia \u00e0 popula\u00e7\u00e3o municipal nessa faixa et\u00e1ria, são esperados aproximadamente ',
        color='0F172A',
        size=10,
    )
    _run(p_demo, _format_optional_decimal(demografia["casos_esperados"], 0), color='334155', size=10, bold=True)
    _run(p_demo, ' casos de doen\u00e7a de Parkinson no munic\u00edpio. ', color='0F172A', size=10)
    _run(
        p_demo,
        f'No ano de {demografia["ano_observado"]}, a farm\u00e1cia analisada registra ',
        color='0F172A',
        size=10,
    )
    _run(p_demo, _format_int_pt(demografia["qtd_cpfs_distintos_observado"]), color='334155', size=10, bold=True)
    _run(
        p_demo,
        ' CPFs distintos com dispensa\u00e7\u00e3o de medicamentos para Parkinson, volume equivalente a ',
        color='0F172A',
        size=10,
    )
    _run(p_demo, _format_optional_decimal(demografia["razao_observado_esperado"], 2), color='334155', size=10, bold=True)
    _run(p_demo, ' vezes a estimativa epidemiol\u00f3gica de casos esperados para todo o munic\u00edpio, ou ', color='0F172A', size=10)
    _run(p_demo, _format_optional_percent(abs(demografia["percentual_superior"]), 2), color='334155', size=10, bold=True)
    if demografia["percentual_superior"] >= 0:
        _run(p_demo, ' acima dessa estimativa, ', color='0F172A', size=10)
    else:
        _run(p_demo, ' abaixo dessa estimativa, ', color='0F172A', size=10)
    _run(
        p_demo,
        'refor\u00e7ando a atipicidade do padr\u00e3o observado.',
        color='0F172A',
        size=10,
    )
    p_memoria = doc.add_paragraph()
    _run(
        p_memoria,
        'Para explicitar a memória de cálculo da comparação demográfica, apresenta-se a seguir a população municipal utilizada, a prevalência epidemiológica de referência, a estimativa de casos esperados no município e o volume de CPFs observados na farmácia.',
        color='0F172A',
        size=10,
    )
    _add_parkinson_demografia_table(doc, demografia, tabela_num)
    p_figuras = doc.add_paragraph()
    _run(
        p_figuras,
        'As figuras seguintes apresentam, de forma visual, os dois componentes dessa comparação: primeiro, a composição etária da população municipal utilizada como base da estimativa; em seguida, o confronto entre os casos esperados no município e os CPFs observados na farmácia.',
        color='0F172A',
        size=10,
    )
    _add_figura_parkinson_faixas_etarias(doc, demografia, figure_number=4)
    _add_figura_parkinson_comparacao(doc, demografia, figure_number=5)


def _add_incompatibilidade_patologica_text(
    doc,
    num: str,
    razao_social: str,
    clinico_comp: dict[str, Any],
    tabela_inicio_num: int,
    bookmark_name: str | None = None,
):
    """Adiciona texto analitico de incompatibilidade patologica e ranking interno."""
    periodo_desc = clinico_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(clinico_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(clinico_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(clinico_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(clinico_comp["multiplicador_brasil"], 2)
    ranking_patologias = clinico_comp.get("ranking_patologias") or []
    heading = doc.add_heading(f'{num} Vendas de medicamentos com incompatibilidade patológica', level=2)
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O comportamento esperado, no âmbito do PFPB, é que alguns medicamentos destinados a doenças específicas sejam distribuídos guardando correlação com o perfil demográfico do beneficiário, como idade ou sexo. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Neste rol, o Sentinela realiza levantamento detalhado de medicamentos para doença de Parkinson (incomum em pessoas com menos de 50 anos), osteoporose (incomum em homens biológicos), diabetes (incomum em pessoas abaixo de 20 anos) e hipertensão (incomum em pessoas abaixo de 20 anos).',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se, {periodo_desc}, percentual atípico de vendas desses medicamentos, correspondente a ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas monitoradas pelo indicador. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' a mediana dos percentuais de vendas com essa mesma criticidade realizadas pelas farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias localizadas em seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)

    if not ranking_patologias:
        return

    p3 = doc.add_paragraph()
    _run(
        p3,
        'Para identificar a origem clínica do alerta, o detalhamento anual do indicador foi ordenado internamente por excesso de CPFs incompatíveis em relação ao esperado para a região, quantidade de CPFs incompatíveis, valor incompatível pago e percentual de CPFs incompatíveis. ',
        color='0F172A',
        size=10,
    )
    _run(
        p3,
        'Assim, são descritas abaixo apenas as patologias que apresentaram incompatibilidades efetivas para o CNPJ analisado.',
        color='0F172A',
        size=10,
    )

    tabela_atual_num = tabela_inicio_num
    for idx, item in enumerate(ranking_patologias, start=1):
        doc.add_heading(f'{num}.{idx} {item["titulo"]}: {item["criterio"]}', level=3)
        pct_agregado_fmt = _format_optional_ratio_percent(
            _ratio(item.get("qtd_cpfs_incompativeis"), item.get("qtd_cpfs_distintos")),
            2,
        )
        percentil_fmt = _format_optional_ratio_percent(item.get("maior_percentil_regional_qtd_cpfs_incompativeis"), 2)
        participacao_fmt = _format_optional_ratio_percent(item.get("maior_participacao_cpfs_incompativeis_regiao"), 2)
        rank = item.get("melhor_rank_regional_qtd_cpfs_incompativeis")
        rank_fmt = _format_int_pt(rank) if rank is not None else "não calculado"

        p_item = doc.add_paragraph()
        is_parkinson_menor_50 = _is_parkinson_menor_50_item(item)
        if is_parkinson_menor_50:
            _run(
                p_item,
                'No período analisado, as apurações anuais registram dispensações de medicamentos para doença de Parkinson a ',
                color='0F172A',
                size=10,
            )
            _run(p_item, _format_int_pt(item["qtd_cpfs_distintos"]), color='334155', size=10, bold=True)
            _run(
                p_item,
                ' CPFs distintos. Embora a doença seja predominantemente associada a faixas etárias mais elevadas, ',
                color='0F172A',
                size=10,
            )
            _run(p_item, _format_int_pt(item["qtd_cpfs_incompativeis"]), color='334155', size=10, bold=True)
            _run(p_item, ' desses beneficiários, equivalentes a ', color='0F172A', size=10)
            _run(p_item, pct_agregado_fmt, color='334155', size=10, bold=True)
            _run(
                p_item,
                ', têm menos de 50 anos na data da dispensação. As operações associadas a esse grupo somam ',
                color='0F172A',
                size=10,
            )
            _run(p_item, _format_int_pt(item["qtd_autorizacoes_incompativeis"]), color='334155', size=10, bold=True)
            _run(p_item, ' autorizações e ', color='0F172A', size=10)
            _run(p_item, _format_brl_pt(item["valor_incompativel_pago"]), color='334155', size=10, bold=True)
            _run(p_item, ' em valores pagos pelo programa.', color='0F172A', size=10)
        else:
            _run(
                p_item,
                f'No período analisado, as apurações anuais registraram dispensações de medicamentos para {item["objeto"]} a ',
                color='0F172A',
                size=10,
            )
            _run(p_item, _format_int_pt(item["qtd_cpfs_distintos"]), color='334155', size=10, bold=True)
            _run(p_item, ' CPFs distintos. Desse universo, ', color='0F172A', size=10)
            _run(p_item, _format_int_pt(item["qtd_cpfs_incompativeis"]), color='334155', size=10, bold=True)
            _run(p_item, ' CPFs, equivalentes a ', color='0F172A', size=10)
            _run(p_item, pct_agregado_fmt, color='334155', size=10, bold=True)
            _run(
                p_item,
                f', enquadraram-se no critério de incompatibilidade clínica ({item["criterio"]}). As operações associadas a esse grupo somaram ',
                color='0F172A',
                size=10,
            )
            _run(p_item, _format_int_pt(item["qtd_autorizacoes_incompativeis"]), color='334155', size=10, bold=True)
            _run(p_item, ' autorizações e ', color='0F172A', size=10)
            _run(p_item, _format_brl_pt(item["valor_incompativel_pago"]), color='334155', size=10, bold=True)
            _run(p_item, ' em valores pagos pelo programa.', color='0F172A', size=10)

        _add_clinica_evolucao_anual_table(doc, item, tabela_atual_num)
        tabela_atual_num += 1
        if is_parkinson_menor_50:
            _add_parkinson_demografia_text(doc, item, tabela_atual_num)
            tabela_atual_num += 1

        p_municipio = doc.add_paragraph()
        _run(
            p_municipio,
            'A comparação territorial mais próxima permite avaliar se o achado se distribui entre os estabelecimentos do município ou se está concentrado na farmácia analisada. ',
            color='0F172A',
            size=10,
        )
        _run(
            p_municipio,
            'As tabelas seguintes resumem a materialidade financeira observada na farmácia, nos demais estabelecimentos municipais e no conjunto do município, além de listar as farmácias com maior valor incompatível para o mesmo recorte clínico.',
            color='0F172A',
            size=10,
        )
        _add_clinica_municipio_resumo_table(doc, item, tabela_atual_num)
        tabela_atual_num += 1
        _add_clinica_municipio_top20_table(doc, item, tabela_atual_num)
        tabela_atual_num += 1

        p_rank = doc.add_paragraph()
        _run(p_rank, 'No contexto regional, o estabelecimento ocupa a posição ', color='0F172A', size=10)
        _run(p_rank, rank_fmt, color='334155', size=10, bold=True)
        _run(p_rank, ' por quantidade de CPFs incompatíveis e situa-se no percentil regional ', color='0F172A', size=10)
        _run(p_rank, percentil_fmt, color='334155', size=10, bold=True)
        _run(p_rank, '. Além disso, concentra ', color='0F172A', size=10)
        _run(p_rank, participacao_fmt, color='334155', size=10, bold=True)
        _run(
            p_rank,
            ' dos CPFs incompatíveis apurados na região para esse recorte clínico, indicando que o achado regional está fortemente associado ao estabelecimento analisado.',
            color='0F172A',
            size=10,
        )


def _build_teto_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de vendas no teto maximo."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "teto maximo")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_teto"),
        "mediana_regiao": as_float("med_teto_reg"),
        "mediana_uf": as_float("med_teto_uf"),
        "mediana_brasil": as_float("med_teto_br"),
        "multiplicador_regiao": as_float("risco_teto_reg"),
        "multiplicador_uf": as_float("risco_teto_uf"),
        "multiplicador_brasil": as_float("risco_teto_br"),
    }


def _add_teto_text(doc, num: str, razao_social: str, teto_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de vendas no teto maximo usando a matriz atual."""
    periodo_desc = teto_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(teto_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(teto_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(teto_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(teto_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Vendas no “teto máximo” para clientes da Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O Programa Farmácia Popular do Brasil estabelece limites máximos de quantitativo de retirada mensal pelo cidadão de medicamentos, de acordo com cada um de seus princípios ativos. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Retiradas de medicamentos por um CPF no limite máximo mensal são consideradas, para fins de monitoramento, uma “venda no teto”. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'A expectativa da análise é que o percentual levantado para vendas de medicamentos “no teto” pelo estabelecimento acompanhe o padrão das demais farmácias localizadas na mesma região. Percentual muito acima da mediana da região sugere a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB foram realizadas no “teto máximo”. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com essa configuração das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias localizadas em seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_polimedicamento_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de quatro ou mais itens por cupom."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "polimedicamento")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_polimedicamento"),
        "mediana_regiao": as_float("med_polimedicamento_reg"),
        "mediana_uf": as_float("med_polimedicamento_uf"),
        "mediana_brasil": as_float("med_polimedicamento_br"),
        "multiplicador_regiao": as_float("risco_polimedicamento_reg"),
        "multiplicador_uf": as_float("risco_polimedicamento_uf"),
        "multiplicador_brasil": as_float("risco_polimedicamento_br"),
    }


def _add_polimedicamento_text(doc, num: str, razao_social: str, polimedicamento_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de cupons com quatro ou mais medicamentos usando a matriz atual."""
    periodo_desc = polimedicamento_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(polimedicamento_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(polimedicamento_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(polimedicamento_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(polimedicamento_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Vendas de quatro ou mais itens de medicamentos por cupom realizadas pela Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O comportamento esperado, no âmbito do PFPB, é que a grande maioria das vendas contenha apenas um ou dois itens de medicamentos. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'A ocorrência de cupons de vendas com quatro ou mais medicamentos foge do padrão epidemiológico esperado. Nesse sentido, cupons de vendas com essa composição e emitidos por uma farmácia em padrão muito acima dos demais estabelecimentos de sua região sugerem a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB correspondem a cupons de venda contendo quatro ou mais medicamentos. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com o mesmo perfil das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias localizadas em seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_ticket_medio_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de ticket medio."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "ticket medio")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "valor": as_float("val_ticket_medio"),
        "mediana_regiao": as_float("med_ticket_reg"),
        "mediana_uf": as_float("med_ticket_uf"),
        "mediana_brasil": as_float("med_ticket_br"),
        "multiplicador_regiao": as_float("risco_ticket_reg"),
        "multiplicador_uf": as_float("risco_ticket_uf"),
        "multiplicador_brasil": as_float("risco_ticket_br"),
    }


def _add_ticket_medio_text(doc, num: str, razao_social: str, ticket_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de ticket medio usando a matriz atual."""
    periodo_desc = ticket_comp["periodo_desc"]
    valor_fmt = _format_decimal_pt(ticket_comp["valor"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(ticket_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(ticket_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(ticket_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Valor do “ticket médio” dos medicamentos vendidos pela Farmácia {razao_social}, muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O comportamento esperado, no âmbito do PFPB, é de que o valor financeiro médio (“ticket médio”) das dispensações de medicamentos de uma farmácia para seus clientes acompanhe o padrão dos estabelecimentos de sua região, em um determinado período. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'A premissa é de baixa elasticidade-preço no mercado de medicamentos na região atendida pela farmácia credenciada. Nesse sentido, valores de ticket médio muito acima do padrão dos demais estabelecimentos sugerem a ocorrência de priorização de itens mais caros ou de uma maximização indevida nas vendas.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que o valor de ticket médio por ela registrado, {periodo_desc}, foi de ', color='0F172A', size=10)
    _run(p2, f'R$ {valor_fmt}', color='334155', size=10, bold=True)
    _run(p2, '. Tal valor corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' a mediana dos valores das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o valor equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o valor mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_receita_paciente_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de faturamento medio mensal por cliente."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "receita por paciente")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "valor": as_float("val_receita_paciente"),
        "mediana_regiao": as_float("med_receita_paciente_reg"),
        "mediana_uf": as_float("med_receita_paciente_uf"),
        "mediana_brasil": as_float("med_receita_paciente_br"),
        "multiplicador_regiao": as_float("risco_receita_paciente_reg"),
        "multiplicador_uf": as_float("risco_receita_paciente_uf"),
        "multiplicador_brasil": as_float("risco_receita_paciente_br"),
    }


def _add_receita_paciente_text(doc, num: str, razao_social: str, receita_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de faturamento medio mensal por cliente usando a matriz atual."""
    periodo_desc = receita_comp["periodo_desc"]
    valor_fmt = _format_decimal_pt(receita_comp["valor"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(receita_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(receita_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(receita_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Faturamento médio mensal por cliente, obtido pela Farmácia {razao_social}, muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O comportamento esperado, no âmbito do PFPB, é que o gasto médio mensal por cliente (CPF) em um estabelecimento farmacêutico, em determinado período, acompanhe o padrão das demais farmácias localizadas em sua mesma região. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Nesse sentido, receita média mensal por paciente obtida por uma farmácia que esteja muito acima do padrão dos demais estabelecimentos sugere a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que o valor médio mensal por cliente, {periodo_desc}, foi de ', color='0F172A', size=10)
    _run(p2, f'R$ {valor_fmt}', color='334155', size=10, bold=True)
    _run(p2, '. Tal valor corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o valor mediano das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o valor equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o valor mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_per_capita_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de faturamento mensal per capita."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "per capita")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "valor": as_float("val_per_capita"),
        "mediana_regiao": as_float("med_per_capita_reg"),
        "mediana_uf": as_float("med_per_capita_uf"),
        "mediana_brasil": as_float("med_per_capita_br"),
        "multiplicador_regiao": as_float("risco_per_capita_reg"),
        "multiplicador_uf": as_float("risco_per_capita_uf"),
        "multiplicador_brasil": as_float("risco_per_capita_br"),
    }


def _add_per_capita_text(doc, num: str, razao_social: str, per_capita_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de faturamento mensal per capita usando a matriz atual."""
    periodo_desc = per_capita_comp["periodo_desc"]
    valor_fmt = _format_decimal_pt(per_capita_comp["valor"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(per_capita_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(per_capita_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(per_capita_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Faturamento mensal per capita, obtido pela Farmácia {razao_social}, muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'O comportamento esperado para o faturamento mensal de uma farmácia, advindo do PFPB, é que ele guarde uma proporção razoável com a população do município onde está localizada. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Quando o estabelecimento apresenta valores de vendas per capita mensal muito desproporcionais às farmácias de sua região, tal comportamento sugere forte probabilidade de que ele esteja captando e utilizando CPFs de pessoas residentes em outras regiões.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que seu faturamento mensal per capita, {periodo_desc}, foi de ', color='0F172A', size=10)
    _run(p2, f'R$ {valor_fmt}', color='334155', size=10, bold=True)
    _run(p2, '. Tal valor corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o valor mediano das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o valor equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o valor mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_alto_custo_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de medicamentos de alto custo."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "alto custo")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_alto_custo"),
        "mediana_regiao": as_float("med_alto_custo_reg"),
        "mediana_uf": as_float("med_alto_custo_uf"),
        "mediana_brasil": as_float("med_alto_custo_br"),
        "multiplicador_regiao": as_float("risco_alto_custo_reg"),
        "multiplicador_uf": as_float("risco_alto_custo_uf"),
        "multiplicador_brasil": as_float("risco_alto_custo_br"),
    }


def _add_alto_custo_text(doc, num: str, razao_social: str, alto_custo_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de medicamentos de alto custo usando a matriz atual."""
    periodo_desc = alto_custo_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(alto_custo_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(alto_custo_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(alto_custo_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(alto_custo_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Vendas de medicamentos de alto custo realizadas pela Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'Os preços dos medicamentos, no âmbito do PFPB, variam em virtude dos seus princípios ativos. O comportamento esperado é que as vendas de uma farmácia apresentem um mix equilibrado de produtos, e não uma concentração atípica apenas em itens de alto valor financeiro. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Nesse sentido, a expectativa é que o percentual de venda pela farmácia de medicamentos de alto custo, correspondentes aos 10% mais caros da tabela do PFPB, acompanhe o padrão dos demais estabelecimentos localizados na mesma região.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB correspondem a medicamentos de alto custo. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com o mesmo perfil das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_vendas_rapidas_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de vendas em menos de 60 segundos."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "vendas rapidas")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_vendas_rapidas"),
        "mediana_regiao": as_float("med_vendas_rapidas_reg"),
        "mediana_uf": as_float("med_vendas_rapidas_uf"),
        "mediana_brasil": as_float("med_vendas_rapidas_br"),
        "multiplicador_regiao": as_float("risco_vendas_rapidas_reg"),
        "multiplicador_uf": as_float("risco_vendas_rapidas_uf"),
        "multiplicador_brasil": as_float("risco_vendas_rapidas_br"),
    }


def _add_vendas_rapidas_text(doc, num: str, razao_social: str, vendas_rapidas_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de vendas em menos de 60 segundos usando a matriz atual."""
    periodo_desc = vendas_rapidas_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(vendas_rapidas_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(vendas_rapidas_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(vendas_rapidas_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(vendas_rapidas_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(f'{num} Vendas de medicamentos em tempo inferior a 60 segundos', level=2)
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'No âmbito do PFPB, o comportamento esperado é que a dispensação de medicamento para o cidadão, no balcão da farmácia, seja realizada em alguns minutos, tendo em vista o tempo envolvido em um atendimento humano padrão, que envolve etapas logísticas como conferência de documentação, busca do produto e assinatura',
        color='0F172A',
        size=10,
    )
    _footnote_ref(
        doc,
        p1,
        16,
        'O art. 25 do Anexo LXXVIII da Portaria de Consolidação nº 5, de 28.09.2017, dispõe da seguinte maneira: “o paciente, obrigatoriamente, deve assinar o cupom vinculado, sendo que uma via deve ser mantida pelo estabelecimento e a outra entregue ao paciente.”',
    )
    _run(p1, '. ', color='0F172A', size=10)
    _run(
        p1,
        'Nesse sentido, dispensações de medicamentos sucessivas em intervalo de tempo inferior a 60 segundos e acima do padrão dos demais estabelecimentos localizados na mesma região sugerem a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB foram realizadas em tempo inferior a 60 segundos. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com essa mesma criticidade das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_recorrencia_sistemica_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de recorrencia sistemica de 30 dias."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "recorrencia sistemica")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_recorrencia_sistemica"),
        "mediana_regiao": as_float("med_recorrencia_sistemica_reg"),
        "mediana_uf": as_float("med_recorrencia_sistemica_uf"),
        "mediana_brasil": as_float("med_recorrencia_sistemica_br"),
        "multiplicador_regiao": as_float("risco_recorrencia_sistemica_reg"),
        "multiplicador_uf": as_float("risco_recorrencia_sistemica_uf"),
        "multiplicador_brasil": as_float("risco_recorrencia_sistemica_br"),
    }


def _add_recorrencia_sistemica_text(doc, num: str, razao_social: str, recorrencia_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de recorrencia sistemica de 30 dias usando a matriz atual."""
    periodo_desc = recorrencia_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(recorrencia_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(recorrencia_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(recorrencia_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(recorrencia_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Vendas de medicamentos com precisão absoluta de 30 dias realizadas pela Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'No âmbito do PFPB, os medicamentos têm limite exato de 30 dias para serem retirados pelos clientes. Uma tentativa de uma nova retirada de um medicamento, dentro desse prazo, ocasiona um bloqueio administrativo da sua dispensação. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Espera-se que as retiradas de medicamentos sigam um comportamento de consumo real, no qual as datas de retirada costumam variar alguns dias ao longo dos meses, e não que sejam realizadas sistematicamente com precisão absoluta de 30 em 30 dias. Nesse sentido, a identificação de vendas de medicamentos realizadas precisamente no prazo de 30 dias e com percentual acima do padrão dos demais estabelecimentos localizados na mesma região sugere indício de agendamento automatizado para vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB foram realizadas com prazos precisos de 30 dias. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com essa mesma criticidade das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_dias_pico_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de vendas em dias de pico."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "dias de pico")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_pico"),
        "mediana_regiao": as_float("med_pico_reg"),
        "mediana_uf": as_float("med_pico_uf"),
        "mediana_brasil": as_float("med_pico_br"),
        "multiplicador_regiao": as_float("risco_pico_reg"),
        "multiplicador_uf": as_float("risco_pico_uf"),
        "multiplicador_brasil": as_float("risco_pico_br"),
    }


def _add_dias_pico_text(doc, num: str, razao_social: str, dias_pico_comp: dict[str, Any], bookmark_name: str | None = None):
    """Adiciona texto analitico de vendas em dias de pico usando a matriz atual."""
    periodo_desc = dias_pico_comp["periodo_desc"]
    percentual_fmt = _format_decimal_pt(dias_pico_comp["percentual"], 2)
    multiplicador_reg_fmt = _format_decimal_pt(dias_pico_comp["multiplicador_regiao"], 2)
    multiplicador_uf_fmt = _format_decimal_pt(dias_pico_comp["multiplicador_uf"], 2)
    multiplicador_br_fmt = _format_decimal_pt(dias_pico_comp["multiplicador_brasil"], 2)

    heading = doc.add_heading(
        f'{num} Vendas de medicamentos em dias de pico realizadas pela Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'No âmbito do PFPB, o comportamento esperado para as farmácias é de que suas vendas sejam distribuídas de forma homogênea ao longo do mês. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'A identificação de registros baixos de venda na maior parte do tempo e altos volumes em dias específicos sugere a ocorrência de “lançamentos em lote”. Nesse sentido, a identificação de vendas em dias de pico realizadas por uma farmácia acima do padrão dos demais estabelecimentos localizados na mesma região sugere a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_fmt}%', color='334155', size=10, bold=True)
    _run(p2, ' das vendas de medicamentos por ela efetivadas no âmbito do PFPB foram realizadas em dias de pico. Tal percentual corresponde a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_reg_fmt} {_vez_ou_vezes(multiplicador_reg_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano de vendas com essa mesma criticidade das farmácias de sua região. ', color='0F172A', size=10)
    _run(p2, 'Ampliando-se o comparativo geográfico, o percentual equivale a ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_uf_fmt} {_vez_ou_vezes(multiplicador_uf_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o percentual mediano das farmácias de seu Estado e ', color='0F172A', size=10)
    _run(p2, f'{multiplicador_br_fmt} {_vez_ou_vezes(multiplicador_br_fmt)}', color='334155', size=10, bold=True)
    _run(p2, ' o das farmácias de todo o Brasil.', color='0F172A', size=10)


def _build_dispersao_geografica_context(
    cnpj: str,
    data_inicio: Optional[date],
    data_fim: Optional[date],
) -> dict[str, Any] | None:
    """Busca os dados atuais da matriz para o texto de dispersao geografica."""
    row = _get_matriz_row_context(cnpj, data_inicio, data_fim, "dispersao geografica")
    if row is None:
        return None

    def as_float(key: str) -> float:
        value = row.get(key)
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    if data_inicio and data_fim:
        periodo_desc = (
            f'no ano de {data_inicio.year}'
            if data_inicio.year == data_fim.year
            else f'no período de {data_inicio.year} a {data_fim.year}'
        )
    else:
        periodo_desc = 'no período analisado'

    cnpj_limpo = _cnpj_digits(cnpj)
    farmacias = get_df_dados_farmacia()
    _require_columns(farmacias, {"id_cnpj", "cnpj", "uf"}, "Cache de farmacias")
    farmacias = farmacias.with_columns(
        pl.col("cnpj").cast(pl.Utf8).str.replace_all(r"\D", "").alias("_cnpj_limpo")
    )
    farmacia_rows = farmacias.filter(pl.col("_cnpj_limpo") == cnpj_limpo)
    if farmacia_rows.is_empty():
        raise RuntimeError(f"CNPJ {cnpj_limpo} sem id_cnpj no cache de farmacias para dispersao geografica.")
    farmacia_row = farmacia_rows.row(0, named=True)
    id_cnpj = int(farmacia_row["id_cnpj"])

    origem_uf = (
        scan_geografico_origem_uf()
        .filter(pl.col("id_cnpj") == id_cnpj)
        .collect()
    )
    _require_columns(
        origem_uf,
        {
            "id_cnpj",
            "ano_base",
            "uf_farmacia",
            "uf_paciente",
            "is_outra_uf",
            "qtd_autorizacoes",
            "valor_autorizado",
        },
        "Cache geografico de origem por UF",
    )
    if data_inicio:
        origem_uf = origem_uf.filter(pl.col("ano_base") >= data_inicio.year)
    if data_fim:
        origem_uf = origem_uf.filter(pl.col("ano_base") <= data_fim.year)
    if origem_uf.is_empty():
        raise RuntimeError(
            f"Cache geografico_origem_uf sem registros para id_cnpj={id_cnpj} no periodo da Nota Tecnica."
        )

    origem_uf_agg = (
        origem_uf
        .group_by(["uf_farmacia", "uf_paciente", "is_outra_uf"])
        .agg([
            pl.sum("qtd_autorizacoes").alias("qtd_autorizacoes"),
            pl.sum("valor_autorizado").alias("valor_autorizado"),
        ])
        .sort("valor_autorizado", descending=True)
    )
    total_valor_origem = float(origem_uf_agg.get_column("valor_autorizado").sum() or 0.0)
    total_valor_outra_uf = float(
        origem_uf_agg
        .filter(pl.col("is_outra_uf") == True)
        .get_column("valor_autorizado")
        .sum()
        or 0.0
    )
    total_autorizacoes_origem = int(origem_uf_agg.get_column("qtd_autorizacoes").sum() or 0)
    total_autorizacoes_outra_uf = int(
        origem_uf_agg
        .filter(pl.col("is_outra_uf") == True)
        .get_column("qtd_autorizacoes")
        .sum()
        or 0
    )
    if total_valor_origem <= 0:
        raise RuntimeError(f"Cache geografico_origem_uf sem valor autorizado positivo para id_cnpj={id_cnpj}.")

    percentual_financeiro_outra_uf = (total_valor_outra_uf / total_valor_origem) * 100.0
    origem_rows = []
    for origem_row in origem_uf_agg.iter_rows(named=True):
        valor = float(origem_row["valor_autorizado"] or 0.0)
        is_outra = bool(origem_row["is_outra_uf"])
        origem_rows.append({
            "uf_farmacia": str(origem_row["uf_farmacia"] or "").strip().upper(),
            "uf_paciente": str(origem_row["uf_paciente"] or "").strip().upper(),
            "is_outra_uf": is_outra,
            "qtd_autorizacoes": int(origem_row["qtd_autorizacoes"] or 0),
            "valor_autorizado": valor,
            "percentual_sobre_total": (valor / total_valor_origem) * 100.0,
            "percentual_sobre_outra_uf": (
                (valor / total_valor_outra_uf) * 100.0
                if is_outra and total_valor_outra_uf > 0
                else None
            ),
        })

    return {
        "periodo_desc": periodo_desc,
        "percentual": as_float("pct_geografico"),
        "percentual_financeiro_outra_uf": percentual_financeiro_outra_uf,
        "mediana_regiao": as_float("med_geografico_reg"),
        "mediana_uf": as_float("med_geografico_uf"),
        "mediana_brasil": as_float("med_geografico_br"),
        "multiplicador_regiao": as_float("risco_geografico_reg"),
        "multiplicador_uf": as_float("risco_geografico_uf"),
        "multiplicador_brasil": as_float("risco_geografico_br"),
        "id_cnpj": id_cnpj,
        "uf_farmacia": str(farmacia_row["uf"] or "").strip().upper(),
        "origem_uf_rows": origem_rows,
        "total_valor_origem": total_valor_origem,
        "total_valor_outra_uf": total_valor_outra_uf,
        "total_autorizacoes_origem": total_autorizacoes_origem,
        "total_autorizacoes_outra_uf": total_autorizacoes_outra_uf,
    }


def _add_dispersao_geografica_origem_uf_table(doc, razao_social: str, dispersao_comp: dict[str, Any], tabela_num: int):
    rows = dispersao_comp.get("origem_uf_rows") or []
    if not rows:
        raise RuntimeError("Distribuicao geografica por UF obrigatoria ausente para a Nota Tecnica.")

    rows_tabela = rows[:12]
    restantes = rows[12:]
    if restantes:
        valor_restante = sum(float(row["valor_autorizado"] or 0.0) for row in restantes)
        autorizacoes_restante = sum(int(row["qtd_autorizacoes"] or 0) for row in restantes)
        total_valor = float(dispersao_comp["total_valor_origem"] or 0.0)
        total_outra = float(dispersao_comp["total_valor_outra_uf"] or 0.0)
        valor_outra_restante = sum(
            float(row["valor_autorizado"] or 0.0)
            for row in restantes
            if bool(row.get("is_outra_uf"))
        )
        rows_tabela.append({
            "uf_paciente": "Demais UFs",
            "is_outra_uf": any(bool(row.get("is_outra_uf")) for row in restantes),
            "qtd_autorizacoes": autorizacoes_restante,
            "valor_autorizado": valor_restante,
            "percentual_sobre_total": (valor_restante / total_valor) * 100.0 if total_valor > 0 else 0.0,
            "percentual_sobre_outra_uf": (
                (valor_outra_restante / total_outra) * 100.0
                if total_outra > 0 and valor_outra_restante > 0
                else None
            ),
        })

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.keep_together = True
    p_title.paragraph_format.space_before = Pt(16)
    p_title.paragraph_format.space_after = Pt(8)
    _run(
        p_title,
        f'Tabela {tabela_num} - Distribuição das autorizações por UF de residência do paciente vinculadas à Farmácia {razao_social}.',
        color='334155',
        size=8,
        bold=True,
    )

    table = doc.add_table(rows=len(rows_tabela) + 1, cols=5)
    table.style = 'Table Grid'
    _set_table_fixed_widths(
        table,
        [Inches(1.02), Inches(1.12), Inches(1.40), Inches(1.20), Inches(1.26)],
    )
    tbl_pr = table._tbl.tblPr
    table_jc = tbl_pr.find(qn('w:jc'))
    if table_jc is None:
        table_jc = OxmlElement('w:jc')
        tbl_pr.append(table_jc)
    table_jc.set(qn('w:val'), 'center')

    headers = [
        "UF paciente",
        "Autorizações",
        "Valor autorizado",
        "% do valor total",
        "% do valor em outras UFs",
    ]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=6.8, bold=True)
        _cell_bg(cell, 'E2E8F0')

    for row_idx, row in enumerate(rows_tabela, start=1):
        cells = table.rows[row_idx].cells
        pct_outra = row.get("percentual_sobre_outra_uf")
        values = [
            row["uf_paciente"],
            _format_int_pt(row["qtd_autorizacoes"]),
            _format_brl_pt(row["valor_autorizado"]),
            f'{_format_decimal_pt(float(row["percentual_sobre_total"]), 2)}%',
            f'{_format_decimal_pt(float(pct_outra), 2)}%' if pct_outra is not None else '—',
        ]
        fill = 'F8FAFC' if row_idx % 2 else 'FFFFFF'
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _cell_bg(cell, fill)
            _run(para, value, color='0F172A', size=6.8, bold=col_idx in (2, 3))

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.keep_together = True
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(8)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir das autorizações registradas no SAV/PFPB e da UF de residência do beneficiário constante na base de CPFs da Receita Federal do Brasil, utilizada no indicador geográfico.',
        color='64748B',
        size=7,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_dispersao_geografica_text(doc, num: str, razao_social: str, dispersao_comp: dict[str, Any], tabela_num: int, bookmark_name: str | None = None):
    """Adiciona texto analitico de vendas para residentes em outros Estados usando a matriz atual."""
    periodo_desc = dispersao_comp["periodo_desc"]
    percentual_financeiro_fmt = _format_decimal_pt(dispersao_comp["percentual_financeiro_outra_uf"], 2)
    total_aut_outra = _format_int_pt(dispersao_comp["total_autorizacoes_outra_uf"])
    total_valor_outra = _format_brl_pt(dispersao_comp["total_valor_outra_uf"])

    heading = doc.add_heading(
        f'{num} Vendas de medicamentos para pessoas residentes em outros Estados realizadas pela Farmácia {razao_social} com percentual sobre suas vendas totais muito superior ao dos estabelecimentos de sua região',
        level=2,
    )
    if bookmark_name:
        _add_bookmark(heading, bookmark_name)

    p1 = doc.add_paragraph()
    _run(
        p1,
        'No âmbito do PFPB, o comportamento esperado é que a grande maioria dos clientes atendidos pelas farmácias resida no mesmo Estado do estabelecimento. ',
        color='0F172A',
        size=10,
    )
    _run(
        p1,
        'Para tal verificação, é realizado o comparativo entre o endereço do beneficiário, contido na base do Cadastro de Pessoa Física (CPF), e o endereço de registro do próprio estabelecimento, contido no Cadastro Nacional de Pessoas Jurídicas (CNPJ). A identificação de vendas de medicamentos para pessoas de outros Estados acima do padrão dos demais estabelecimentos localizados na mesma região sugere a ocorrência de vendas fictícias.',
        color='0F172A',
        size=10,
    )

    p2 = doc.add_paragraph()
    _run(p2, f'Em relação à Farmácia {razao_social}, verificou-se que, {periodo_desc}, ', color='0F172A', size=10)
    _run(p2, f'{percentual_financeiro_fmt}%', color='334155', size=10, bold=True)
    _run(
        p2,
        ' do valor autorizado em vendas de medicamentos no âmbito do PFPB esteve associado a pessoas residentes em outros Estados. Esse percentual é calculado pela razão entre o valor autorizado para beneficiários de outras UFs e o valor autorizado total do estabelecimento no período analisado.',
        color='0F172A',
        size=10,
    )

    p3 = doc.add_paragraph()
    _run(
        p3,
        'No detalhamento por UF de residência dos beneficiários, foram identificadas ',
        color='0F172A',
        size=10,
    )
    _run(p3, total_aut_outra, color='334155', size=10, bold=True)
    _run(p3, ' autorizações para pessoas residentes em outros Estados, correspondentes a ', color='0F172A', size=10)
    _run(p3, total_valor_outra, color='334155', size=10, bold=True)
    _run(
        p3,
        ' em valor autorizado. A distribuição financeira por UF, apresentada na tabela e no mapa a seguir, permite identificar as origens que mais contribuíram para o comportamento atípico observado.',
        color='0F172A',
        size=10,
    )

    _add_dispersao_geografica_origem_uf_table(doc, razao_social, dispersao_comp, tabela_num)
    _add_mapa_geografico_origem_uf(doc, razao_social, dispersao_comp)
