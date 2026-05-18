import io
import html
from collections.abc import Sequence
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .nota_tecnica_docx_utils import _format_block_footnote, _format_block_title, _format_picture_paragraph, _run
from .nota_tecnica_formatters import _format_decimal_pt


def _axis_currency_label(value: float) -> str:
    """Formata valores compactos para eixo financeiro."""
    if value >= 1_000_000:
        return f'R$ {_format_decimal_pt(value / 1_000_000, 1)} mi'
    if value >= 1_000:
        return f'R$ {_format_decimal_pt(value / 1_000, 0)} mil'
    return f'R$ {_format_decimal_pt(value, 0)}'


def _nice_axis_max(value: float) -> float:
    """Calcula limite superior arredondado para o eixo Y."""
    if value <= 0:
        return 1.0
    magnitude = 10 ** (len(str(int(value))) - 1)
    normalized = value / magnitude
    if normalized <= 1:
        nice = 1
    elif normalized <= 2:
        nice = 2
    elif normalized <= 5:
        nice = 5
    else:
        nice = 10
    return nice * magnitude


def _svg_escape(value: Any) -> str:
    """Escapa texto para uso seguro em SVG."""
    escaped = html.escape(str(value or ""), quote=True)
    return escaped.encode("ascii", "xmlcharrefreplace").decode("ascii")


def _svg_to_png_stream(svg: str) -> io.BytesIO:
    """Converte SVG para PNG mantendo a insercao no DOCX compativel."""
    import resvg_py

    png_result: Any = resvg_py.svg_to_bytes(svg_string=svg, background="white")
    if not isinstance(png_result, bytes):
        raise RuntimeError(f"SVG to PNG conversion returned unsupported type: {type(png_result)!r}")
    png_bytes = png_result
    if not png_bytes:
        raise RuntimeError("SVG to PNG conversion returned empty PNG bytes")
    stream = io.BytesIO(png_bytes)
    stream.seek(0)
    return stream


def _svg_currency_axis_label(value: float) -> str:
    return _svg_escape(_axis_currency_label(value))


def _svg_point_path(points: list[tuple[float, float]]) -> str:
    if not points:
        return ""
    first_x, first_y = points[0]
    commands = [f"M {first_x:.2f} {first_y:.2f}"]
    commands.extend(f"L {x:.2f} {y:.2f}" for x, y in points[1:])
    return " ".join(commands)


def _svg_smooth_path(points: list[tuple[float, float]]) -> str:
    if not points:
        return ""
    if len(points) == 1:
        x, y = points[0]
        return f"M {x:.2f} {y:.2f}"

    commands = [f"M {points[0][0]:.2f} {points[0][1]:.2f}"]
    for idx in range(1, len(points)):
        x0, y0 = points[idx - 1]
        x1, y1 = points[idx]
        dx = (x1 - x0) * 0.5
        commands.append(
            f"C {x0 + dx:.2f} {y0:.2f}, {x1 - dx:.2f} {y1:.2f}, {x1:.2f} {y1:.2f}"
        )
    return " ".join(commands)


def _percentile_curve_anchors(x_values: list[int], y_values: list[float]) -> list[tuple[float, float]]:
    paired = sorted((float(x), y) for x, y in zip(x_values, y_values))
    if not paired:
        return [(1.0, 0.0), (100.0, 0.0)]

    unique: list[tuple[float, float]] = []
    seen: set[float] = set()
    for x, y in paired:
        if x in seen:
            continue
        seen.add(x)
        unique.append((x, y))

    cumulative: list[tuple[float, float]] = []
    current_y = 0.0
    for x, y in unique:
        current_y = max(current_y, y)
        cumulative.append((x, current_y))

    anchors: list[tuple[float, float]] = [(cumulative[0][0], cumulative[0][1])]
    start = 0
    tolerance = 1e-9
    for idx in range(1, len(cumulative) + 1):
        if idx == len(cumulative) or abs(cumulative[idx][1] - cumulative[start][1]) > tolerance:
            end = idx - 1
            center_x = (cumulative[start][0] + cumulative[end][0]) / 2
            anchors.append((center_x, cumulative[start][1]))
            start = idx
    anchors.append((cumulative[-1][0], cumulative[-1][1]))

    deduped: list[tuple[float, float]] = []
    for x, y in sorted(anchors):
        if deduped and abs(deduped[-1][0] - x) <= tolerance:
            deduped[-1] = (x, max(deduped[-1][1], y))
        else:
            deduped.append((x, y))
    return deduped if len(deduped) >= 2 else cumulative


def _add_gradient_bar_segment(
    ax,
    *,
    x_center: float,
    bottom: float,
    height: float,
    width: float,
    bottom_color: str,
    top_color: str,
    np,
    rounded: bool = False,
    zorder: int = 3,
):
    """Desenha um segmento de barra com degrade vertical recortado por patch."""
    if height <= 0:
        return

    from matplotlib.colors import to_rgba
    from matplotlib.patches import FancyBboxPatch, Rectangle

    x0 = x_center - width / 2
    y0 = bottom
    y1 = bottom + height

    if rounded:
        patch = FancyBboxPatch(
            (x0, y0),
            width,
            height,
            boxstyle=f"round,pad=0,rounding_size={min(width * 0.18, height * 0.18)}",
            linewidth=0,
            facecolor="none",
            transform=ax.transData,
            zorder=zorder,
        )
    else:
        patch = Rectangle(
            (x0, y0),
            width,
            height,
            linewidth=0,
            facecolor="none",
            transform=ax.transData,
            zorder=zorder,
        )
    ax.add_patch(patch)

    c0 = np.array(to_rgba(bottom_color))
    c1 = np.array(to_rgba(top_color))
    gradient = np.zeros((256, 1, 4), dtype=float)
    for idx, ratio in enumerate(np.linspace(0, 1, 256)):
        gradient[idx, 0, :] = c0 * (1 - ratio) + c1 * ratio

    image = ax.imshow(
        gradient,
        extent=[x0, x0 + width, y0, y1],
        origin="lower",
        aspect="auto",
        interpolation="bicubic",
        zorder=zorder,
    )
    image.set_clip_path(patch)


def _build_evolucao_financeira_chart(evolucao_comp: dict[str, Any]) -> io.BytesIO:
    """Gera grafico PNG de evolucao financeira com matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib.patches import Patch
    from matplotlib.ticker import FuncFormatter

    rows = evolucao_comp["rows"]
    labels = [row.get("semestre_fmt") or row.get("semestre") or "" for row in rows]
    regular_values = [max(float(row.get("regular") or 0.0), 0.0) for row in rows]
    irregular_values = [max(float(row.get("irregular") or 0.0), 0.0) for row in rows]
    total_values = [
        max(float(row.get("total") or (regular_values[idx] + irregular_values[idx])), 0.0)
        for idx, row in enumerate(rows)
    ]

    fig, ax = plt.subplots(figsize=(10.2, 5.0), dpi=130)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    text_color = "#334155"
    muted = "#64748B"
    grid_color = "#CBD5E1"
    regular_bottom = "#34C79A"
    regular_top = "#A7F3D0"
    irregular_bottom = "#F05A6E"
    irregular_top = "#FDA4AF"

    x_positions = np.arange(len(rows), dtype=float)
    bar_width = min(0.64, max(0.38, 7.0 / max(len(rows), 1)))

    for idx, x_pos in enumerate(x_positions):
        regular = regular_values[idx]
        irregular = irregular_values[idx]
        total = total_values[idx]
        if total <= 0:
            continue
        _add_gradient_bar_segment(
            ax,
            x_center=x_pos,
            bottom=0.0,
            height=regular,
            width=bar_width,
            bottom_color=regular_bottom,
            top_color=regular_top,
            np=np,
            rounded=irregular <= 0,
            zorder=3,
        )
        _add_gradient_bar_segment(
            ax,
            x_center=x_pos,
            bottom=regular,
            height=irregular,
            width=bar_width,
            bottom_color=irregular_bottom,
            top_color=irregular_top,
            np=np,
            rounded=True,
            zorder=4,
        )

    axis_max = _nice_axis_max(max(total_values or [1.0]) * 1.12)
    ax.set_ylim(0, axis_max)
    ax.set_xlim(-0.65, len(rows) - 0.35 if rows else 0.65)

    ax.set_title(
        "Evolução semestral das transferências e vendas sem comprovação",
        fontsize=12,
        fontweight="bold",
        color=text_color,
        pad=18,
    )
    ax.set_ylabel("Valor movimentado", fontsize=10.5, color=muted, labelpad=10)
    ax.set_xticks(x_positions)
    ax.set_xticklabels(labels, rotation=35, ha="right", fontsize=9.4, color=muted)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: _axis_currency_label(float(value))))
    ax.tick_params(axis="y", colors=muted, labelsize=9.5, length=0)
    ax.tick_params(axis="x", length=0)

    ax.grid(axis="y", color=grid_color, linestyle=(0, (5, 7)), linewidth=0.8, alpha=0.72)
    ax.grid(axis="x", visible=False)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#E2E8F0")
    ax.spines["bottom"].set_color("#E2E8F0")

    legend_handles = [
        Patch(facecolor=regular_top, edgecolor="none", label="Vendas regulares"),
        Patch(facecolor=irregular_bottom, edgecolor="none", label="Vendas sem comprovação"),
    ]
    legend = ax.legend(
        handles=legend_handles,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.02),
        ncol=2,
        frameon=False,
        fontsize=9.8,
        handlelength=1.4,
        columnspacing=1.8,
    )
    for text in legend.get_texts():
        text.set_color(muted)

    fig.tight_layout(pad=1.7)
    stream = io.BytesIO()
    fig.savefig(stream, format="png", dpi=130, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    stream.seek(0)
    return stream


def _monotone_smooth_curve(x_values: list[int], y_values: list[float], *, np) -> tuple[list[float], list[float]]:
    """Gera pontos intermediarios para uma curva visual suave e monotona."""
    if len(x_values) < 3 or len(x_values) != len(y_values):
        return [float(value) for value in x_values], list(y_values)

    x = np.asarray(x_values, dtype=float)
    y = np.asarray(y_values, dtype=float)
    order = np.argsort(x)
    x = x[order]
    y = y[order]

    unique_x, unique_idx = np.unique(x, return_index=True)
    x = unique_x
    y = y[unique_idx]
    n = len(x)
    if n < 3:
        return x.tolist(), y.tolist()

    y = np.maximum.accumulate(y)

    # Percentis de regioes pequenas costumam gerar muitos valores repetidos.
    # Para a figura, criamos ancoras no centro de cada patamar e interpolamos
    # entre elas com smoothstep. Isso adiciona pontos e remove o aspecto de escada.
    anchors_x = [float(x[0])]
    anchors_y = [float(y[0])]
    start = 0
    tolerance = 1e-9
    for idx in range(1, n + 1):
        if idx == n or abs(float(y[idx]) - float(y[start])) > tolerance:
            end = idx - 1
            center_x = float((x[start] + x[end]) / 2)
            anchors_x.append(center_x)
            anchors_y.append(float(y[start]))
            start = idx
    anchors_x.append(float(x[-1]))
    anchors_y.append(float(y[-1]))

    anchors_x = np.asarray(anchors_x, dtype=float)
    anchors_y = np.asarray(anchors_y, dtype=float)
    order = np.argsort(anchors_x)
    anchors_x = anchors_x[order]
    anchors_y = np.maximum.accumulate(anchors_y[order])

    unique_anchor_x, unique_anchor_idx = np.unique(anchors_x, return_index=True)
    anchors_x = unique_anchor_x
    anchors_y = anchors_y[unique_anchor_idx]
    if len(anchors_x) < 2:
        return x.tolist(), y.tolist()

    x_smooth = np.linspace(float(anchors_x[0]), float(anchors_x[-1]), max(1200, len(anchors_x) * 80))
    segment_idx = np.searchsorted(anchors_x, x_smooth, side="right") - 1
    segment_idx = np.clip(segment_idx, 0, len(anchors_x) - 2)
    x0 = anchors_x[segment_idx]
    x1 = anchors_x[segment_idx + 1]
    y0 = anchors_y[segment_idx]
    y1 = anchors_y[segment_idx + 1]
    t = (x_smooth - x0) / np.maximum(x1 - x0, 1e-9)
    t = np.clip(t, 0.0, 1.0)
    eased = (3 * t**2) - (2 * t**3)
    y_smooth = y0 + (y1 - y0) * eased

    y_smooth = np.maximum.accumulate(y_smooth)
    y_smooth = np.clip(y_smooth, min(float(np.nanmin(y)), 0.0), max(float(np.nanmax(y)), 100.0))
    return x_smooth.tolist(), y_smooth.tolist()


def _add_gradient_area(ax, x_values: Sequence[float], y_values: Sequence[float], *, np, color_hex: str, y_max: float):
    """Preenche a area sob a curva com degrade vertical recortado pelo poligono."""
    from matplotlib.colors import to_rgba
    from matplotlib.path import Path
    from matplotlib.patches import PathPatch

    x = np.asarray(x_values, dtype=float)
    y = np.asarray(y_values, dtype=float)
    if len(x) < 2:
        return

    rgba = np.array(to_rgba(color_hex))
    gradient = np.zeros((256, 1, 4), dtype=float)
    gradient[:, 0, :3] = rgba[:3]
    gradient[:, 0, 3] = np.linspace(0.03, 0.24, 256)

    image = ax.imshow(
        gradient,
        extent=[float(np.nanmin(x)), float(np.nanmax(x)), 0, y_max],
        origin="lower",
        aspect="auto",
        interpolation="bicubic",
        zorder=2,
    )

    vertices = [(float(x[0]), 0.0), *zip(x.astype(float), y.astype(float)), (float(x[-1]), 0.0), (float(x[0]), 0.0)]
    codes = [Path.MOVETO] + [Path.LINETO] * len(x) + [Path.LINETO, Path.CLOSEPOLY]
    clip_path = PathPatch(Path(vertices, codes), facecolor="none", edgecolor="none", transform=ax.transData)
    ax.add_patch(clip_path)
    image.set_clip_path(clip_path)


def _build_percentil_risco_chart(percentil_comp: dict[str, Any]) -> io.BytesIO:
    """Gera grafico PNG da posicao percentilica do CNPJ com matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib.ticker import FuncFormatter

    percentiles = percentil_comp.get("percentiles") or []
    x_values = [int(point.get("percentile") or 0) for point in percentiles]
    y_values = [float(point.get("score") or 0.0) for point in percentiles]
    current_value = float(percentil_comp.get("current_value") or 0.0)
    percentile_rank = int(percentil_comp.get("percentile_rank") or 100)
    metric_label = percentil_comp.get("metric_label") or "% de vendas sem comprovação"

    fig, ax = plt.subplots(figsize=(10.2, 5.0), dpi=130)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    line_color = "#E11D48"
    area_color = "#F43F5E"
    text_color = "#334155"
    muted = "#64748B"
    grid_color = "#CBD5E1"

    x_curve, y_curve = _monotone_smooth_curve(x_values, y_values, np=np)
    ax.plot(x_curve, y_curve, color=line_color, linewidth=2.8, solid_capstyle="round", zorder=3)

    marker_y = current_value
    ax.axvline(percentile_rank, color=line_color, linestyle=(0, (5, 5)), linewidth=1.5, alpha=0.72, zorder=1)
    ax.axhline(marker_y, color=line_color, linestyle=(0, (5, 5)), linewidth=1.3, alpha=0.42, zorder=1)
    ax.scatter([percentile_rank], [marker_y], s=94, color="white", edgecolor=line_color, linewidth=2.4, zorder=5)
    ax.scatter([percentile_rank], [marker_y], s=28, color=line_color, zorder=6)

    current_txt = f"{_format_decimal_pt(current_value, 1)}%"
    badge_txt = f"Estabelecimento\nPercentil {percentile_rank} - {current_txt}"
    ax.annotate(
        badge_txt,
        xy=(percentile_rank, marker_y),
        xytext=(18, 28),
        textcoords="offset points",
        fontsize=10,
        fontweight="bold",
        color=line_color,
        bbox=dict(boxstyle="round,pad=0.55,rounding_size=0.25", fc="white", ec=line_color, lw=1.35),
        arrowprops=dict(arrowstyle="-", color=line_color, lw=1.2, shrinkA=0, shrinkB=8),
        zorder=7,
    )

    ax.set_title(
        "Distribuição percentílica regional do percentual de vendas sem comprovação",
        fontsize=12,
        fontweight="bold",
        color=text_color,
        pad=18,
    )
    ax.set_xlabel("Percentil dos estabelecimentos da Região de Saúde", fontsize=10.5, color=muted, labelpad=10)
    ax.set_ylabel(metric_label, fontsize=10.5, color=muted, labelpad=10)

    max_y = max([current_value, *y_values, 1.0])
    upper = min(100.0, max_y * 1.16 if max_y < 90 else 100.0)
    ax.set_xlim(1, 100)
    ax.set_ylim(0, upper)
    _add_gradient_area(ax, x_curve, y_curve, np=np, color_hex=area_color, y_max=upper)
    ax.set_xticks([1, 20, 40, 60, 80, 100])
    ax.set_xticklabels(["1%", "20%", "40%", "60%", "80%", "100%"])
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{_format_decimal_pt(value, 0)}%"))

    ax.grid(axis="y", color=grid_color, linestyle=(0, (5, 7)), linewidth=0.8, alpha=0.72)
    ax.grid(axis="x", visible=False)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#E2E8F0")
    ax.spines["bottom"].set_color("#E2E8F0")
    ax.tick_params(axis="both", colors=muted, labelsize=9.5, length=0)

    fig.tight_layout(pad=1.7)
    stream = io.BytesIO()
    fig.savefig(stream, format="png", dpi=130, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    stream.seek(0)
    return stream


def _build_posicionamento_regional_chart(posicionamento_comp: dict[str, Any]) -> io.BytesIO:
    """Gera scatter PNG do posicionamento regional com matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib.ticker import FuncFormatter

    rows = posicionamento_comp.get("rows") or []
    current = posicionamento_comp.get("current") or {}
    others = [row for row in rows if not row.get("is_current")]

    other_x = [float(row.get("total_mov") or 0.0) for row in others]
    other_y = [float(row.get("pct_sem_comprovacao") or 0.0) for row in others]
    current_x = float(current.get("total_mov") or 0.0)
    current_y = float(current.get("pct_sem_comprovacao") or 0.0)

    all_x = [float(row.get("total_mov") or 0.0) for row in rows]
    all_y = [float(row.get("pct_sem_comprovacao") or 0.0) for row in rows]
    median_y = float(np.median(all_y)) if all_y else 0.0
    x_max = _nice_axis_max(max([*all_x, 1.0]) * 1.12)
    y_max = min(100.0, max([*all_y, 1.0]) * 1.16 if max([*all_y, 1.0]) < 90 else 100.0)
    y_max = max(y_max, 10.0)

    fig, ax = plt.subplots(figsize=(10.2, 5.0), dpi=130)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    text_color = "#334155"
    muted = "#64748B"
    grid_color = "#CBD5E1"
    other_color = "#94A3B8"
    current_color = "#F05A6E"

    ax.scatter(
        other_x,
        other_y,
        s=42,
        color=other_color,
        alpha=0.58,
        edgecolors="white",
        linewidths=0.8,
        zorder=3,
        label="Outras farmácias",
    )
    ax.scatter(
        [current_x],
        [current_y],
        s=170,
        color=current_color,
        alpha=0.96,
        edgecolors="white",
        linewidths=2.0,
        zorder=5,
        label="Estabelecimento analisado",
    )
    ax.scatter(
        [current_x],
        [current_y],
        s=260,
        facecolors="none",
        edgecolors=current_color,
        linewidths=1.5,
        alpha=0.28,
        zorder=4,
    )

    ax.axhline(
        median_y,
        color="#F59E0B",
        linestyle=(0, (5, 5)),
        linewidth=1.2,
        alpha=0.7,
        zorder=1,
    )
    ax.text(
        0.99,
        median_y,
        f"Mediana regional: {_format_decimal_pt(median_y, 2)}%",
        transform=ax.get_yaxis_transform(),
        ha="right",
        va="bottom",
        fontsize=8.8,
        color="#B45309",
    )

    label = f"Estabelecimento\n{_axis_currency_label(current_x)} | {_format_decimal_pt(current_y, 1)}%"
    label_dx = -22 if current_x > x_max * 0.72 else 20
    label_dy = -26 if current_y > y_max * 0.76 else 22
    ax.annotate(
        label,
        xy=(current_x, current_y),
        xytext=(label_dx, label_dy),
        textcoords="offset points",
        fontsize=9.4,
        fontweight="bold",
        color=current_color,
        ha="right" if label_dx < 0 else "left",
        va="top" if label_dy < 0 else "bottom",
        bbox=dict(boxstyle="round,pad=0.5,rounding_size=0.22", fc="white", ec=current_color, lw=1.2),
        arrowprops=dict(arrowstyle="-", color=current_color, lw=1.1, shrinkA=0, shrinkB=9),
        zorder=7,
    )

    ax.set_xlim(0, x_max)
    ax.set_ylim(0, y_max)

    ax.set_title(
        "Posicionamento regional por volume e percentual sem comprovação",
        fontsize=12,
        fontweight="bold",
        color=text_color,
        pad=18,
    )
    ax.set_xlabel("Valor total movimentado", fontsize=10.5, color=muted, labelpad=10)
    ax.set_ylabel(posicionamento_comp.get("metric_label") or "% de dispensações sem comprovação", fontsize=10.5, color=muted, labelpad=10)
    ax.xaxis.set_major_formatter(FuncFormatter(lambda value, _: _axis_currency_label(float(value))))
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{_format_decimal_pt(value, 0)}%"))
    ax.tick_params(axis="both", colors=muted, labelsize=9.5, length=0)

    ax.grid(axis="both", color=grid_color, linestyle=(0, (5, 7)), linewidth=0.8, alpha=0.62)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#E2E8F0")
    ax.spines["bottom"].set_color("#E2E8F0")

    legend = ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, 1.02),
        ncol=2,
        frameon=False,
        fontsize=9.6,
        columnspacing=1.8,
        handletextpad=0.45,
    )
    for text in legend.get_texts():
        text.set_color(muted)

    fig.tight_layout(pad=1.7)
    stream = io.BytesIO()
    fig.savefig(stream, format="png", dpi=130, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    stream.seek(0)
    return stream


def _build_evolucao_financeira_chart_svg(evolucao_comp: dict[str, Any]) -> str:
    """Gera SVG da evolucao financeira para conversao posterior em PNG."""
    rows = evolucao_comp["rows"]
    width, height = 1326, 650
    left, right, top, bottom = 120, 56, 110, 120
    plot_w = width - left - right
    plot_h = height - top - bottom
    plot_bottom = top + plot_h
    text_color = "#334155"
    muted = "#64748B"
    grid = "#CBD5E1"

    max_total = max(float(row.get("total") or 0.0) for row in rows) if rows else 1.0
    axis_max = _nice_axis_max(max_total * 1.10)
    tick_count = 5
    n = max(len(rows), 1)
    slot = plot_w / n
    bar_w = min(88, max(24, slot * 0.82))

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        "<defs>",
        '<linearGradient id="regularBar" x1="0" y1="0" x2="0" y2="1"><stop offset="0%" stop-color="#B8F5DC"/><stop offset="100%" stop-color="#34D399"/></linearGradient>',
        '<linearGradient id="irregularBar" x1="0" y1="0" x2="0" y2="1"><stop offset="0%" stop-color="#FDA4AF"/><stop offset="100%" stop-color="#F0526B"/></linearGradient>',
        "</defs>",
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        f'<text x="{width / 2:.0f}" y="42" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="22" font-weight="700" fill="{text_color}">Evolu&#231;&#227;o semestral das transfer&#234;ncias e vendas sem comprova&#231;&#227;o</text>',
        f'<rect x="{width / 2 - 215:.0f}" y="68" width="24" height="12" fill="url(#regularBar)"/>',
        f'<text x="{width / 2 - 181:.0f}" y="79" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">Vendas regulares</text>',
        f'<rect x="{width / 2 + 20:.0f}" y="68" width="24" height="12" fill="url(#irregularBar)"/>',
        f'<text x="{width / 2 + 54:.0f}" y="79" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">Vendas sem comprova&#231;&#227;o</text>',
        f'<text x="31" y="{top + plot_h / 2:.0f}" transform="rotate(-90 31 {top + plot_h / 2:.0f})" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="16" fill="{muted}">Valor movimentado</text>',
    ]

    for idx in range(tick_count + 1):
        value = axis_max * idx / tick_count
        y = plot_bottom - (value / axis_max) * plot_h
        parts.append(f'<line x1="{left}" y1="{y:.2f}" x2="{width - right}" y2="{y:.2f}" stroke="{grid}" stroke-width="1" stroke-dasharray="6 8" opacity="0.75"/>')
        parts.append(f'<text x="{left - 14}" y="{y + 5:.2f}" text-anchor="end" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">{_svg_currency_axis_label(value)}</text>')

    for idx, row in enumerate(rows):
        center_x = left + slot * idx + slot / 2
        x = center_x - bar_w / 2
        regular = max(float(row.get("regular") or 0.0), 0.0)
        irregular = max(float(row.get("irregular") or 0.0), 0.0)
        total = max(float(row.get("total") or (regular + irregular)), 0.0)
        if total > 0:
            total_h = max(2.0, (total / axis_max) * plot_h)
            regular_h = (regular / axis_max) * plot_h
            irregular_h = max(0.0, total_h - regular_h)
            if regular_h > 0:
                y_regular = plot_bottom - regular_h
                parts.append(f'<rect x="{x:.2f}" y="{y_regular:.2f}" width="{bar_w:.2f}" height="{regular_h:.2f}" fill="url(#regularBar)"/>')
            if irregular_h > 0:
                y_total = plot_bottom - total_h
                parts.append(f'<rect x="{x:.2f}" y="{y_total:.2f}" width="{bar_w:.2f}" height="{irregular_h:.2f}" fill="url(#irregularBar)"/>')

        label = _svg_escape(row.get("semestre_fmt") or row.get("semestre") or "")
        label = label.replace("&#186;", "&#176;").replace("&#170;", "&#176;")
        parts.append(f'<text x="{center_x + 8:.2f}" y="{plot_bottom + 42}" transform="rotate(-38 {center_x + 8:.2f} {plot_bottom + 42})" text-anchor="end" font-family="Segoe UI, Arial, sans-serif" font-size="14" font-weight="700" fill="{muted}">{label}</text>')

    parts.append(f'<line x1="{left}" y1="{plot_bottom}" x2="{width - right}" y2="{plot_bottom}" stroke="#E2E8F0" stroke-width="1"/>')
    parts.append("</svg>")
    return "".join(parts)


def _build_percentil_risco_chart_svg(percentil_comp: dict[str, Any]) -> str:
    """Gera SVG do grafico de percentil para conversao posterior em PNG."""
    percentiles = percentil_comp.get("percentiles") or []
    x_values = [int(point.get("percentile") or 0) for point in percentiles]
    y_values = [float(point.get("score") or 0.0) for point in percentiles]
    current_value = float(percentil_comp.get("current_value") or 0.0)
    percentile_rank = int(percentil_comp.get("percentile_rank") or 100)
    metric_label = _svg_escape(percentil_comp.get("metric_label") or "% de vendas sem comprovacao")

    width, height = 1326, 650
    left, right, top, bottom = 96, 56, 106, 90
    plot_w = width - left - right
    plot_h = height - top - bottom
    plot_bottom = top + plot_h
    line_color = "#E11D48"
    area_color = "#F43F5E"
    text_color = "#334155"
    muted = "#64748B"
    grid = "#CBD5E1"
    max_y = max([current_value, *y_values, 1.0])
    upper = min(100.0, max_y * 1.16 if max_y < 90 else 100.0)

    def sx(value: float) -> float:
        return left + ((value - 1.0) / 99.0) * plot_w

    def sy(value: float) -> float:
        return plot_bottom - (max(0.0, min(value, upper)) / upper) * plot_h

    anchors = _percentile_curve_anchors(x_values, y_values)
    curve = [(sx(x), sy(y)) for x, y in anchors]
    curve_path = _svg_smooth_path(curve)
    area_path = f"{curve_path} L {curve[-1][0]:.2f} {plot_bottom:.2f} L {curve[0][0]:.2f} {plot_bottom:.2f} Z" if curve else ""
    marker_x = sx(float(percentile_rank))
    marker_y = sy(current_value)
    badge_x = min(marker_x + 28, width - right - 245)
    badge_y = max(marker_y - 62, top + 12)
    current_txt = _svg_escape(f"{_format_decimal_pt(current_value, 1)}%")

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        "<defs>",
        '<linearGradient id="percentilArea" x1="0" y1="0" x2="0" y2="1"><stop offset="0%" stop-color="#F43F5E" stop-opacity="0.24"/><stop offset="100%" stop-color="#F43F5E" stop-opacity="0.03"/></linearGradient>',
        "</defs>",
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        f'<text x="{width / 2:.0f}" y="42" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="22" font-weight="700" fill="{text_color}">Distribui&#231;&#227;o percent&#237;lica regional do percentual de vendas sem comprova&#231;&#227;o</text>',
    ]
    for idx in range(6):
        value = upper * idx / 5
        y = sy(value)
        parts.append(f'<line x1="{left}" y1="{y:.2f}" x2="{width - right}" y2="{y:.2f}" stroke="{grid}" stroke-width="1" stroke-dasharray="6 8" opacity="0.75"/>')
        parts.append(f'<text x="{left - 14}" y="{y + 5:.2f}" text-anchor="end" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">{_svg_escape(_format_decimal_pt(value, 0))}%</text>')

    if area_path:
        parts.append(f'<path d="{area_path}" fill="url(#percentilArea)"/>')
    if curve_path:
        parts.append(f'<path d="{curve_path}" fill="none" stroke="{line_color}" stroke-width="4" stroke-linecap="round" stroke-linejoin="round"/>')

    parts.extend([
        f'<line x1="{marker_x:.2f}" y1="{top}" x2="{marker_x:.2f}" y2="{plot_bottom}" stroke="{line_color}" stroke-width="2" stroke-dasharray="6 6" opacity="0.7"/>',
        f'<line x1="{left}" y1="{marker_y:.2f}" x2="{width - right}" y2="{marker_y:.2f}" stroke="{line_color}" stroke-width="2" stroke-dasharray="6 6" opacity="0.36"/>',
        f'<circle cx="{marker_x:.2f}" cy="{marker_y:.2f}" r="12" fill="#FFFFFF" stroke="{line_color}" stroke-width="4"/>',
        f'<circle cx="{marker_x:.2f}" cy="{marker_y:.2f}" r="5" fill="{line_color}"/>',
        f'<line x1="{marker_x:.2f}" y1="{marker_y:.2f}" x2="{badge_x:.2f}" y2="{badge_y + 42:.2f}" stroke="{line_color}" stroke-width="2" opacity="0.8"/>',
        f'<rect x="{badge_x:.2f}" y="{badge_y:.2f}" width="230" height="64" rx="9" fill="#FFFFFF" stroke="{line_color}" stroke-width="2"/>',
        f'<text x="{badge_x + 14:.2f}" y="{badge_y + 25:.2f}" font-family="Segoe UI, Arial, sans-serif" font-size="16" font-weight="700" fill="{line_color}">Estabelecimento</text>',
        f'<text x="{badge_x + 14:.2f}" y="{badge_y + 49:.2f}" font-family="Segoe UI, Arial, sans-serif" font-size="16" font-weight="700" fill="{line_color}">Percentil {percentile_rank} - {current_txt}</text>',
    ])
    for tick in [1, 20, 40, 60, 80, 100]:
        x = sx(float(tick))
        parts.append(f'<text x="{x:.2f}" y="{plot_bottom + 34}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">{tick}%</text>')
    parts.extend([
        f'<line x1="{left}" y1="{plot_bottom}" x2="{width - right}" y2="{plot_bottom}" stroke="#E2E8F0" stroke-width="1"/>',
        f'<line x1="{left}" y1="{top}" x2="{left}" y2="{plot_bottom}" stroke="#E2E8F0" stroke-width="1"/>',
        f'<text x="{width / 2:.0f}" y="{height - 18}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="16" fill="{muted}">Percentil dos estabelecimentos da Regi&#227;o de Sa&#250;de</text>',
        f'<text x="24" y="{top + plot_h / 2:.0f}" transform="rotate(-90 24 {top + plot_h / 2:.0f})" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="16" fill="{muted}">{metric_label}</text>',
        "</svg>",
    ])
    return "".join(parts)


def _build_posicionamento_regional_chart_svg(posicionamento_comp: dict[str, Any]) -> str:
    """Gera SVG do posicionamento regional para conversao posterior em PNG."""
    rows = posicionamento_comp.get("rows") or []
    current = posicionamento_comp.get("current") or {}
    others = [row for row in rows if not row.get("is_current")]
    all_x = [float(row.get("total_mov") or 0.0) for row in rows]
    all_y = [float(row.get("pct_sem_comprovacao") or 0.0) for row in rows]
    current_x = float(current.get("total_mov") or 0.0)
    current_y = float(current.get("pct_sem_comprovacao") or 0.0)
    sorted_y = sorted(all_y)
    mid = len(sorted_y) // 2
    median_y = ((sorted_y[mid - 1] + sorted_y[mid]) / 2) if sorted_y and len(sorted_y) % 2 == 0 else (sorted_y[mid] if sorted_y else 0.0)
    x_max = _nice_axis_max(max([*all_x, 1.0]) * 1.12)
    max_y = max([*all_y, 1.0])
    y_max = min(100.0, max_y * 1.16 if max_y < 90 else 100.0)
    y_max = max(y_max, 10.0)

    width, height = 1326, 650
    left, right, top, bottom = 112, 56, 106, 90
    plot_w = width - left - right
    plot_h = height - top - bottom
    plot_bottom = top + plot_h
    text_color = "#334155"
    muted = "#64748B"
    grid = "#CBD5E1"
    other_color = "#94A3B8"
    current_color = "#F05A6E"

    def sx(value: float) -> float:
        return left + (max(0.0, min(value, x_max)) / x_max) * plot_w

    def sy(value: float) -> float:
        return plot_bottom - (max(0.0, min(value, y_max)) / y_max) * plot_h

    marker_x = sx(current_x)
    marker_y = sy(current_y)
    label_left = marker_x > left + plot_w * 0.72
    label_top = marker_y < top + plot_h * 0.24
    badge_w, badge_h = 250, 64
    badge_x = marker_x - badge_w - 28 if label_left else marker_x + 28
    badge_y = marker_y + 22 if label_top else marker_y - badge_h - 22
    badge_x = max(left + 8, min(badge_x, width - right - badge_w))
    badge_y = max(top + 8, min(badge_y, plot_bottom - badge_h - 8))

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
        f'<text x="{width / 2:.0f}" y="42" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="22" font-weight="700" fill="{text_color}">Posicionamento regional por volume e percentual sem comprova&#231;&#227;o</text>',
    ]
    for idx in range(6):
        value = y_max * idx / 5
        y = sy(value)
        parts.append(f'<line x1="{left}" y1="{y:.2f}" x2="{width - right}" y2="{y:.2f}" stroke="{grid}" stroke-width="1" stroke-dasharray="6 8" opacity="0.65"/>')
        parts.append(f'<text x="{left - 14}" y="{y + 5:.2f}" text-anchor="end" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">{_svg_escape(_format_decimal_pt(value, 0))}%</text>')
    for idx in range(6):
        value = x_max * idx / 5
        x = sx(value)
        parts.append(f'<line x1="{x:.2f}" y1="{top}" x2="{x:.2f}" y2="{plot_bottom}" stroke="{grid}" stroke-width="1" stroke-dasharray="6 8" opacity="0.45"/>')
        parts.append(f'<text x="{x:.2f}" y="{plot_bottom + 34}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">{_svg_currency_axis_label(value)}</text>')

    median_screen_y = sy(median_y)
    parts.extend([
        f'<line x1="{left}" y1="{median_screen_y:.2f}" x2="{width - right}" y2="{median_screen_y:.2f}" stroke="#F59E0B" stroke-width="2" stroke-dasharray="6 6" opacity="0.72"/>',
        f'<text x="{width - right - 8}" y="{median_screen_y - 8:.2f}" text-anchor="end" font-family="Segoe UI, Arial, sans-serif" font-size="14" fill="#B45309">Mediana regional: {_svg_escape(_format_decimal_pt(median_y, 2))}%</text>',
    ])
    for row in others:
        x = sx(float(row.get("total_mov") or 0.0))
        y = sy(float(row.get("pct_sem_comprovacao") or 0.0))
        parts.append(f'<circle cx="{x:.2f}" cy="{y:.2f}" r="6" fill="{other_color}" opacity="0.58" stroke="#FFFFFF" stroke-width="1.4"/>')

    current_label = f"{_axis_currency_label(current_x)} | {_format_decimal_pt(current_y, 1)}%"
    parts.extend([
        f'<circle cx="{marker_x:.2f}" cy="{marker_y:.2f}" r="19" fill="none" stroke="{current_color}" stroke-width="2" opacity="0.28"/>',
        f'<circle cx="{marker_x:.2f}" cy="{marker_y:.2f}" r="13" fill="{current_color}" opacity="0.96" stroke="#FFFFFF" stroke-width="3"/>',
        f'<line x1="{marker_x:.2f}" y1="{marker_y:.2f}" x2="{badge_x + (0 if label_left else badge_w):.2f}" y2="{badge_y + badge_h / 2:.2f}" stroke="{current_color}" stroke-width="2" opacity="0.85"/>',
        f'<rect x="{badge_x:.2f}" y="{badge_y:.2f}" width="{badge_w}" height="{badge_h}" rx="9" fill="#FFFFFF" stroke="{current_color}" stroke-width="2"/>',
        f'<text x="{badge_x + 14:.2f}" y="{badge_y + 25:.2f}" font-family="Segoe UI, Arial, sans-serif" font-size="16" font-weight="700" fill="{current_color}">Estabelecimento</text>',
        f'<text x="{badge_x + 14:.2f}" y="{badge_y + 49:.2f}" font-family="Segoe UI, Arial, sans-serif" font-size="16" font-weight="700" fill="{current_color}">{_svg_escape(current_label)}</text>',
        f'<circle cx="{width / 2 - 170}" cy="73" r="7" fill="{other_color}" opacity="0.58" stroke="#FFFFFF" stroke-width="1.4"/>',
        f'<text x="{width / 2 - 154}" y="78" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">Outras farm&#225;cias</text>',
        f'<circle cx="{width / 2 + 42}" cy="73" r="9" fill="{current_color}" stroke="#FFFFFF" stroke-width="2"/>',
        f'<text x="{width / 2 + 60}" y="78" font-family="Segoe UI, Arial, sans-serif" font-size="15" fill="{muted}">Estabelecimento analisado</text>',
        f'<line x1="{left}" y1="{plot_bottom}" x2="{width - right}" y2="{plot_bottom}" stroke="#E2E8F0" stroke-width="1"/>',
        f'<line x1="{left}" y1="{top}" x2="{left}" y2="{plot_bottom}" stroke="#E2E8F0" stroke-width="1"/>',
        f'<text x="{width / 2:.0f}" y="{height - 18}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="16" fill="{muted}">Valor total movimentado</text>',
        f'<text x="24" y="{top + plot_h / 2:.0f}" transform="rotate(-90 24 {top + plot_h / 2:.0f})" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="16" fill="{muted}">{_svg_escape(posicionamento_comp.get("metric_label") or "% de dispensacoes sem comprovacao")}</text>',
        "</svg>",
    ])
    return "".join(parts)


def _build_evolucao_financeira_chart_svg_png(evolucao_comp: dict[str, Any]) -> io.BytesIO:
    return _svg_to_png_stream(_build_evolucao_financeira_chart_svg(evolucao_comp))


def _build_percentil_risco_chart_svg_png(percentil_comp: dict[str, Any]) -> io.BytesIO:
    return _svg_to_png_stream(_build_percentil_risco_chart_svg(percentil_comp))


def _build_posicionamento_regional_chart_svg_png(posicionamento_comp: dict[str, Any]) -> io.BytesIO:
    return _svg_to_png_stream(_build_posicionamento_regional_chart_svg(posicionamento_comp))


def _build_evolucao_financeira_chart_prefer_svg(evolucao_comp: dict[str, Any]) -> io.BytesIO:
    try:
        return _build_evolucao_financeira_chart_svg_png(evolucao_comp)
    except Exception:
        return _build_evolucao_financeira_chart(evolucao_comp)


def _build_percentil_risco_chart_prefer_svg(percentil_comp: dict[str, Any]) -> io.BytesIO:
    try:
        return _build_percentil_risco_chart_svg_png(percentil_comp)
    except Exception:
        return _build_percentil_risco_chart(percentil_comp)


def _build_posicionamento_regional_chart_prefer_svg(posicionamento_comp: dict[str, Any]) -> io.BytesIO:
    try:
        return _build_posicionamento_regional_chart_svg_png(posicionamento_comp)
    except Exception:
        return _build_posicionamento_regional_chart(posicionamento_comp)


def _format_figure_title(paragraph):
    _format_block_title(paragraph, space_before=18, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _format_figure_footnote(paragraph):
    _format_block_footnote(paragraph, space_before=5, space_after=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_figura_posicionamento_regional(doc, razao_social: str, cnpj_fmt: str, posicionamento_comp: dict[str, Any], figure_number: int = 1):
    """Insere figura de posicionamento regional no documento."""
    p_title = doc.add_paragraph()
    _format_figure_title(p_title)
    _run(
        p_title,
        f'Figura {figure_number:02d} - Posicionamento regional da Farmácia {razao_social} (CNPJ {cnpj_fmt}) em relação aos estabelecimentos da mesma Região de Saúde.',
        color='0F172A',
        size=9,
        bold=True,
    )

    chart_stream = _build_posicionamento_regional_chart_prefer_svg(posicionamento_comp)
    p_img = doc.add_paragraph()
    _format_picture_paragraph(p_img)
    run = p_img.add_run()
    run.add_picture(chart_stream, width=Inches(7.1))

    p_foot = doc.add_paragraph()
    _format_figure_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Dispensações informadas no SAV e NF-e de aquisição de medicamentos.',
        color='64748B',
        size=8,
    )
def _add_figura_percentil_risco(doc, razao_social: str, cnpj_fmt: str, percentil_comp: dict[str, Any], figure_number: int = 1):
    """Insere figura de percentil de risco no documento."""
    p_title = doc.add_paragraph()
    _format_figure_title(p_title)
    _run(
        p_title,
        f'Figura {figure_number:02d} - Posição percentílica da Farmácia {razao_social} (CNPJ {cnpj_fmt}) quanto ao percentual de vendas sem comprovação na Região de Saúde.',
        color='0F172A',
        size=9,
        bold=True,
    )

    chart_stream = _build_percentil_risco_chart_prefer_svg(percentil_comp)
    p_img = doc.add_paragraph()
    _format_picture_paragraph(p_img)
    run = p_img.add_run()
    run.add_picture(chart_stream, width=Inches(7.1))

    p_foot = doc.add_paragraph()
    _format_figure_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Dispensações informadas no SAV e NF-e de aquisição de medicamentos.',
        color='64748B',
        size=8,
    )
def _add_figura_evolucao_financeira(doc, razao_social: str, cnpj_fmt: str, evolucao_comp: dict[str, Any], figure_number: int = 1):
    """Insere figura da evolucao financeira no documento."""
    p_title = doc.add_paragraph()
    _format_figure_title(p_title)
    _run(
        p_title,
        f'Figura {figure_number:02d} - Evolução semestral dos recursos recebidos e das "vendas sem comprovação" da Farmácia {razao_social} (CNPJ {cnpj_fmt}).',
        color='0F172A',
        size=9,
        bold=True,
    )

    chart_stream = _build_evolucao_financeira_chart_prefer_svg(evolucao_comp)
    p_img = doc.add_paragraph()
    _format_picture_paragraph(p_img)
    run = p_img.add_run()
    run.add_picture(chart_stream, width=Inches(7.1))

    p_foot = doc.add_paragraph()
    _format_figure_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Dispensações informadas no SAV e NF-e de aquisição de medicamentos.',
        color='64748B',
        size=8,
    )
