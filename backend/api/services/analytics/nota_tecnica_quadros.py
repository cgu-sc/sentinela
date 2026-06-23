from datetime import date
from decimal import Decimal
from typing import Any

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .nota_tecnica_docx_utils import (
    _add_bookmark,
    _cell_bg,
    _format_block_footnote,
    _format_block_title,
    _keep_small_table_together,
    _repeat_table_header,
    _run,
    _set_table_fixed_widths,
    _set_table_open_borders,
)
from .indicator_rules import get_volume_atipico_aumento_minimo
from .nota_tecnica_formatters import (
    _format_cpf_cnpj,
    _format_decimal_pt,
    _format_patologia_pt,
    _title_case_pt,
)

_CNAES_FARMACEUTICOS = {"4771701", "4771702"}


def _format_quadro_title(paragraph, *, space_before: float = 16, space_after: float = 8):
    _format_block_title(
        paragraph,
        space_before=space_before,
        space_after=space_after,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _format_quadro_footnote(paragraph, *, space_before: float = 5, space_after: float = 18):
    _format_block_footnote(
        paragraph,
        space_before=space_before,
        space_after=space_after,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _format_cnae_secundario_destacado(cnaes_secundarios: Any) -> str:
    if cnaes_secundarios is None:
        raise RuntimeError("Campo cnaes_secundarios ausente no contexto do Quadro 01.")
    if not isinstance(cnaes_secundarios, list):
        raise RuntimeError("Campo cnaes_secundarios deve ser uma lista no contexto do Quadro 01.")
    if not cnaes_secundarios:
        return "—"

    def normalize_item(item: Any) -> tuple[int, str, str | None]:
        if not isinstance(item, dict):
            raise RuntimeError("Item de cnaes_secundarios deve ser um objeto.")
        raw_id = item.get("id_cnae")
        if raw_id is None:
            raise RuntimeError("Item de cnaes_secundarios sem id_cnae.")
        code = "".join(char for char in str(raw_id) if char.isdigit())
        if not code:
            raise RuntimeError("Item de cnaes_secundarios com id_cnae invalido.")
        description = item.get("descricao")
        if description is not None and not isinstance(description, str):
            raise RuntimeError("Item de cnaes_secundarios com descricao invalida.")
        return int(code), code, description.strip() if description else None

    normalized = [normalize_item(item) for item in cnaes_secundarios]
    selected = min(
        (
            item
            for item in normalized
            if item[1] in _CNAES_FARMACEUTICOS
        ),
        default=None,
        key=lambda item: item[0],
    )
    if selected is None:
        selected = min(normalized, key=lambda item: item[0])

    _, code, description = selected
    return f"{code} - {description}" if description else code


def _add_quadro_socios_volume_atipico(doc, socios_volume_atipico: list[dict[str, Any]], tabela_num: int):
    """Adiciona tabela com ingressos societarios proximos a aumentos atipicos."""
    if not socios_volume_atipico:
        return

    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Ingressos societários próximos a semestres com aumento atípico das transferências',
        color='334155',
        size=12,
        bold=True,
    )

    table = doc.add_table(rows=len(socios_volume_atipico) + 1, cols=5)
    _set_table_open_borders(table)
    _set_table_fixed_widths(table, [Inches(1.98), Inches(1.26), Inches(1.26), Inches(1.02), Inches(1.78)])

    headers = [
        'Sócio',
        'Entrada na sociedade',
        'Semestre do aumento',
        'Crescimento',
        'Distância temporal',
    ]
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=10, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    for row_idx, item in enumerate(socios_volume_atipico, start=1):
        cells = table.rows[row_idx].cells
        distancia = int(item.get("distancia_semestres") or 0)
        if distancia == 0:
            distancia_txt = 'Mesmo semestre'
        elif distancia == 1:
            distancia_txt = '1 semestre após a entrada'
        else:
            distancia_txt = f'{distancia} semestres após a entrada'

        taxa = item.get("taxa_crescimento_pct")
        taxa_txt = f'+{_format_decimal_pt(float(taxa), 1)}%' if taxa is not None else '-'
        entrada_txt = item.get("entrada_txt") or '—'
        if entrada_txt != '—':
            entrada_txt = entrada_txt[:1].upper() + entrada_txt[1:]

        values = [
            item.get("nome_socio") or '—',
            entrada_txt,
            item.get("semestre_fmt") or '—',
            taxa_txt,
            distancia_txt,
        ]
        for col_idx, value in enumerate(values):
            para = cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (1, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            _run(para, value, color='0F172A', size=10, bold=col_idx == 3 and taxa is not None)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir do quadro societário cadastral e da evolução semestral das transferências do PFPB. A distância temporal considera o mesmo semestre de entrada e até dois semestres posteriores.',
        color='64748B',
        size=10,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_quadro_comparativo_regional(doc, regional_comp: dict[str, Any], cnpj_data: dict, periodo_txt: str, tabela_num: int):
    """Adiciona tabela compacta comparando a farmacia auditada com a mediana regional."""
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Comparativo do percentual de vendas sem comprovação da farmácia auditada em relação à Região de Saúde',
        color='334155',
        size=12,
        bold=True,
    )

    table = doc.add_table(rows=5, cols=2)
    _set_table_open_borders(table)
    _set_table_fixed_widths(table, [Inches(4.86), Inches(2.44)])

    hdr_cells = table.rows[0].cells
    _run(hdr_cells[0].paragraphs[0], 'Métrica', color='0F172A', size=10, bold=True)
    _run(hdr_cells[1].paragraphs[0], 'Valor', color='0F172A', size=10, bold=True)
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_bg(cell, 'E2E8F0')

    rows = [
        ('Percentual da farmácia auditada', f'{_format_decimal_pt(float(cnpj_data["percValSemComp"]), 2)}%'),
        ('Mediana regional', f'{_format_decimal_pt(regional_comp["mediana_regional"], 2)}%'),
        ('Multiplicador sobre a mediana regional', f'{_format_decimal_pt(regional_comp["multiplicador"], 2)} vezes'),
        ('Farmácias consideradas na região', f'{regional_comp["qtd_farmacias"]}'),
    ]

    for idx, (label, value) in enumerate(rows, start=1):
        cells = table.rows[idx].cells
        _run(cells[0].paragraphs[0], label, color='475569', size=10)
        value_para = cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(value_para, value, color='0F172A', size=10, bold=True)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(p_foot, f'Fonte: Sentinela, com base no SAV/PFPB e em NF-e, no período analisado ({periodo_txt}).', color='64748B', size=10)
    _keep_small_table_together(p_title, table, [p_foot])


def _add_tabela_gtins_sem_comprovacao(doc, razao_social: str, cnpj_fmt: str, gtin_comp: dict[str, Any], periodo_txt: str, tabela_num: int):
    """Adiciona tabela com todos os GTINs com vendas sem comprovacao no periodo."""
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Relação de medicamentos supostamente distribuídos pela Farmácia {razao_social} (CNPJ {cnpj_fmt}), sem estoque amparado em notas fiscais de aquisição, no período de {periodo_txt}.',
        color='334155',
        size=12,
        bold=True,
    )

    rows_data = gtin_comp["rows"]
    table = doc.add_table(rows=len(rows_data) + 2, cols=6)
    _set_table_fixed_widths(
        table,
        [Inches(0.84), Inches(2.42), Inches(0.98), Inches(0.98), Inches(1.18), Inches(0.90)],
    )
    _set_table_open_borders(table)
    _repeat_table_header(table.rows[0])

    headers = [
        'GTIN',
        'Descrição',
        'Patologia associada',
        'Qtd de medicamentos sem comprovação',
        'Valor de vendas sem comprovação (R$)',
        'Percentual sem comprovação',
    ]
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=9, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    for row_idx, item in enumerate(rows_data, start=1):
        cells = table.rows[row_idx].cells
        _run(cells[0].paragraphs[0], item["gtin"], color='0F172A', size=9)
        _run(
            cells[1].paragraphs[0],
            _title_case_pt(item["descricao"]),
            color='0F172A',
            size=9,
        )
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(
            cells[2].paragraphs[0],
            _format_patologia_pt(item["patologia"]),
            color='0F172A',
            size=9,
        )
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[3].paragraphs[0], f'{item["qtd_sem_comprovacao"]:,}'.replace(',', '.'), color='0F172A', size=9)
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(cells[4].paragraphs[0], f'R$ {_format_decimal_pt(item["valor_sem_comprovacao"], 2)}', color='0F172A', size=9)
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[5].paragraphs[0], f'{_format_decimal_pt(item["pct_sem_comprovacao"], 2)}%', color='0F172A', size=9)

    total_cells = table.rows[-1].cells
    total_cells[0].merge(total_cells[2])
    _run(total_cells[0].paragraphs[0], 'TOTAIS', color='0F172A', size=9, bold=True)
    total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(total_cells[3].paragraphs[0], f'{gtin_comp["total_qtd"]:,}'.replace(',', '.'), color='0F172A', size=9, bold=True)
    total_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(total_cells[4].paragraphs[0], f'R$ {_format_decimal_pt(gtin_comp["total_valor"], 2)}', color='0F172A', size=9, bold=True)
    total_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(total_cells[5].paragraphs[0], f'{_format_decimal_pt(gtin_comp["total_pct_sem_comprovacao"], 2)}%', color='0F172A', size=9, bold=True)
    for cell in total_cells:
        _cell_bg(cell, 'F8FAFC')

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(p_foot, f'Fonte: informações sobre as dispensações informadas mensalmente pelas farmácias no Sistema Autorizador de Vendas do PFPB, no período de {periodo_txt}.', color='64748B', size=10)


def _add_quadro_evolucao_financeira(
    doc,
    razao_social: str,
    cnpj_fmt: str,
    evolucao_comp: dict[str, Any],
    tabela_num: int,
):
    """Adiciona tabela semestral de transferencias e vendas sem comprovacao."""
    periodo_semestres = (
        f'no {evolucao_comp["primeiro_semestre_fmt"]}'
        if evolucao_comp["primeiro_semestre"] == evolucao_comp["ultimo_semestre"]
        else f'do {evolucao_comp["primeiro_semestre_fmt"]} ao {evolucao_comp["ultimo_semestre_fmt"]}'
    )
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Evolução semestral do faturamento junto ao Ministério da Saúde e das “vendas sem comprovação” da Farmácia {razao_social} (CNPJ {cnpj_fmt}), {periodo_semestres}.',
        color='334155',
        size=12,
        bold=True,
    )
    _add_bookmark(p_title, "tabela_evolucao_financeira")

    rows_data = evolucao_comp["rows"]
    table = doc.add_table(rows=len(rows_data) + 2, cols=6)
    _set_table_fixed_widths(table, [Inches(1.06), Inches(1.22), Inches(1.22), Inches(1.32), Inches(1.07), Inches(1.41)])
    _set_table_open_borders(table)
    _repeat_table_header(table.rows[0])

    headers = [
        'Semestre',
        'Total',
        'Vendas regulares',
        'Vendas sem comprovação',
        '% sem comprovação',
        'Aumento atípico do total de vendas',
    ]
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=9, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    limite_aumento_volume_atipico = get_volume_atipico_aumento_minimo()
    for row_idx, item in enumerate(rows_data, start=1):
        cells = table.rows[row_idx].cells
        has_aumento_atipico = item.get("volume_atipico") and item.get("taxa_crescimento_pct") is not None
        has_sem_comprovacao_relevante = item["irregular"] >= limite_aumento_volume_atipico
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[0].paragraphs[0], item.get("semestre_fmt") or item["semestre"], color='0F172A', size=9)
        for col_idx, key in enumerate(("total", "regular", "irregular"), start=1):
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(cells[col_idx].paragraphs[0], f'R$ {_format_decimal_pt(item[key], 2)}', color='0F172A', size=9)
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        pct_color = '334155' if item["pct_irregular"] >= 50 else '0F172A'
        _run(
            cells[4].paragraphs[0],
            f'{_format_decimal_pt(item["pct_irregular"], 2)}%',
            color=pct_color,
            size=9,
            bold=item["pct_irregular"] >= 50,
        )
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if has_aumento_atipico:
            _run(
                cells[5].paragraphs[0],
                f'+{_format_decimal_pt(item["taxa_crescimento_pct"], 2)}%',
                color='991B1B',
                size=9,
                bold=True,
            )
        else:
            _run(cells[5].paragraphs[0], '-', color='64748B', size=9)
        if has_aumento_atipico or has_sem_comprovacao_relevante:
            for cell in cells:
                _cell_bg(cell, 'FEE2E2')

    total_cells = table.rows[-1].cells
    _run(total_cells[0].paragraphs[0], 'TOTAL', color='0F172A', size=9, bold=True)
    total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    totals = [
        f'R$ {_format_decimal_pt(evolucao_comp["total"], 2)}',
        f'R$ {_format_decimal_pt(evolucao_comp["regular"], 2)}',
        f'R$ {_format_decimal_pt(evolucao_comp["irregular"], 2)}',
        f'{_format_decimal_pt(evolucao_comp["pct_irregular"], 2)}%',
        '-',
    ]
    for col_idx, value in enumerate(totals, start=1):
        total_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx < 4 else WD_ALIGN_PARAGRAPH.CENTER
        _run(total_cells[col_idx].paragraphs[0], value, color='0F172A', size=9, bold=True)
    for cell in total_cells:
        _cell_bg(cell, 'F8FAFC')

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Sistema Sentinela, com base nas dispensações registradas no SAV/PFPB e nas notas fiscais eletrônicas de aquisição de medicamentos.',
        color='64748B',
        size=10,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_tabela_medicamentos_aumento_atipico(doc, medicamentos_aumento_atipico: list[dict[str, Any]], tabela_num: int):
    """Adiciona tabela com medicamentos que mais contribuiram para semestres atipicos."""
    if not medicamentos_aumento_atipico:
        return

    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Medicamentos associados aos semestres com aumento atípico de volume financeiro',
        color='334155',
        size=12,
        bold=True,
    )

    table = doc.add_table(rows=len(medicamentos_aumento_atipico) + 1, cols=6)
    _set_table_fixed_widths(table, [Inches(1.08), Inches(0.97), Inches(2.42), Inches(0.80), Inches(0.98), Inches(1.05)])
    _set_table_open_borders(table)
    _repeat_table_header(table.rows[0])

    headers = [
        'Semestre',
        'GTIN',
        'Medicamento',
        'Valor anterior',
        'Valor no semestre',
        'Aumento no semestre',
    ]
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=9, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    semestre_block_index: dict[str, int] = {}
    next_block_index = 0

    for row_idx, item in enumerate(medicamentos_aumento_atipico, start=1):
        cells = table.rows[row_idx].cells
        semestre_key = str(item.get("semestre_fmt") or item.get("semestre") or "")
        if semestre_key not in semestre_block_index:
            semestre_block_index[semestre_key] = next_block_index
            next_block_index += 1
        semestre_bg = 'E2E8F0' if semestre_block_index[semestre_key] % 2 == 0 else 'F8FAFC'
        text_color = '0F172A'
        increase_color = '334155'
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[0].paragraphs[0], item["semestre_fmt"], color='334155', size=9, bold=True)
        _cell_bg(cells[0], semestre_bg)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[1].paragraphs[0], item["gtin"], color=text_color, size=9)
        _run(cells[2].paragraphs[0], item["descricao"], color=text_color, size=9)

        for col_idx, key in enumerate(("valor_anterior", "valor_atual"), start=3):
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(cells[col_idx].paragraphs[0], f'R$ {_format_decimal_pt(item[key], 2)}', color=text_color, size=9)

        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        aumento_txt = f'R$ {_format_decimal_pt(item["aumento_valor"], 2)}'
        if item.get("aumento_relativo_pct") is not None:
            aumento_txt += f'\n(+{_format_decimal_pt(item["aumento_relativo_pct"], 1)}%)'
        _run(
            cells[5].paragraphs[0],
            aumento_txt,
            color=increase_color,
            size=9,
            bold=True,
        )

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(
        p_foot,
        'Fonte: Sentinela, a partir da evolução mensal por GTIN. O percentual indicado no campo “Aumento no semestre” representa o crescimento do GTIN em relação ao semestre anterior e é exibido apenas quando o valor anterior é superior a R$ 1,00. GTINs com participação individual menor ou igual a 0,1% no aumento positivo do semestre são consolidados em uma linha agregada.',
        color='64748B',
        size=10,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_quadro_identificacao(doc, data: dict, capital_social: Decimal, periodo_txt: str):
    """Adiciona o Quadro 01 com as informações detalhadas da farmácia."""
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(p_title, f"Quadro 01 - Informações detalhadas da Farmácia {data.get('razao_social') or ''}", color='334155', size=12, bold=True)
    _run(p_title, f"\n(CNPJ {data.get('cnpj_fmt') or ''})", color='475569', size=12)

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    
    # Configura larguras (Total 7.3 inches)
    col_label_w = Inches(2.20)
    col_value_w = Inches(5.10)
    _set_table_fixed_widths(tbl, [col_label_w, col_value_w])

    data_abertura = data.get('data_abertura')
    abertura_txt = data_abertura.strftime('%d/%m/%Y') if data_abertura is not None else '—'

    rows_to_add = [
        ('CNPJ', data.get('cnpj_fmt')),
        ('Razão Social', data.get('razao_social')),
        ('Nome Fantasia', data.get('nome_fantasia') or '—'),
        ('Natureza Jurídica', data.get('natureza_juridica') or '—'),
        ('CNAE Principal', f"{data.get('id_cnae_principal') or ''} - {data.get('cnae_principal') or ''}"),
        ('CNAE Secundário', _format_cnae_secundario_destacado(data.get('cnaes_secundarios'))),
        ('Abertura', abertura_txt),
        ('Situação', data.get('situacao_rf') or '—'),
        ('Porte', data.get('porte_empresa') or '—'),
        ('Capital Social *', f"R$ {capital_social:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')),
        ('Endereço', data.get('endereco_completo') or '—'),
        ('Bairro', data.get('bairro') or '—'),
        ('Município/UF', f"{data.get('municipio') or '—'} / {data.get('uf') or '—'}"),
        ('CEP', data.get('cep') or '—'),
        ('Telefone(s)', ' / '.join(filter(None, [data.get('telefone_1'), data.get('telefone_2')])) or '—'),
        ('E-mail', data.get('email') or '—'),
    ]

    for label, value in rows_to_add:
        row = tbl.add_row()
        row.cells[0].width = col_label_w
        row.cells[1].width = col_value_w
        
        # Estilo Rótulo
        c0 = row.cells[0]
        _cell_bg(c0, 'F8FAFC')
        p0 = c0.paragraphs[0]
        _run(p0, label, color='475569', size=10, bold=True)
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)

        # Estilo Valor
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        _run(p1, str(value) if value else '—', color='0F172A', size=10)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)

    # Nota de Rodapé do Quadro
    total_mov = Decimal(str(data.get('total_mov') or 0.0))
    cap_social_dec = Decimal(str(capital_social))
    relacao_pct = (total_mov / cap_social_dec * 100) if capital_social > 0 else 0
    vezes = (total_mov / cap_social_dec) if capital_social > 0 else 0
    
    p_nota = doc.add_paragraph()
    p_nota.paragraph_format.space_before = Pt(6)
    p_nota.paragraph_format.space_after = Pt(2)
    relacao_txt = f"{relacao_pct:,.2f}%".replace(',', 'X').replace('.', ',').replace('X', '.')
    total_mov_txt = f"R$ {total_mov:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    _run(p_nota, f"* A relação do valor de vendas no âmbito do PFPB sobre o capital social é de {relacao_txt}, ou seja, ela recebeu ", color='475569', size=8)
    _run(p_nota, total_mov_txt, color='0F172A', size=8, bold=True)
    _run(p_nota, f" do Programa, no período de {periodo_txt}, o que corresponde ", color='475569', size=8)
    _run(p_nota, f"{vezes:,.1f} vezes".replace(',', 'X').replace('.', ',').replace('X', '.'), color='0F172A', size=8, bold=True)
    _run(p_nota, " o valor do seu capital social.", color='475569', size=8)

    p_fonte = doc.add_paragraph()
    _format_quadro_footnote(p_fonte)
    dt_extracao = data.get('data_processamento')
    dt_extracao_txt = dt_extracao.strftime('%d/%m/%Y') if dt_extracao else date.today().strftime('%d/%m/%Y')
    _run(p_fonte, f"Fonte: Dados registrados no Cadastro Nacional de Pessoas Jurídicas da RFB, com atualização em {dt_extracao_txt}.", color='94A3B8', size=10, italic=True)

    # ── Quadro Societário Atual ──────────────────────────────────────────
    p_socio_intro = doc.add_paragraph()
    p_socio_intro.paragraph_format.space_before = Pt(12)
    _run(p_socio_intro, f"O quadro societário atual da Farmácia {data.get('razao_social') or ''} conta com os seguintes sócios:", color='0F172A', size=10)

    if data.get('socios_ativos'):
        for s in data['socios_ativos']:
            p_s = doc.add_paragraph(style='List Bullet')
            p_s.paragraph_format.left_indent = Inches(0.5)
            cpf_fmt = _format_cpf_cnpj(s.cpf_cnpj_socio)
            entrada_fmt = s.data_entrada_sociedade.strftime('%d/%m/%Y') if s.data_entrada_sociedade else '—'
            _run(p_s, f"{s.nome_socio}, CPF: {cpf_fmt} (entrada em {entrada_fmt})", color='0F172A', size=10)
    else:
        p_s = doc.add_paragraph(style='List Bullet')
        _run(p_s, "Informação de sócios não disponível ou nenhum sócio ativo identificado.", color='475569', size=10, italic=True)

    p_socio_alerta = doc.add_paragraph()
    p_socio_alerta.paragraph_format.space_before = Pt(6)
    p_socio_alerta.paragraph_format.space_after = Pt(2)
    _run(p_socio_alerta, 'ATENÇÃO: ', color='DC2626', size=10, bold=True, italic=True)
    _run(
        p_socio_alerta,
        'Caso sejam identificadas criticidades em relação ao estabelecimento (sócio com características de “laranja”, endereço inexistente, alteração de endereço para outro município, sócio com vínculo empregatício, etc.), sugere-se trazer tais apontamentos na sequência do parágrafo anterior.',
        color='DC2626',
        size=10,
        italic=True,
    )

    # Contexto trabalhista/eSocial é renderizado fora do quadro cadastral.


def _add_quadro_esocial(doc, razao_social: str, cnpj_fmt: str, esocial_comp: dict[str, Any], tabela_num: int):
    """Adiciona quadro de vínculos trabalhistas anuais do eSocial."""
    rows_data = esocial_comp.get("rows") or []
    if not rows_data:
        return

    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} - Vínculos trabalhistas identificados durante o ano no eSocial para a Farmácia {razao_social} (CNPJ {cnpj_fmt})',
        color='334155',
        size=12,
        bold=True,
    )

    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    _set_table_open_borders(table)
    _set_table_fixed_widths(table, [Inches(1.20), Inches(3.05), Inches(3.05)])

    headers = [
        'Ano',
        'Farmacêuticos com vínculo',
        'Outros trabalhadores com vínculo',
    ]
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=10, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    for row_idx, item in enumerate(rows_data, start=1):
        cells = table.rows[row_idx].cells
        qtd_trabalhadores_vinculo = int(item.get("qtd_trabalhadores_vinculo_ano") or 0)
        qtd_farmaceuticos_vinculo = int(item.get("qtd_farmaceuticos_vinculo_ano") or 0)
        qtd_outros_vinculo = max(qtd_trabalhadores_vinculo - qtd_farmaceuticos_vinculo, 0)
        values = [
            str(item.get("ano_base") or "—"),
            str(qtd_farmaceuticos_vinculo),
            str(qtd_outros_vinculo),
        ]
        for col_idx, value in enumerate(values):
            para = cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(para, value, color='0F172A', size=10)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(
        p_foot,
        f'Fonte: Sentinela, a partir de dados do eSocial. Data de carga mais recente da base eSocial disponível: {esocial_comp.get("dt_carga_fonte_txt") or "—"}.',
        color='64748B',
        size=10,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_quadro_esocial_trabalhadores(doc, razao_social: str, cnpj_fmt: str, esocial_comp: dict[str, Any]):
    """Adiciona quadro detalhado de trabalhadores e CBOs identificados no eSocial."""
    rows_data = esocial_comp.get("trabalhador_detalhe_rows") or []
    if not rows_data:
        return

    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Quadro 01-B - Vínculos trabalhistas identificados no eSocial para a Farmácia {razao_social} (CNPJ {cnpj_fmt})',
        color='334155',
        size=12,
        bold=True,
    )

    table = doc.add_table(rows=len(rows_data) + 1, cols=6)
    table.style = 'Table Grid'
    _set_table_fixed_widths(table, [Inches(0.56), Inches(1.25), Inches(0.76), Inches(2.72), Inches(1.00), Inches(1.01)])

    headers = ['Ano', 'CPF', 'CBO', 'Título CBO', 'Admissão', 'Rescisão']
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=10, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    for row_idx, item in enumerate(rows_data, start=1):
        cells = table.rows[row_idx].cells
        cbo = item.get("cbo")
        values = [
            str(item.get("ano_base") or "—"),
            item.get("cpf_trabalhador") or "—",
            f"{int(cbo):06d}" if cbo is not None else "—",
            item.get("titulo_cbo") or "—",
            item.get("dt_admissao_txt") or "—",
            item.get("dt_rescisao_txt") or "—",
        ]
        for col_idx, value in enumerate(values):
            para = cells[col_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 3 else WD_ALIGN_PARAGRAPH.CENTER
            _run(para, value, color='0F172A', size=10)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    detalhe_modo = esocial_comp.get("trabalhador_detalhe_modo")
    if detalhe_modo == "anos_criticos":
        escopo = (
            "Como o período analisado possui mais de 15 trabalhadores distintos, "
            "o quadro apresenta apenas os vínculos dos anos com condição crítica identificada."
        )
    else:
        escopo = "O quadro apresenta a lista completa de vínculos trabalhistas identificados no período analisado."
    _run(
        p_foot,
        f'Fonte: Sentinela, a partir de dados do eSocial. {escopo}',
        color='64748B',
        size=10,
    )
    _keep_small_table_together(p_title, table, [p_foot])


def _add_quadro_53(doc, razao_social, cnpj_fmt, cnpj_data, periodo_txt, tabela_num: int):
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(p_title, f'Tabela {tabela_num} – Dispensações de medicamentos informadas no Sistema Autorizador de Vendas (SAV) pela Farmácia {razao_social} (CNPJ {cnpj_fmt}), sem comprovação em notas fiscais de aquisição (período de {periodo_txt}).', color='334155', size=12, bold=True)
    
    table = doc.add_table(rows=4, cols=3)
    _set_table_open_borders(table)
    _set_table_fixed_widths(table, [Inches(3.60), Inches(1.85), Inches(1.85)])
    
    hdr_cells = table.rows[0].cells
    _run(hdr_cells[0].paragraphs[0], 'Situação', color='0F172A', size=10, bold=True)
    _run(hdr_cells[1].paragraphs[0], 'Valor em R$', color='0F172A', size=10, bold=True)
    _run(hdr_cells[2].paragraphs[0], 'Quantidades de Medicamentos', color='0F172A', size=10, bold=True)
    
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_bg(cell, 'E2E8F0')

    r1 = table.rows[1].cells
    _run(r1[0].paragraphs[0], 'Dispensações totais informadas no SAV pela farmácia', color='475569', size=10)
    _run(r1[1].paragraphs[0], f'{cnpj_data.get("totalMov", 0):,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'), color='0F172A', size=10, bold=True)
    _run(r1[2].paragraphs[0], f'{cnpj_data.get("totalQtde", 0):,.0f}'.replace(',', '.'), color='0F172A', size=10, bold=True)
    
    r2 = table.rows[2].cells
    _run(r2[0].paragraphs[0], 'Valor de dispensações sem comprovação (R$)', color='475569', size=10)
    _run(r2[1].paragraphs[0], f'{cnpj_data.get("valSemComp", 0):,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.'), color='0F172A', size=10, bold=True)
    _run(r2[2].paragraphs[0], f'{cnpj_data.get("qtdeSemComp", 0):,.0f}'.replace(',', '.'), color='0F172A', size=10, bold=True)

    r3 = table.rows[3].cells
    _run(r3[0].paragraphs[0], '% de vendas no Programa Farmácia Popular sem comprovação', color='475569', size=10)
    _run(r3[1].paragraphs[0], f'{cnpj_data.get("percValSemComp", 0):.2f}%'.replace('.', ','), color='0F172A', size=10, bold=True)
    _run(r3[2].paragraphs[0], f'{cnpj_data.get("percQtdeSemComp", 0):.2f}%'.replace('.', ','), color='0F172A', size=10, bold=True)
    
    for row in table.rows[1:]:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(p_foot, 'Fonte: Relatório de Autorizações Consolidadas, emitido pelo Departamento de Assistência Farmacêutica - DAF/SCTICS/MS, e base de dados das notas fiscais eletrônicas (NF-e), mantida pela Receita Federal do Brasil.', color='64748B', size=10)
    _keep_small_table_together(p_title, table, [p_foot])


def _add_tabela_repasses_anuais(
    doc,
    razao_social: str,
    cnpj_fmt: str,
    repasses_ctx: dict[str, Any],
    tabela_num: int,
) -> None:
    """Tabela de ordens bancárias do SIAFI agregadas por ano para o CNPJ."""
    rows_data = repasses_ctx["rows"]
    periodo_fmt = repasses_ctx["periodo_fmt"]

    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        f'Tabela {tabela_num} – Valores consolidados de ordens bancárias recebidas pela Farmácia {razao_social} (CNPJ {cnpj_fmt}), {periodo_fmt}.',
        color='334155',
        size=12,
        bold=True,
    )

    # cabeçalho + linhas de dados + linha de total
    table = doc.add_table(rows=len(rows_data) + 2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_widths(table, [Inches(2.80), Inches(1.10), Inches(2.40)])
    _set_table_open_borders(table)
    _repeat_table_header(table.rows[0])

    headers = ['Órgão Repassador do Recurso', 'Ano do Repasse', 'Valor total das Ordens Bancárias recebidas']
    for idx, header in enumerate(headers):
        para = table.rows[0].cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=9, bold=True)
        _cell_bg(table.rows[0].cells[idx], 'E2E8F0')

    for row_idx, item in enumerate(rows_data, start=1):
        cells = table.rows[row_idx].cells
        _run(cells[0].paragraphs[0], '36000 - MINISTÉRIO DA SAÚDE', color='0F172A', size=9)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cells[1].paragraphs[0], str(item["ano"]), color='0F172A', size=9)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(cells[2].paragraphs[0], f'R$ {_format_decimal_pt(item["valor"], 2)}', color='0F172A', size=9)

    total_cells = table.rows[-1].cells
    total_cells[0].merge(total_cells[1])
    total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(total_cells[0].paragraphs[0], 'Total', color='0F172A', size=9, bold=True)
    total_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(total_cells[2].paragraphs[0], f'R$ {_format_decimal_pt(repasses_ctx["total"], 2)}', color='0F172A', size=9, bold=True)
    for cell in total_cells:
        _cell_bg(cell, 'F8FAFC')

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    p_foot = doc.add_paragraph()
    _format_quadro_footnote(p_foot)
    _run(p_foot, 'Fonte: Sistema Integrado de Administração Financeira do Governo Federal (SIAFI).', color='64748B', size=10)
    _keep_small_table_together(p_title, table, [p_foot])

    p_atencao = doc.add_paragraph()
    p_atencao.paragraph_format.space_before = Pt(8)
    p_atencao.paragraph_format.space_after = Pt(4)
    _run(p_atencao, 'ATENÇÃO: ', color='C0392B', size=10, bold=True, italic=True)
    _run(
        p_atencao,
        f'Como o Sistema Sentinela se concentra apenas nos medicamentos previstos no rol do PFPB, o comportamento esperado é de que o valor total identificado pelo Sistema de faturamento da empresa junto ao MS (valor total da segunda coluna da Tabela {tabela_num - 1}) seja inferior ou no máximo igual ao valor total das ordens bancárias recebidas do MS (total do quadro anterior, que inclui pagamentos de medicamentos e também fraldas geriátricas e absorventes higiênicos). Valores totais de ordens bancárias menores do que os faturados indicam a possibilidade de glosa por parte do Ministério da Saúde e, consequentemente, de não efetivação total ou parcial de tentativas de fraude ao Programa. Neste caso, sugere-se alterar o texto padrão da "conclusão e encaminhamento" sugerido no item "8" desta Nota Técnica.',
        color='C0392B',
        size=10,
        italic=True,
    )
