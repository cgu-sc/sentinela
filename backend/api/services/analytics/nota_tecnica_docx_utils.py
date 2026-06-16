from itertools import count

from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


_BOOKMARK_ID_COUNTER = count(1)


def _add_bookmark(paragraph, name: str):
    """Adiciona um bookmark interno do Word ao paragrafo informado."""
    if not name:
        raise RuntimeError("Nome de bookmark obrigatorio para Nota Tecnica.")

    bookmark_id = str(next(_BOOKMARK_ID_COUNTER))
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bookmark_id)
    start.set(qn('w:name'), name)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bookmark_id)

    p = paragraph._p
    insert_idx = 1 if p.pPr is not None else 0
    p.insert(insert_idx, start)
    p.append(end)


def _add_internal_hyperlink(
    paragraph,
    text: str,
    anchor: str,
    *,
    color: str = '2563EB',
    size: float = 7,
    bold: bool = False,
    underline: bool = True,
):
    """Adiciona um hyperlink interno apontando para um bookmark do documento."""
    if not anchor:
        raise RuntimeError("Destino de hyperlink interno obrigatorio para Nota Tecnica.")

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    if bold:
        r_pr.append(OxmlElement('w:b'))

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    r_pr.append(color_el)

    underline_el = OxmlElement('w:u')
    underline_el.set(qn('w:val'), 'single' if underline else 'none')
    r_pr.append(underline_el)

    size_el = OxmlElement('w:sz')
    size_el.set(qn('w:val'), str(int(round(size * 2))))
    r_pr.append(size_el)

    run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_external_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    color: str = '2563EB',
    size: float = 10,
    bold: bool = False,
    underline: bool = True,
):
    """Adiciona um hyperlink externo ao paragrafo informado."""
    if not url:
        raise RuntimeError("URL de hyperlink externo obrigatoria para Nota Tecnica.")

    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship_id)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    if bold:
        r_pr.append(OxmlElement('w:b'))

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    r_pr.append(color_el)

    underline_el = OxmlElement('w:u')
    underline_el.set(qn('w:val'), 'single' if underline else 'none')
    r_pr.append(underline_el)

    size_el = OxmlElement('w:sz')
    size_el.set(qn('w:val'), str(int(round(size * 2))))
    r_pr.append(size_el)

    run.append(r_pr)
    t = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _cell_bg(cell, fill_hex: str):
    """Define cor de fundo de uma celula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _cell_bg_run(run, fill_hex: str):
    """Define cor de fundo para um run especifico."""
    r = run._r
    rPr = r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    rPr.append(shd)


def _cell_borders(cell, left=None, right=None, top=None, bottom=None):
    """Define bordas coloridas em lados especificos da celula. None = sem borda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side, spec in (('left', left), ('right', right), ('top', top), ('bottom', bottom)):
        el = OxmlElement(f'w:{side}')
        if spec:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), spec['sz'])
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), spec['color'])
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _tbl_no_borders(tbl):
    """Remove todas as bordas visiveis de uma tabela no nivel da tabela."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_table_open_borders(tbl, *, color: str = '0F172A', sz: str = '4'):
    """Aplica bordas apenas horizontais (topo, cabecalho, linhas, base), sem bordas laterais/verticais."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    horizontal = {'top', 'bottom', 'insideH'}
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        if side in horizontal:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _row_cant_split(row):
    """Evita que uma linha da tabela seja quebrada entre paginas pelo Word."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)


def _repeat_table_header(row):
    """Marca uma linha de tabela para repetir como cabecalho nas paginas seguintes."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        tbl_header = OxmlElement('w:tblHeader')
        tbl_header.set(qn('w:val'), 'true')
        trPr.append(tbl_header)


def _keep_small_table_together(title_para, table, trailing_paragraphs=None):
    """Mantem titulo e tabela curta juntos quando houver espaco na pagina."""
    title_para.paragraph_format.keep_with_next = True
    title_para.paragraph_format.keep_together = True

    table_paragraphs = []
    for row in table.rows:
        _row_cant_split(row)
        for cell in row.cells:
            table_paragraphs.extend(cell.paragraphs)

    for para in table_paragraphs:
        para.paragraph_format.keep_together = True
        para.paragraph_format.keep_with_next = True

    trailing_paragraphs = trailing_paragraphs or []
    for para in trailing_paragraphs:
        para.paragraph_format.keep_together = True
        para.paragraph_format.keep_with_next = True

    if trailing_paragraphs:
        trailing_paragraphs[-1].paragraph_format.keep_with_next = False
    elif table_paragraphs:
        table_paragraphs[-1].paragraph_format.keep_with_next = False


def _format_block_title(
    paragraph,
    *,
    space_before: float = 16,
    space_after: float = 8,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def _format_block_footnote(
    paragraph,
    *,
    space_before: float = 5,
    space_after: float = 18,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def _format_picture_paragraph(paragraph, *, keep_with_next: bool = True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)


def _rgb(hex6: str) -> RGBColor:
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


def _run(para, text: str, *, color: str = '0F172A', size: float = 10, bold=False, italic=False, underline=False):
    if bold and color in {'0F172A', '475569', '64748B'}:
        color = '334155'
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    return run


def _get_or_add_footnotes_part(doc):
    """Obtem ou cria a parte XML de footnotes reais do Word."""
    try:
        return doc.part.part_related_by(RT.FOOTNOTES)
    except KeyError:
        xml = (
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            '</w:footnotes>'
        )
        part = XmlPart(
            PackURI('/word/footnotes.xml'),
            CT.WML_FOOTNOTES,
            parse_xml(xml),
            doc.part.package,
        )
        doc.part.relate_to(part, RT.FOOTNOTES)
        return part


def _ensure_footnote(doc, number: int, text: str):
    """Garante a existencia de uma nota de rodape numerada."""
    part = _get_or_add_footnotes_part(doc)
    footnotes = part.element
    for existing in footnotes.findall(qn('w:footnote')):
        if existing.get(qn('w:id')) == str(number):
            return

    footnote = OxmlElement('w:footnote')
    footnote.set(qn('w:id'), str(number))
    p = OxmlElement('w:p')
    p_pr = OxmlElement('w:pPr')
    p_style = OxmlElement('w:pStyle')
    p_style.set(qn('w:val'), 'FootnoteText')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(p_style)
    p_pr.append(spacing)
    p.append(p_pr)

    r_ref = OxmlElement('w:r')
    r_ref_pr = OxmlElement('w:rPr')
    r_ref_style = OxmlElement('w:rStyle')
    r_ref_style.set(qn('w:val'), 'FootnoteReference')
    r_ref_vert = OxmlElement('w:vertAlign')
    r_ref_vert.set(qn('w:val'), 'superscript')
    r_ref_size = OxmlElement('w:sz')
    r_ref_size.set(qn('w:val'), '20')
    r_ref_pr.append(r_ref_style)
    r_ref_pr.append(r_ref_vert)
    r_ref_pr.append(r_ref_size)
    r_ref.append(r_ref_pr)
    r_ref_note = OxmlElement('w:footnoteRef')
    r_ref.append(r_ref_note)
    p.append(r_ref)

    r_text = OxmlElement('w:r')
    r_text_pr = OxmlElement('w:rPr')
    r_text_size = OxmlElement('w:sz')
    r_text_size.set(qn('w:val'), '16')
    r_text_pr.append(r_text_size)
    r_text.append(r_text_pr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = f' {text}'
    r_text.append(t)
    p.append(r_text)

    footnote.append(p)
    footnotes.append(footnote)


def _footnote_ref(doc, para, number: int, text: str):
    """Insere referencia de footnote real e garante seu texto."""
    _ensure_footnote(doc, number, text)
    run = para.add_run()
    r = run._r
    r_pr = r.get_or_add_rPr()
    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'FootnoteReference')
    r_pr.append(r_style)
    r_vert = OxmlElement('w:vertAlign')
    r_vert.set(qn('w:val'), 'superscript')
    r_pr.append(r_vert)
    r_size = OxmlElement('w:sz')
    r_size.set(qn('w:val'), '20')
    r_pr.append(r_size)
    ref = OxmlElement('w:footnoteReference')
    ref.set(qn('w:id'), str(number))
    r.append(ref)
    return run


def _set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width.twips)))
    tc_w.set(qn('w:type'), 'dxa')


def _set_table_fixed_widths(table, widths):
    """Fixa a grade da tabela para o Word nao redistribuir as colunas."""
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_layout = tbl_pr.find(qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = OxmlElement('w:tblLayout')
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn('w:type'), 'fixed')

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(int(width.twips) for width in widths)))
    tbl_w.set(qn('w:type'), 'dxa')

    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl_grid = OxmlElement('w:tblGrid')
    for width in widths:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(width.twips)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for idx, width in enumerate(widths):
        table.columns[idx].width = width
        for cell in table.columns[idx].cells:
            _set_cell_width(cell, width)


def _write_cell(cell, text: str, *, size: float = 6.4, bold: bool = False, color: str = '0F172A', align=None):
    if bold and color in {'0F172A', '475569', '64748B'}:
        color = '334155'
    p = cell.paragraphs[0]
    p.text = ''
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    _run(p, text, color=color, size=size, bold=bold)


def _append_table_row_fast(table, values, widths, *, alignments=None, sizes=None, color: str = '0F172A', fill: str | None = None):
    tr = OxmlElement('w:tr')
    alignments = alignments or []
    sizes = sizes or []

    for idx, value in enumerate(values):
        tc = OxmlElement('w:tc')
        tc_pr = OxmlElement('w:tcPr')

        if idx < len(widths):
            tc_w = OxmlElement('w:tcW')
            tc_w.set(qn('w:w'), str(int(widths[idx].twips)))
            tc_w.set(qn('w:type'), 'dxa')
            tc_pr.append(tc_w)

        if fill:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tc_pr.append(shd)

        tc.append(tc_pr)

        p = OxmlElement('w:p')
        p_pr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        p_pr.append(spacing)

        align = alignments[idx] if idx < len(alignments) else None
        if align:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), align)
            p_pr.append(jc)

        p.append(p_pr)

        def append_run(text_value, *, run_color=color, run_size=None, line_break=False):
            r = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')
            color_el = OxmlElement('w:color')
            color_el.set(qn('w:val'), run_color)
            r_pr.append(color_el)
            sz = OxmlElement('w:sz')
            size = run_size if run_size is not None else (sizes[idx] if idx < len(sizes) else 6.4)
            sz.set(qn('w:val'), str(int(round(size * 2))))
            r_pr.append(sz)
            r.append(r_pr)

            if line_break:
                r.append(OxmlElement('w:br'))
            else:
                text_el = OxmlElement('w:t')
                text_str = str(text_value)
                if text_str.startswith(' ') or text_str.endswith(' '):
                    text_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                text_el.text = text_str
                r.append(text_el)
            p.append(r)

        if isinstance(value, list):
            for segment in value:
                if isinstance(segment, dict):
                    append_run(
                        segment.get('text', ''),
                        run_color=segment.get('color', color),
                        run_size=segment.get('size'),
                        line_break=bool(segment.get('break')),
                    )
                else:
                    append_run(segment)
        else:
            append_run(value)

        tc.append(p)
        tr.append(tc)

    table._tbl.append(tr)
