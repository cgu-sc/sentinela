from datetime import datetime
from typing import Any

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .farmacia import get_movimentacao_data
from .nota_tecnica_docx_utils import (
    _append_table_row_fast,
    _cell_bg,
    _format_block_footnote,
    _format_block_title,
    _repeat_table_header,
    _row_cant_split,
    _run,
    _set_cell_width,
    _set_table_fixed_widths,
    _write_cell,
)
from .nota_tecnica_formatters import _format_decimal_pt


def _model_to_dict(model: Any) -> dict[str, Any]:
    if isinstance(model, dict):
        return model
    if hasattr(model, "model_dump"):
        return model.model_dump()
    if hasattr(model, "dict"):
        return model.dict()
    return {}


def _resolve_engine(db_or_engine: Any) -> Any:
    if hasattr(db_or_engine, "connect"):
        return db_or_engine
    if hasattr(db_or_engine, "get_bind"):
        return db_or_engine.get_bind()
    if hasattr(db_or_engine, "bind"):
        return db_or_engine.bind
    return db_or_engine


def _parse_period_date(value: Any) -> tuple[datetime, str] | None:
    if not value:
        return None
    if hasattr(value, "strftime"):
        return datetime(value.year, value.month, value.day), value.strftime("%d/%m/%Y")

    text = str(value).strip()
    if not text or text in {"-", "—"}:
        return None

    for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
        try:
            parsed = datetime.strptime(text[:10], fmt)
            return parsed, parsed.strftime("%d/%m/%Y")
        except ValueError:
            continue
    return None


def _pick_period_boundary(current: tuple[datetime, str] | None, value: Any, direction: str) -> tuple[datetime, str] | None:
    parsed = _parse_period_date(value)
    if parsed is None:
        return current
    if current is None:
        return parsed
    if direction == "min" and parsed[0] < current[0]:
        return parsed
    if direction == "max" and parsed[0] > current[0]:
        return parsed
    return current


def _build_sections(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    for row in rows:
        tipo_linha = row.get("tipo_linha")
        if tipo_linha == "header_medicamento":
            if current:
                sections.append(current)
            current = {
                "gtin": row.get("gtin"),
                "medicamento": "",
                "rows": [],
                "subtotal": None,
            }
            continue

        if current is None:
            continue

        if tipo_linha == "resumo_parcial":
            current["subtotal"] = row
            if not current["medicamento"]:
                current["medicamento"] = row.get("medicamento") or ""
            continue

        if tipo_linha == "header_colunas":
            continue

        if tipo_linha in {"venda_normal", "venda_irregular"} and not current["medicamento"]:
            current["medicamento"] = row.get("medicamento") or ""
        current["rows"].append(row)

    if current:
        sections.append(current)
    return sections


def _build_anexo_ii_context(cnpj: str, db_or_engine: Any) -> dict[str, Any]:
    """Monta a memoria de calculo das vendas sem comprovacao para o Anexo II."""
    movimentacao = get_movimentacao_data(cnpj, _resolve_engine(db_or_engine))
    summary = _model_to_dict(getattr(movimentacao, "summary", None))
    rows = [_model_to_dict(row) for row in (getattr(movimentacao, "rows", None) or [])]
    total_irregular = float(summary.get("valor_irregular") or 0.0)

    consolidados: list[dict[str, Any]] = []
    detalhes: list[dict[str, Any]] = []
    for section in _build_sections(rows):
        subtotal = section.get("subtotal") or {}
        irregular_rows = [
            row
            for row in section.get("rows", [])
            if row.get("tipo_linha") == "venda_irregular" and float(row.get("valor_irregular") or 0.0) > 0
        ]
        valor_irregular = float(subtotal.get("valor_irregular") or 0.0)
        if valor_irregular <= 0 and irregular_rows:
            valor_irregular = sum(float(row.get("valor_irregular") or 0.0) for row in irregular_rows)
        if valor_irregular <= 0:
            continue

        min_start = None
        max_end = None
        for row in irregular_rows:
            min_start = _pick_period_boundary(min_start, row.get("periodo_inicio_irregular"), "min")
            max_end = _pick_period_boundary(max_end, row.get("periodo_final"), "max")

        periodo = f"{min_start[1]} a {max_end[1]}" if min_start and max_end else "—"
        estoque_final = subtotal.get("estoque_final")
        if estoque_final is None:
            for row in reversed(section.get("rows", [])):
                if row.get("estoque_final") is not None:
                    estoque_final = row.get("estoque_final")
                    break

        consolidado = {
            "gtin": str(section.get("gtin") or subtotal.get("gtin") or ""),
            "medicamento": section.get("medicamento") or subtotal.get("medicamento") or "NÃO IDENTIFICADO",
            "periodo_sem_comprovacao": periodo,
            "estoque_final": int(estoque_final or 0),
            "vendas": int(subtotal.get("vendas") or sum(int(row.get("vendas") or 0) for row in section.get("rows", []))),
            "vendas_irregular": int(
                subtotal.get("vendas_irregular") or sum(int(row.get("vendas_irregular") or 0) for row in irregular_rows)
            ),
            "valor": round(float(subtotal.get("valor") or sum(float(row.get("valor") or 0.0) for row in section.get("rows", []))), 2),
            "valor_irregular": round(valor_irregular, 2),
            "pct_prejuizo_total": (valor_irregular / total_irregular * 100) if total_irregular > 0 else 0.0,
        }
        detail_rows = [
            row
            for row in section.get("rows", [])
            if row.get("tipo_linha") in {"venda_normal", "venda_irregular"}
        ]
        consolidados.append(consolidado)
        detalhes.append({**consolidado, "rows": detail_rows})

    consolidados.sort(key=lambda item: item["valor_irregular"], reverse=True)
    detalhes.sort(key=lambda item: item["valor_irregular"], reverse=True)

    return {
        "summary": summary,
        "consolidado": consolidados,
        "detalhes": detalhes,
        "total_gtins_irregulares": len(consolidados),
    }


def _add_table_header(table, headers: list[str], widths: list[Any], *, size: float = 7.0):
    hdr = table.rows[0]
    _repeat_table_header(hdr)
    _row_cant_split(hdr)
    for idx, label in enumerate(headers):
        cell = hdr.cells[idx]
        _set_cell_width(cell, widths[idx])
        _cell_bg(cell, "E2E8F0")
        _write_cell(cell, label, size=size, bold=True, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_anexo_ii_detalhamento(doc, detalhes: list[dict[str, Any]], timing: Any = None):
    if not detalhes:
        return

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(12)
    p_intro.paragraph_format.space_after = Pt(6)
    _run(
        p_intro,
        "Detalhamento da memória de cálculo dos GTINs que apresentaram vendas sem comprovação",
        color="0F172A",
        size=9,
        bold=True,
    )

    headers = [
        "Período inicial",
        "Início da não comprovação",
        "Período final",
        "Estoque inicial",
        "Estoque final",
        "Caixas vendidas",
        "Caixas sem comprovação",
        "Valor vendido",
        "Valor sem comprovação",
        "NF-e consideradas",
    ]
    widths = [
        Inches(0.74),
        Inches(0.82),
        Inches(0.74),
        Inches(0.62),
        Inches(0.62),
        Inches(0.64),
        Inches(0.72),
        Inches(0.78),
        Inches(0.82),
        Inches(3.55),
    ]

    for idx, detalhe in enumerate(detalhes, start=1):
        rows = detalhe.get("rows") or []
        if not rows:
            continue

        p_title = doc.add_paragraph()
        _format_block_title(p_title, space_before=16, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _run(
            p_title,
            f"Quadro II.{idx + 2} - Memória de cálculo do GTIN {detalhe.get('gtin') or ''} - {detalhe.get('medicamento') or 'NÃO IDENTIFICADO'}",
            color="334155",
            size=8,
            bold=True,
        )

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        _set_table_fixed_widths(table, widths)
        _add_table_header(table, headers, widths, size=7.6)

        for item in rows:
            has_irregular = float(item.get("valor_irregular") or 0.0) > 0
            values = [
                item.get("periodo_inicial") or "-",
                item.get("periodo_inicio_irregular") or "-",
                item.get("periodo_final") or "-",
                f'{int(item.get("estoque_inicial") or 0):,}'.replace(",", "."),
                f'{int(item.get("estoque_final") or 0):,}'.replace(",", "."),
                f'{int(item.get("vendas") or 0):,}'.replace(",", "."),
                f'{int(item.get("vendas_irregular") or 0):,}'.replace(",", "."),
                f'R$ {_format_decimal_pt(float(item.get("valor") or 0.0), 2)}',
                f'R$ {_format_decimal_pt(float(item.get("valor_irregular") or 0.0), 2)}',
                item.get("notas") or "-",
            ]
            _append_table_row_fast(
                table,
                values,
                widths,
                alignments=["center"] * 9 + ["left"],
                sizes=[7.3] * 9 + [6.9],
                color="0F172A",
                fill="FEF2F2" if has_irregular else None,
            )

        subtotal_row = table.add_row()
        _row_cant_split(subtotal_row)
        label_cell = subtotal_row.cells[0].merge(subtotal_row.cells[4])
        _cell_bg(label_cell, "F8FAFC")
        _write_cell(label_cell, "Subtotal do GTIN", size=7.7, bold=True, color="475569", align=WD_ALIGN_PARAGRAPH.RIGHT)
        subtotal_values = [
            f'{int(detalhe.get("vendas") or 0):,}'.replace(",", "."),
            f'{int(detalhe.get("vendas_irregular") or 0):,}'.replace(",", "."),
            f'R$ {_format_decimal_pt(float(detalhe.get("valor") or 0.0), 2)}',
            f'R$ {_format_decimal_pt(float(detalhe.get("valor_irregular") or 0.0), 2)}',
            "",
        ]
        for offset, value in enumerate(subtotal_values, start=5):
            cell = subtotal_row.cells[offset]
            _cell_bg(cell, "F8FAFC")
            _write_cell(cell, value, size=7.7, bold=True, color="475569", align=WD_ALIGN_PARAGRAPH.CENTER)
        if timing:
            timing.mark(f"anexo II detalhe GTIN {idx} ({len(rows)} linhas)")


def _add_anexo_ii_memoria_calculo(
    doc,
    razao_social: str,
    cnpj_fmt: str,
    periodo_txt: str,
    anexo_ii_comp: dict[str, Any],
    timing: Any = None,
):
    """Adiciona o Anexo II com resumo, consolidado e detalhamento da memoria de calculo."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    summary = anexo_ii_comp.get("summary") or {}
    consolidados = anexo_ii_comp.get("consolidado") or []
    detalhes = anexo_ii_comp.get("detalhes") or []

    doc.add_heading("ANEXO II - MEMÓRIA DE CÁLCULO DAS VENDAS SEM COMPROVAÇÃO", level=1)
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    _run(
        p_intro,
        f"O presente anexo apresenta a memória de cálculo utilizada para identificação das dispensações de medicamentos sem comprovação de estoque da Farmácia {razao_social} (CNPJ {cnpj_fmt}), no período {periodo_txt}, a partir do confronto entre as vendas informadas no SAV e as notas fiscais eletrônicas de aquisição de medicamentos.",
        color="0F172A",
        size=9,
    )

    p_title = doc.add_paragraph()
    _format_block_title(p_title, space_before=16, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_title, "Quadro II.1 - Resumo geral da memória de cálculo", color="334155", size=8, bold=True)

    summary_headers = [
        "Total de medicamentos dispensados",
        "Medicamentos dispensados sem comprovação",
        "Valor total informado no SAV",
        "Valor sem comprovação",
        "Percentual sem comprovação",
    ]
    summary_widths = [Inches(1.65), Inches(1.85), Inches(2.0), Inches(1.8), Inches(1.8)]
    summary_table = doc.add_table(rows=2, cols=len(summary_headers))
    summary_table.style = "Table Grid"
    summary_table.autofit = False
    _set_table_fixed_widths(summary_table, summary_widths)
    _add_table_header(summary_table, summary_headers, summary_widths, size=7.4)

    summary_values = [
        f'{int(summary.get("total_vendas") or 0):,}'.replace(",", "."),
        f'{int(summary.get("total_vendas_irregular") or 0):,}'.replace(",", "."),
        f'R$ {_format_decimal_pt(float(summary.get("valor_total") or 0.0), 2)}',
        f'R$ {_format_decimal_pt(float(summary.get("valor_irregular") or 0.0), 2)}',
        f'{_format_decimal_pt(float(summary.get("pct_irregular") or 0.0), 2)}%',
    ]
    row = summary_table.rows[1]
    _row_cant_split(row)
    for idx, value in enumerate(summary_values):
        _set_cell_width(row.cells[idx], summary_widths[idx])
        _write_cell(row.cells[idx], value, size=7.5, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    if timing:
        timing.mark("anexo II resumo geral")

    p_title2 = doc.add_paragraph()
    _format_block_title(p_title2, space_before=16, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p_title2, "Quadro II.2 - Medicamentos com vendas sem comprovação, por GTIN", color="334155", size=8, bold=True)

    headers = [
        "GTIN",
        "Medicamento / princípio ativo",
        "Período sem comprovação",
        "Estoque final estimado",
        "Total de caixas vendidas",
        "Caixas sem comprovação",
        "Valor total vendido",
        "Valor sem comprovação",
        "% do prejuízo total",
    ]
    widths = [
        Inches(1.0),
        Inches(2.15),
        Inches(1.25),
        Inches(0.75),
        Inches(0.8),
        Inches(0.85),
        Inches(1.0),
        Inches(1.0),
        Inches(0.75),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _set_table_fixed_widths(table, widths)
    _add_table_header(table, headers, widths)

    if consolidados:
        for item in consolidados:
            row = table.add_row()
            _row_cant_split(row)
            values = [
                item["gtin"],
                item["medicamento"],
                item["periodo_sem_comprovacao"],
                f'{int(item["estoque_final"]):,}'.replace(",", "."),
                f'{int(item["vendas"]):,}'.replace(",", "."),
                f'{int(item["vendas_irregular"]):,}'.replace(",", "."),
                f'R$ {_format_decimal_pt(float(item["valor"]), 2)}',
                f'R$ {_format_decimal_pt(float(item["valor_irregular"]), 2)}',
                f'{_format_decimal_pt(float(item["pct_prejuizo_total"]), 1)}%',
            ]
            for idx, value in enumerate(values):
                cell = row.cells[idx]
                _write_cell(cell, value, size=6.7, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        row = table.add_row()
        _row_cant_split(row)
        merged = row.cells[0].merge(row.cells[-1])
        _write_cell(
            merged,
            "Não foram identificados medicamentos com vendas sem comprovação na memória de cálculo disponível.",
            size=7.2,
            color="475569",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    if timing:
        timing.mark(f"anexo II quadro consolidado ({len(consolidados)} GTINs)")

    p_foot = doc.add_paragraph()
    _format_block_footnote(p_foot, space_before=5, space_after=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _run(
        p_foot,
        "Fonte: Dispensações informadas no SAV e NF-e de aquisição de medicamentos.",
        color="64748B",
        size=8,
    )

    _add_anexo_ii_detalhamento(doc, detalhes, timing=timing)
    if timing:
        total_linhas = sum(len(item.get("rows") or []) for item in detalhes)
        timing.mark(f"anexo II detalhamento total ({len(detalhes)} GTINs, {total_linhas} linhas)")
