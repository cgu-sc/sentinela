from datetime import date
from typing import Any

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .nota_tecnica_docx_utils import (
    _cell_bg,
    _repeat_table_header,
    _row_cant_split,
    _run,
    _set_cell_width,
    _set_table_fixed_widths,
    _write_cell,
    _write_cell_fast,
)
from .nota_tecnica_formatters import (
    _format_cpf_cnpj,
    _format_date_pt,
    _format_decimal_pt,
    _title_case_pt,
)

def _build_falecidos_grupos(transacoes: list[Any]) -> list[dict[str, Any]]:
    """Agrupa transacoes de falecidos por CPF, seguindo o modelo da aba Mortality."""
    grupos: dict[str, dict[str, Any]] = {}
    for t in transacoes:
        cpf = str(getattr(t, "cpf", "") or "").zfill(11)
        if cpf not in grupos:
            grupos[cpf] = {
                "cpf": cpf,
                "nome": _title_case_pt(getattr(t, "nome_falecido", None)),
                "municipio": _title_case_pt(getattr(t, "municipio_falecido", None)) if getattr(t, "municipio_falecido", None) else "—",
                "uf": getattr(t, "uf_falecido", None) or "—",
                "dt_obito": getattr(t, "dt_obito", None),
                "outros_cnpj": getattr(t, "outros_estabelecimentos", None),
                "transacoes": [],
                "total_valor": 0.0,
                "max_dias": 0,
            }

        grupo = grupos[cpf]
        valor = float(getattr(t, "valor_total_autorizacao", 0.0) or 0.0)
        dias = int(getattr(t, "dias_apos_obito", 0) or 0)
        grupo["transacoes"].append(t)
        grupo["total_valor"] += valor
        grupo["max_dias"] = max(grupo["max_dias"], dias)

    def sort_key(item: dict[str, Any]):
        datas_validas = [
            data_autorizacao
            for t in item["transacoes"]
            if isinstance((data_autorizacao := getattr(t, "data_autorizacao", None)), date)
        ]
        primeira_data = min(datas_validas, default=date.max)
        return (item["cpf"], primeira_data)

    grupos_lista = sorted(grupos.values(), key=sort_key)
    for grupo in grupos_lista:
        grupo["transacoes"].sort(
            key=lambda t: (
                getattr(t, "data_autorizacao", None)
                if isinstance(getattr(t, "data_autorizacao", None), date)
                else date.max,
                getattr(t, "num_autorizacao", "") or "",
            )
        )
    return grupos_lista


def _write_person_cell(cell, nome: str, cpf: str):
    p = cell.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, nome, color='0F172A', size=8.2)
    p.add_run('\n')
    _run(p, _format_cpf_cnpj(cpf), color='64748B', size=7.0)


def _add_anexo_iii_falecidos(doc, razao_social: str, cnpj_fmt: str, falecidos_comp: dict[str, Any], timing: Any = None):
    """Adiciona o Anexo III com detalhamento das vendas para pessoas falecidas."""
    transacoes = falecidos_comp.get("transacoes") or []
    if not transacoes:
        return

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ''
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.orientation = WD_ORIENT.PORTRAIT
    if section.page_width > section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width

    grupos = _build_falecidos_grupos(transacoes)
    total_autorizacoes = int(falecidos_comp.get("total_autorizacoes") or len(transacoes))
    cpfs_distintos = int(falecidos_comp.get("cpfs_distintos") or len(grupos))
    valor_total = float(falecidos_comp.get("valor_total") or 0.0)
    if timing:
        timing.mark(f"anexo III agrupamento ({len(grupos)} CPFs, {len(transacoes)} transacoes)")

    doc.add_heading('ANEXO III - DETALHAMENTO DE VENDAS PARA PESSOAS FALECIDAS', level=1)
    p_intro = doc.add_paragraph()
    _run(
        p_intro,
        f'Detalhamento de transações agrupadas por CPF relativas à Farmácia {razao_social} (CNPJ {cnpj_fmt}), '
        f'{falecidos_comp.get("periodo_desc") or "no período analisado"}.',
        color='0F172A',
        size=9,
    )

    headers = [
        'Pessoa falecida',
        'Município/UF',
        'Dt. óbito',
        'Data da venda',
        'Valor (R$)',
        'Dias após óbito',
    ]
    widths = [
        Inches(2.95), Inches(1.48), Inches(1.0), Inches(1.12), Inches(0.6), Inches(0.45),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    _set_table_fixed_widths(table, widths)

    hdr = table.rows[0]
    _repeat_table_header(hdr)
    _row_cant_split(hdr)
    for idx, label in enumerate(headers):
        cell = hdr.cells[idx]
        _set_cell_width(cell, widths[idx])
        _cell_bg(cell, 'E2E8F0')
        _write_cell(cell, label, size=7.3, bold=True, color='0F172A', align=WD_ALIGN_PARAGRAPH.CENTER)
    if timing:
        timing.mark("anexo III cabecalho tabela")

    for grupo_idx, grupo in enumerate(grupos, start=1):
        for t in grupo["transacoes"]:
            row = table.add_row()
            values = [
                None,
                f'{grupo["municipio"]}/{grupo["uf"]}',
                _format_date_pt(getattr(t, "dt_obito", None)),
                _format_date_pt(getattr(t, "data_autorizacao", None)),
                _format_decimal_pt(float(getattr(t, "valor_total_autorizacao", 0.0) or 0.0), 2),
                f'{int(getattr(t, "dias_apos_obito", 0) or 0)} d',
            ]
            for idx, value in enumerate(values):
                cell = row.cells[idx]
                if idx == 0:
                    cpf = str(getattr(t, "cpf", "") or "")
                    _write_person_cell(cell, _title_case_pt(getattr(t, "nome_falecido", None)), cpf)
                    continue
                align = WD_ALIGN_PARAGRAPH.RIGHT if idx in {4, 5} else WD_ALIGN_PARAGRAPH.LEFT
                if idx in {2, 3}:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                _write_cell_fast(cell, value or "—", size=7.4, color='0F172A', align=align)

        subtotal_row = table.add_row()
        _row_cant_split(subtotal_row)
        label_cell = subtotal_row.cells[0].merge(subtotal_row.cells[3])
        _cell_bg(label_cell, 'F8FAFC')
        _write_cell(
            label_cell,
            f'Subtotal - {len(grupo["transacoes"])} autorização(ões)',
            size=7.4,
            bold=True,
            color='475569',
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        value_cell = subtotal_row.cells[4]
        _cell_bg(value_cell, 'F8FAFC')
        _write_cell(value_cell, _format_decimal_pt(grupo["total_valor"], 2), size=7.4, bold=True, color='475569', align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_bg(subtotal_row.cells[5], 'F8FAFC')
        _write_cell(subtotal_row.cells[5], '', size=7.4)
        if timing:
            timing.mark(f"anexo III CPF {grupo_idx} ({len(grupo['transacoes'])} transacoes)")

    total_row = table.add_row()
    _row_cant_split(total_row)
    total_label = total_row.cells[0].merge(total_row.cells[3])
    _cell_bg(total_label, 'E2E8F0')
    _write_cell(
        total_label,
        f'TOTAL GERAL - {cpfs_distintos} CPF(s) distintos - {total_autorizacoes} autorização(ões)',
        size=7.6,
        bold=True,
        color='0F172A',
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    total_value = total_row.cells[4]
    _cell_bg(total_value, 'E2E8F0')
    _write_cell(total_value, _format_decimal_pt(valor_total, 2), size=7.6, bold=True, color='0F172A', align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_bg(total_row.cells[5], 'E2E8F0')
    _write_cell(total_row.cells[5], '', size=7.6)
    if timing:
        timing.mark("anexo III total geral")
