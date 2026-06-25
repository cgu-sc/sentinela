"""
Validação isolada do fix de farmácia sem repasses na Nota Técnica.
Cria docx em memória e valida os 2 cenários: com e sem repasses.
"""
import sys
import os
import io
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
os.environ.setdefault("PYTHONPATH", r"D:\sentinela;D:\sentinela\backend;D:\sentinela\backend\api\services\analytics")

from backend.api.services.analytics.nota_tecnica_quadros import (
    _add_tabela_repasses_anuais,
)


def test_sem_repasses():
    print("TESTE: farmácia SEM repasses")
    doc = Document()
    repasses_ctx = {
        "rows": [],
        "total": 0.0,
        "periodo_fmt": "2020 a 2024",
        "sem_repasses": True,
    }
    _add_tabela_repasses_anuais(
        doc,
        "FARMACIA TESTE LTDA",
        "12.345.678/0001-90",
        repasses_ctx,
        tabela_num=5,
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    table = doc2.tables[0]
    assert len(table.rows) == 3, f"esperado 3 linhas, obtido {len(table.rows)}"

    info_text = table.rows[1].cells[0].text
    assert "Não foram identificadas" in info_text
    assert table.rows[1].cells[0].text == table.rows[1].cells[1].text == table.rows[1].cells[2].text

    total_val = table.rows[-1].cells[2].text
    assert "0,00" in total_val

    last_para_text = doc2.paragraphs[-1].text
    assert "Não foram identificadas ordens bancárias no período analisado" in last_para_text
    assert "possível repasse dos recursos para o CNPJ da Matriz" in last_para_text
    assert "ausência de pagamentos, glosas ou defasagem temporal" not in last_para_text
    assert "foi consultada a existência" not in last_para_text
    print("  PASS")


def test_com_repasses():
    print("TESTE: farmácia COM repasses")
    doc = Document()
    repasses_ctx = {
        "rows": [
            {"ano": 2020, "valor": 12345.67},
            {"ano": 2021, "valor": 23456.78},
            {"ano": 2022, "valor": 34567.89},
        ],
        "total": 70370.34,
        "periodo_fmt": "2020 a 2022",
        "sem_repasses": False,
    }
    _add_tabela_repasses_anuais(
        doc,
        "FARMACIA TESTE LTDA",
        "12.345.678/0001-90",
        repasses_ctx,
        tabela_num=5,
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    table = doc2.tables[0]
    assert len(table.rows) == 5
    assert "70.370,34" in table.rows[-1].cells[2].text

    last_para_text = doc2.paragraphs[-1].text
    assert "Como o Sistema Sentinela se concentra apenas nos medicamentos" in last_para_text
    assert "glosa" in last_para_text.lower()
    assert "CNPJ da Matriz" not in last_para_text
    print("  PASS")


def test_periodo_fmts():
    print("TESTE: 5 variações de periodo_fmt")
    for fmt in ["2020 a 2024", "2024", "a partir de 2020", "até 2024", "período analisado"]:
        ctx = {"rows": [], "total": 0.0, "periodo_fmt": fmt, "sem_repasses": True}
        doc = Document()
        _add_tabela_repasses_anuais(doc, "F", "1", ctx, tabela_num=1)
        buf = io.BytesIO()
        doc.save(buf)
    print("  PASS")


if __name__ == "__main__":
    test_sem_repasses()
    test_com_repasses()
    test_periodo_fmts()
    print("\n3/3 testes passaram")
