from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .nota_tecnica_docx_utils import _cell_bg, _keep_small_table_together, _run, _set_table_fixed_widths
from .nota_tecnica_formatters import _format_decimal_pt, _format_list_pt
from .nota_tecnica_quadros import _add_quadro_esocial, _add_quadro_esocial_trabalhadores, _format_quadro_title


def _plural(value: int, singular: str, plural: str) -> str:
    return singular if value == 1 else plural


def _month_distance(start_period: Any, end_period: Any) -> int:
    if not hasattr(start_period, "year") or not hasattr(start_period, "month"):
        raise RuntimeError("Periodo inicial sem ano/mes para calcular alerta eSocial.")
    if not hasattr(end_period, "year") or not hasattr(end_period, "month"):
        raise RuntimeError("Periodo final sem ano/mes para calcular alerta eSocial.")

    distance = (int(end_period.year) - int(start_period.year)) * 12 + int(end_period.month) - int(start_period.month)
    if distance < 0:
        raise RuntimeError("Periodo final anterior ao periodo inicial no alerta eSocial.")
    return distance


def _capitalize_first(text: str) -> str:
    if not text:
        return text
    return text[:1].upper() + text[1:]


def _format_annual_summary(rows: list[dict[str, Any]]) -> str:
    partes: list[str] = []
    for row in rows:
        qtd_trab = int(row.get("qtd_trabalhadores_vinculo_ano") or 0)
        qtd_farm = int(row.get("qtd_farmaceuticos_vinculo_ano") or 0)
        partes.append(
            f'{row["ano_base"]} ({qtd_trab} {_plural(qtd_trab, "trabalhador com vínculo", "trabalhadores com vínculo")}, '
            f'{qtd_farm} {_plural(qtd_farm, "farmacêutico com vínculo", "farmacêuticos com vínculo")})'
        )
    return _format_list_pt(partes)


def _format_single_worker_summary(rows: list[dict[str, Any]]) -> str:
    partes: list[str] = []
    for row in rows:
        cbo_txt = row.get("trabalhador_unico_cbo_txt") or "CBO não identificado"
        admissao_txt = row.get("trabalhador_unico_dt_admissao_txt") or "—"
        if admissao_txt != "—":
            partes.append(f'{row["ano_base"]} ({cbo_txt}, admissão em {admissao_txt})')
        else:
            partes.append(f'{row["ano_base"]} ({cbo_txt})')
    return _format_list_pt(partes)


def _add_movimentacao_sem_funcionario_table(
    doc,
    *,
    ultimo_mes_ativo_txt: str,
    periodo_txt: str,
    qtd_meses_txt: str,
    qtd_meses_sem_funcionario: int,
    valor_periodo_txt: str,
    qtd_autorizacoes_periodo_txt: str,
):
    p_title = doc.add_paragraph()
    _format_quadro_title(p_title)
    _run(
        p_title,
        'Quadro 01-C - Síntese da movimentação após o último mês com trabalhador ativo',
        color='334155',
        size=8,
        bold=True,
    )

    rows = [
        ("Último mês com trabalhador ativo no eSocial", _capitalize_first(ultimo_mes_ativo_txt)),
        ("Último mês com movimentação de venda no PFPB", _capitalize_first(periodo_txt)),
        (
            "Período sem trabalhador ativo até a última movimentação",
            f'{qtd_meses_txt} {_plural(qtd_meses_sem_funcionario, "mês", "meses")}',
        ),
        ("Faturamento PFPB no período sem trabalhador ativo", f"R$ {valor_periodo_txt}"),
        ("Autorizações PFPB no período sem trabalhador ativo", qtd_autorizacoes_periodo_txt),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = 'Table Grid'
    _set_table_fixed_widths(table, [Inches(4.29), Inches(2.71)])

    headers = ("Item verificado", "Resultado")
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _cell_bg(cell, 'E2E8F0')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, header, color='0F172A', size=8.5)

    for row_idx, (label, value) in enumerate(rows, start=1):
        label_para = table.rows[row_idx].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(label_para, label, color='0F172A', size=8.5)

        value_para = table.rows[row_idx].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(value_para, value, color='0F172A', size=8.5)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)

    _keep_small_table_together(p_title, table)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)


def _add_legal_context(doc):
    p_leg = doc.add_paragraph()
    p_leg.paragraph_format.space_before = Pt(6)
    _run(
        p_leg,
        'A ausência de vínculo identificado no eSocial para trabalhador com CBO de farmacêutico indica possível fragilidade na regularidade técnico-operacional do estabelecimento, uma vez que a atividade farmacêutica exige responsabilidade e assistência técnica de profissional habilitado durante o horário de funcionamento.',
        color='0F172A',
        size=10,
    )


def _add_movimentacao_sem_funcionario_alerta(doc, esocial_comp: dict[str, Any]):
    alerta = esocial_comp.get("movimentacao_sem_funcionario_alerta")
    if not alerta:
        return

    campos_obrigatorios = {
        "ultimo_periodo_movimentacao",
        "ultimo_periodo_movimentacao_txt",
        "ultimo_mes_trabalhador_ativo",
        "ultimo_mes_trabalhador_ativo_txt",
        "valor_pfpb_periodo_sem_funcionario",
        "qtd_autorizacoes_periodo_sem_funcionario",
    }
    missing_alerta = [campo for campo in campos_obrigatorios if alerta.get(campo) is None]
    if missing_alerta:
        raise RuntimeError(
            "Contexto eSocial sem campos obrigatorios para alerta de movimentacao sem funcionario: "
            + ", ".join(sorted(missing_alerta))
        )

    periodo_txt = str(alerta["ultimo_periodo_movimentacao_txt"])
    ultimo_mes_ativo_txt = str(alerta["ultimo_mes_trabalhador_ativo_txt"])
    qtd_meses_sem_funcionario = _month_distance(
        alerta["ultimo_mes_trabalhador_ativo"],
        alerta["ultimo_periodo_movimentacao"],
    )
    qtd_meses_txt = _format_decimal_pt(float(qtd_meses_sem_funcionario), 0)
    valor_periodo_txt = _format_decimal_pt(float(alerta["valor_pfpb_periodo_sem_funcionario"]), 2)
    qtd_autorizacoes_periodo = int(alerta["qtd_autorizacoes_periodo_sem_funcionario"])
    qtd_autorizacoes_periodo_txt = _format_decimal_pt(float(qtd_autorizacoes_periodo), 0)
    ano_ultima_mov = int(alerta.get("ano_ultima_movimentacao") or 0)
    ano_ref_esocial = int(alerta.get("ano_esocial_referencia_ultima_movimentacao") or 0)
    sem_esocial_no_ano = bool(alerta.get("is_sem_esocial_no_ano_ultima_movimentacao"))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _run(
        p,
        f'Observa-se, adicionalmente, que o último mês com movimentação de venda registrada no PFPB para o estabelecimento foi {periodo_txt}, enquanto o último mês com trabalhador ativo vinculado ao estabelecimento na base do eSocial foi {ultimo_mes_ativo_txt}. ',
        color='0F172A',
        size=10,
    )
    _run(
        p,
        f'Assim, foram identificados {qtd_meses_txt} {_plural(qtd_meses_sem_funcionario, "mês", "meses")} sem trabalhador ativo no estabelecimento até a última movimentação registrada. Nesse período, a farmácia apresentou faturamento de R$ {valor_periodo_txt} e {qtd_autorizacoes_periodo_txt} {_plural(qtd_autorizacoes_periodo, "autorização", "autorizações")} no PFPB.',
        color='0F172A',
        size=10,
    )
    if sem_esocial_no_ano and ano_ultima_mov and ano_ref_esocial:
        _run(
            p,
            f' Não foram identificados vínculos no eSocial para {ano_ultima_mov}; por isso, a verificação considerou o último ano trabalhista disponível, {ano_ref_esocial}.',
            color='0F172A',
            size=10,
        )
    _add_movimentacao_sem_funcionario_table(
        doc,
        ultimo_mes_ativo_txt=ultimo_mes_ativo_txt,
        periodo_txt=periodo_txt,
        qtd_meses_txt=qtd_meses_txt,
        qtd_meses_sem_funcionario=qtd_meses_sem_funcionario,
        valor_periodo_txt=valor_periodo_txt,
        qtd_autorizacoes_periodo_txt=qtd_autorizacoes_periodo_txt,
    )


def _add_esocial_context_text(doc, razao_social: str, cnpj_fmt: str, esocial_comp: dict[str, Any]):
    """Renderiza o contexto trabalhista/eSocial na seção cadastral da Nota Técnica."""
    doc.add_heading('5.2 Vínculos trabalhistas', level=2)

    if not esocial_comp.get("has_data"):
        p = doc.add_paragraph()
        _run(
            p,
            f'Não foram identificados registros anuais de vínculos trabalhistas no eSocial para a Farmácia {razao_social} (CNPJ {cnpj_fmt}) no período analisado.',
            color='0F172A',
            size=10,
        )
        return

    rows = list(esocial_comp.get("rows") or [])
    annual_summary = _format_annual_summary(rows)
    dt_carga_txt = esocial_comp.get("dt_carga_fonte_txt") or "—"
    anos_sem_farm = list(esocial_comp.get("anos_sem_farmaceutico") or [])
    qtd_anos = len(rows)
    qtd_anos_sem_farm = len(anos_sem_farm)

    p_intro = doc.add_paragraph()

    if qtd_anos == 1:
        row = rows[0]
        ano = row["ano_base"]
        qtd_trab_vinculo = int(row.get("qtd_trabalhadores_vinculo_ano") or 0)
        qtd_farm_vinculo = int(row.get("qtd_farmaceuticos_vinculo_ano") or 0)
        qtd_trab_ativo = int(row.get("qtd_trabalhadores_ativos_competencia") or 0)
        qtd_farm_ativo = int(row.get("qtd_farmaceuticos_ativos_competencia") or 0)
        competencia_txt = row.get("competencia_txt") or "competência-base"
        _run(
            p_intro,
            f'Em consulta à base do eSocial disponível no Sentinela, atualizada até {dt_carga_txt}, foram identificados vínculos trabalhistas associados à Farmácia {razao_social} somente durante o ano de {ano}. ',
            color='0F172A',
            size=10,
        )
        if qtd_farm_vinculo > 0:
            _run(
                p_intro,
                f'Para esse exercício, foram apurados {qtd_trab_vinculo} {_plural(qtd_trab_vinculo, "trabalhador com vínculo", "trabalhadores com vínculo")}, com identificação de {qtd_farm_vinculo} {_plural(qtd_farm_vinculo, "trabalhador", "trabalhadores")} com CBO de farmacêutico (223405).',
                color='0F172A',
                size=10,
            )
        else:
            _run(
                p_intro,
                f'Para esse exercício, foram apurados {qtd_trab_vinculo} {_plural(qtd_trab_vinculo, "trabalhador com vínculo", "trabalhadores com vínculo")}, ',
                color='0F172A',
                size=10,
            )
            _run(
                p_intro,
                'nenhum deles com CBO de farmacêutico',
                color='0F172A',
                size=10,
                underline=True,
            )
            _run(
                p_intro,
                ' (223405).',
                color='0F172A',
                size=10,
            )
            _add_legal_context(doc)
        if qtd_trab_ativo > 0 and (qtd_trab_ativo != qtd_trab_vinculo or qtd_farm_ativo != qtd_farm_vinculo):
            p_ativos = doc.add_paragraph()
            p_ativos.paragraph_format.space_before = Pt(6)
            _run(
                p_ativos,
                f'Na competência-base de {competencia_txt}, permaneciam ativos {qtd_trab_ativo} {_plural(qtd_trab_ativo, "trabalhador", "trabalhadores")}, sendo {qtd_farm_ativo} com CBO de farmacêutico (223405).',
                color='0F172A',
                size=10,
            )
    else:
        periodo_anos_txt = esocial_comp.get("periodo_anos_txt") or "período analisado"
        _run(
            p_intro,
            f'Em consulta à base do eSocial disponível no Sentinela, atualizada até {dt_carga_txt}, foram identificados vínculos trabalhistas associados à Farmácia {razao_social} entre {periodo_anos_txt}. ',
            color='0F172A',
            size=10,
        )
        _run(p_intro, 'No período, os registros indicam a seguinte composição anual: ', color='0F172A', size=10)
        _run(p_intro, annual_summary, color='334155', size=10, bold=True)
        _run(p_intro, '.', color='0F172A', size=10)

        if qtd_anos_sem_farm == qtd_anos:
            p_status = doc.add_paragraph()
            p_status.paragraph_format.space_before = Pt(6)
            _run(
                p_status,
                'Em nenhum dos anos considerados foi identificado trabalhador com CBO de farmacêutico (223405).',
                color='0F172A',
                size=10,
            )
            _add_legal_context(doc)
        elif anos_sem_farm:
            anos_txt = _format_list_pt([str(row["ano_base"]) for row in anos_sem_farm])
            p_status = doc.add_paragraph()
            p_status.paragraph_format.space_before = Pt(6)
            _run(
                p_status,
                f'Nos anos de {anos_txt}, não foi identificado trabalhador com CBO de farmacêutico (223405).',
                color='0F172A',
                size=10,
            )
            _add_legal_context(doc)
        else:
            p_status = doc.add_paragraph()
            p_status.paragraph_format.space_before = Pt(6)
            _run(
                p_status,
                'Em todos os anos considerados, houve identificação de ao menos um trabalhador com CBO de farmacêutico (223405). A informação é apresentada como elemento de contexto sobre a força de trabalho registrada para o estabelecimento.',
                color='0F172A',
                size=10,
            )

    anos_um_trabalhador = list(esocial_comp.get("anos_um_trabalhador") or [])
    if anos_um_trabalhador:
        p_unico = doc.add_paragraph()
        p_unico.paragraph_format.space_before = Pt(6)
        _run(
            p_unico,
            'Também se observa que, em ',
            color='0F172A',
            size=10,
        )
        _run(
            p_unico,
            _format_single_worker_summary(anos_um_trabalhador),
            color='334155',
            size=10,
            bold=True,
        )
        _run(
            p_unico,
            ', havia registro anual de apenas um trabalhador vinculado ao estabelecimento.',
            color='0F172A',
            size=10,
        )

    _add_quadro_esocial(doc, razao_social, cnpj_fmt, esocial_comp)
    _add_quadro_esocial_trabalhadores(doc, razao_social, cnpj_fmt, esocial_comp)
    _add_movimentacao_sem_funcionario_alerta(doc, esocial_comp)
